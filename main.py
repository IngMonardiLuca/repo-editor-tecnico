import json
import math
import os
import platform
import re
import shutil
import sys
import tempfile
import traceback
import urllib.parse
import uuid
import pyphen
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_RIGHT, TA_LEFT
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Image as RLImage, Paragraph, Spacer, Table as RLTable, TableStyle, PageBreak, BaseDocTemplate, Frame, PageTemplate
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.platypus.doctemplate import ActionFlowable
from PySide6.QtCore import Qt, QSettings, QStandardPaths, QTimer, QObject, QThread, Signal, Slot, QPropertyAnimation, QUrl, QEvent,QEasingCurve
from PySide6.QtGui import QAction, QKeySequence, QPixmap, QShortcut, QTextListFormat, QTextDocument, QTextCharFormat, QTextCursor, QDesktopServices, QImage,QCursor
from PySide6.QtWidgets import (
    QApplication,
    QDialog,
    QDockWidget,
    QComboBox,
    QFileDialog,
    QFrame,
    QHBoxLayout,
    QInputDialog,
    QLabel,
    QLineEdit,
    QListWidget,
    QListWidgetItem,
    QMainWindow,
    QMenu,
    QMessageBox,
    QPushButton,
    QSizePolicy,
    QScrollArea,
    QSplitter,
    QStatusBar,
    QTableWidget,
    QTableWidgetItem,
    QTextEdit,
    QToolBar,
    QTreeWidget,
    QTreeWidgetItem,
    QTabWidget,
    QVBoxLayout,
    QWidget,
    QHeaderView,
    QAbstractScrollArea,
    QRadioButton,
    QGroupBox,
    QGridLayout,
    QFormLayout,
    QSpinBox,
    QAbstractItemView
)

import matplotlib
matplotlib.use("Agg")
from matplotlib.figure import Figure
from matplotlib.backends.backend_agg import FigureCanvasAgg

TITLE_ROLE = Qt.UserRole
CONTENT_ROLE = Qt.UserRole + 1


def new_block_id() -> str:
    return uuid.uuid4().hex[:10]


def format_reference_tokens(tokens):
    tokens = [t for t in tokens if t]
    if not tokens:
        return ""
    if len(tokens) == 1:
        return tokens[0]
    if len(tokens) == 2:
        return f"{tokens[0]} e {tokens[1]}"
    return ", ".join(tokens[:-1]) + f" e {tokens[-1]}"


def as_posix_path(value: str) -> str:
    if not value:
        return ""
    return str(value).replace("\\", "/")

def validate_folder_name(value: str):
    value = value.strip()

    if not value:
        return False, "Il nome non può essere vuoto."

    invalid_chars = r'\/:*?"<>|'

    for ch in invalid_chars:
        if ch in value:
            return False, f"Il nome contiene un carattere non valido: {ch}"

    if value.endswith(".") or value.endswith(" "):
        return False, "Il nome non può terminare con punto o spazio."

    reserved_names = {
        "CON", "PRN", "AUX", "NUL",
        "COM1", "COM2", "COM3", "COM4", "COM5", "COM6", "COM7", "COM8", "COM9",
        "LPT1", "LPT2", "LPT3", "LPT4", "LPT5", "LPT6", "LPT7", "LPT8", "LPT9",
    }

    if value.upper() in reserved_names:
        return False, "Il nome è riservato dal sistema operativo."

    return True, ""

class CaptionLineEdit(QLineEdit):
    def __init__(self, owner_dialog, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.owner_dialog = owner_dialog

    def keyPressEvent(self, event):
        if event.key() in (Qt.Key_Return, Qt.Key_Enter):
            self.clearFocus()
            self.owner_dialog.setFocus()
            event.accept()
            return
        super().keyPressEvent(event)


class ImagePreviewDialog(QDialog):
    def __init__(self, image_path: str, caption: str = "", parent=None):
        super().__init__(parent)
        self.setWindowTitle("Anteprima immagine")
        self.resize(900, 760)
        self.caption_value = caption

        layout = QVBoxLayout(self)

        self.image_label = QLabel()
        self.image_label.setAlignment(Qt.AlignCenter)
        self.image_label.setStyleSheet("background: #202020; border: 1px solid #444;")
        layout.addWidget(self.image_label)

        self.caption_edit = CaptionLineEdit(self)
        self.caption_edit.setPlaceholderText("Didascalia immagine...")
        self.caption_edit.setText(caption)
        self.caption_edit.setStyleSheet(
            "QLineEdit { background: white; color: black; border: 1px solid #ccc; padding: 6px; }"
        )
        self.caption_edit.textChanged.connect(self.on_caption_changed)
        layout.addWidget(self.caption_edit)

        pixmap = QPixmap(image_path)
        if pixmap.isNull():
            self.image_label.setText("Impossibile caricare l'immagine")
            self.image_label.setStyleSheet("color: white; background: #202020; border: 1px solid #444;")
        else:
            scaled = pixmap.scaled(860, 640, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            self.image_label.setPixmap(scaled)

        self._apply_focus_rule()

    def on_caption_changed(self, text):
        self.caption_value = text

    def _apply_focus_rule(self):
        if not self.caption_edit.text().strip():
            self.caption_edit.setFocus()
        else:
            self.setFocus()

    def showEvent(self, event):
        super().showEvent(event)
        self.activateWindow()
        self._apply_focus_rule()


class MultiImagePreviewDialog(QDialog):
    def __init__(self, images_data, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Anteprima gruppo immagini")
        self.resize(980, 760)

        self.images_data = images_data if images_data else []
        self.current_index = 0
        self.setFocusPolicy(Qt.StrongFocus)

        root = QVBoxLayout(self)

        self.info_label = QLabel("")
        self.info_label.setAlignment(Qt.AlignCenter)
        self.info_label.setStyleSheet("color: black;")
        root.addWidget(self.info_label)

        self.image_label = QLabel()
        self.image_label.setAlignment(Qt.AlignCenter)
        self.image_label.setStyleSheet("background: #202020; border: 1px solid #444; color: white;")
        self.image_label.setMinimumHeight(560)
        self.image_label.setFocusPolicy(Qt.NoFocus)
        root.addWidget(self.image_label)

        self.caption_edit = CaptionLineEdit(self)
        self.caption_edit.setPlaceholderText("Didascalia immagine corrente...")
        self.caption_edit.setStyleSheet(
            "QLineEdit { background: white; color: black; border: 1px solid #ccc; padding: 6px; }"
        )
        self.caption_edit.textChanged.connect(self.on_caption_changed)
        root.addWidget(self.caption_edit)

        hint_label = QLabel("Usa ← e → per sfogliare le immagini. Premi Enter per uscire dalla didascalia.")
        hint_label.setAlignment(Qt.AlignCenter)
        hint_label.setStyleSheet("color: #666;")
        hint_label.setFocusPolicy(Qt.NoFocus)
        root.addWidget(hint_label)

        self.shortcut_left = QShortcut(QKeySequence(Qt.Key_Left), self)
        self.shortcut_right = QShortcut(QKeySequence(Qt.Key_Right), self)
        self.shortcut_left.activated.connect(self.show_prev)
        self.shortcut_right.activated.connect(self.show_next)

        self.refresh()
        self.activateWindow()
        self.setFocus()

    def current_image(self):
        if not self.images_data:
            return None
        if self.current_index < 0 or self.current_index >= len(self.images_data):
            self.current_index = 0
        return self.images_data[self.current_index]

    def _apply_focus_rule(self):
        item = self.current_image()
        if item is None:
            self.setFocus()
            return
        if not item.get("caption", "").strip():
            self.caption_edit.setFocus()
        else:
            self.setFocus()

    def refresh(self):
        total = len(self.images_data)
        item = self.current_image()

        if item is None:
            self.info_label.setText("Nessuna immagine")
            self.image_label.setPixmap(QPixmap())
            self.image_label.setText("Nessuna immagine")
            self.caption_edit.blockSignals(True)
            self.caption_edit.clear()
            self.caption_edit.blockSignals(False)
            self.caption_edit.setEnabled(False)
            return

        self.info_label.setText(f"Immagine {self.current_index + 1} di {total}")

        pixmap = QPixmap(item.get("path", ""))
        if pixmap.isNull():
            self.image_label.setPixmap(QPixmap())
            self.image_label.setText("Impossibile caricare l'immagine")
        else:
            scaled = pixmap.scaled(920, 560, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            self.image_label.setPixmap(scaled)
            self.image_label.setText("")

        self.caption_edit.setEnabled(True)
        self.caption_edit.blockSignals(True)
        self.caption_edit.setText(item.get("caption", ""))
        self.caption_edit.blockSignals(False)
        self._apply_focus_rule()

    def show_prev(self):
        if len(self.images_data) <= 1:
            return
        self.current_index = (self.current_index - 1) % len(self.images_data)
        self.refresh()

    def show_next(self):
        if len(self.images_data) <= 1:
            return
        self.current_index = (self.current_index + 1) % len(self.images_data)
        self.refresh()

    def on_caption_changed(self, text):
        item = self.current_image()
        if item is not None:
            item["caption"] = text

    def showEvent(self, event):
        super().showEvent(event)
        self.activateWindow()
        self._apply_focus_rule()


class MinimalBlockFrame(QFrame):
    def __init__(self):
        super().__init__()
        self.setFrameShape(QFrame.NoFrame)
        self.setStyleSheet("QFrame { background: #f7f7f7; border: 1px solid #e6e6e6; border-radius: 8px; }")


def style_small_button(button):
    button.setFixedHeight(26)
    button.setStyleSheet(
        "QPushButton { border: 1px solid #d6d6d6; border-radius: 4px; background: #fafafa; padding: 0 8px; }"
        "QPushButton:hover { background: #f0f0f0; }"
    )


def style_icon_button(button):
    button.setFixedSize(28, 24)
    button.setStyleSheet(
        "QPushButton { border: 1px solid #d6d6d6; border-radius: 4px; background: #fafafa; }"
        "QPushButton:hover { background: #f0f0f0; }"
    )


class BaseBlockWidget(MinimalBlockFrame):
    block_type_label = "Blocco"

    def __init__(self, parent_editor, data=None):
        super().__init__()
        self.parent_editor = parent_editor
        self.data = self.parent_editor.main_window.ensure_block_defaults(deepcopy(data) if data else {})
        self.collapsed = self.data.get("collapsed", False)

        self.root_layout = QVBoxLayout(self)
        self.root_layout.setContentsMargins(8, 8, 8, 8)
        self.root_layout.setSpacing(6)

        self.header_row = QHBoxLayout()
        self.header_row.setContentsMargins(0, 0, 0, 0)

        self.btn_toggle = QPushButton("▶" if self.collapsed else "▼")
        style_icon_button(self.btn_toggle)
        self.btn_toggle.clicked.connect(self.toggle_collapsed)
        self.header_row.addWidget(self.btn_toggle)

        self.summary_label = QLabel("")
        self.summary_label.setStyleSheet("color: #555;")
        self.header_row.addWidget(self.summary_label, 1)

        self.btn_up = QPushButton("↑")
        self.btn_down = QPushButton("↓")
        self.btn_duplicate = QPushButton("⧉")
        self.btn_delete = QPushButton("✕")
        for btn in (self.btn_up, self.btn_down, self.btn_duplicate, self.btn_delete):
            style_icon_button(btn)

        self.btn_up.clicked.connect(lambda: self.parent_editor.move_block(self, -1))
        self.btn_down.clicked.connect(lambda: self.parent_editor.move_block(self, 1))
        self.btn_duplicate.clicked.connect(lambda: self.parent_editor.duplicate_block(self))
        self.btn_delete.clicked.connect(lambda: self.parent_editor.remove_block(self))

        self.header_row.addWidget(self.btn_up)
        self.header_row.addWidget(self.btn_down)
        self.header_row.addWidget(self.btn_duplicate)
        self.header_row.addWidget(self.btn_delete)
        self.root_layout.addLayout(self.header_row)

        self.content_widget = QWidget()
        self.content_layout = QVBoxLayout(self.content_widget)
        self.content_layout.setContentsMargins(0, 0, 0, 0)
        self.content_layout.setSpacing(6)
        self.root_layout.addWidget(self.content_widget)
        self.content_widget.setVisible(not self.collapsed)

    def toggle_collapsed(self):
        self.collapsed = not self.collapsed
        self.content_widget.setVisible(not self.collapsed)
        self.btn_toggle.setText("▶" if self.collapsed else "▼")
        self.update_summary()
        self.parent_editor.main_window.mark_dirty()

    def update_summary(self):
        self.summary_label.setText(self.build_summary())

    def build_summary(self):
        return self.block_type_label

    def set_collapsed_state(self, collapsed):
        self.collapsed = collapsed
        self.content_widget.setVisible(not self.collapsed)
        self.btn_toggle.setText("▶" if self.collapsed else "▼")
        self.update_summary()

    def export_common_data(self):
        return {
            "id": self.data.get("id"),
            "meta": deepcopy(self.data.get("meta", {})),
            "collapsed": self.collapsed,
        }

class TextBlockWidget(BaseBlockWidget):
    block_type_label = "Testo"

    def __init__(self, parent_editor, data=None):
        base = {"type": "text", "text": "", "html": "", "collapsed": False}
        base.update(data or {})
        super().__init__(parent_editor, base)

        self.toolbar_widget = QWidget()
        toolbar_row = QHBoxLayout(self.toolbar_widget)
        toolbar_row.setContentsMargins(0, 0, 0, 0)

        self.btn_bold = QPushButton("B")
        self.btn_italic = QPushButton("I")
        self.btn_underline = QPushButton("U")
        self.btn_superscript = QPushButton("x²")
        self.btn_subscript = QPushButton("x₂")
        self.btn_bullets = QPushButton("•")
        self.btn_numbers = QPushButton("1.")
        self.btn_justify = QPushButton("J")
        self.btn_center = QPushButton("C")
        self.btn_insert_attachment = QPushButton("Allegato")
        self.btn_ai = QPushButton("AI")
        self.btn_focus = QPushButton("⛶")

        self.btn_bold.setToolTip("Grassetto")
        self.btn_italic.setToolTip("Corsivo")
        self.btn_underline.setToolTip("Sottolineato")
        self.btn_superscript.setToolTip("Apice")
        self.btn_subscript.setToolTip("Pedice")
        self.btn_bullets.setToolTip("Elenco puntato")
        self.btn_numbers.setToolTip("Elenco numerato")
        self.btn_justify.setToolTip("Giustifica")
        self.btn_center.setToolTip("Centra")
        self.btn_insert_attachment.setToolTip("Inserisci allegato")
        self.btn_ai.setToolTip("Assistente AI")
        self.btn_focus.setToolTip("Editor esteso")

        self.font_size_combo = QComboBox()
        self.font_size_combo.addItems(["9", "10", "11", "12", "14", "16", "18"])
        self.font_size_combo.setCurrentText("11")
        self.font_size_combo.setFixedWidth(70)

        style_icon_button(self.btn_bold)
        toolbar_row.addWidget(self.btn_bold)

        style_icon_button(self.btn_italic)
        toolbar_row.addWidget(self.btn_italic)

        style_icon_button(self.btn_underline)
        toolbar_row.addWidget(self.btn_underline)

        self.btn_superscript.setFixedSize(28, 24)
        self.btn_superscript.setStyleSheet(
            "QPushButton { border: 1px solid #d6d6d6; border-radius: 4px; background: #fafafa; }"
            "QPushButton:hover { background: #f0f0f0; }"
            "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
        )
        toolbar_row.addWidget(self.btn_superscript)

        self.btn_subscript.setFixedSize(28, 24)
        self.btn_subscript.setStyleSheet(
            "QPushButton { border: 1px solid #d6d6d6; border-radius: 4px; background: #fafafa; }"
            "QPushButton:hover { background: #f0f0f0; }"
            "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
        )
        toolbar_row.addWidget(self.btn_subscript)

        self.btn_bullets.setFixedSize(28, 24)
        self.btn_bullets.setStyleSheet(
            "QPushButton { border: 1px solid #d6d6d6; border-radius: 4px; background: #fafafa; }"
            "QPushButton:hover { background: #f0f0f0; }"
            "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
        )
        toolbar_row.addWidget(self.btn_bullets)

        self.btn_numbers.setFixedSize(28, 24)
        self.btn_numbers.setStyleSheet(
            "QPushButton { border: 1px solid #d6d6d6; border-radius: 4px; background: #fafafa; }"
            "QPushButton:hover { background: #f0f0f0; }"
            "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
        )
        toolbar_row.addWidget(self.btn_numbers)

        self.btn_justify.setFixedSize(28, 24)
        self.btn_justify.setStyleSheet(
            "QPushButton { border: 1px solid #d6d6d6; border-radius: 4px; background: #fafafa; }"
            "QPushButton:hover { background: #f0f0f0; }"
            "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
        )
        toolbar_row.addWidget(self.btn_justify)

        self.btn_center.setFixedSize(28, 24)
        self.btn_center.setStyleSheet(
            "QPushButton { border: 1px solid #d6d6d6; border-radius: 4px; background: #fafafa; }"
            "QPushButton:hover { background: #f0f0f0; }"
            "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
        )
        toolbar_row.addWidget(self.btn_center)

        toolbar_row.addWidget(QLabel("Dim."))
        toolbar_row.addWidget(self.font_size_combo)

        self.btn_insert_attachment.setFixedWidth(90)
        self.btn_ai.setFixedWidth(40)
        self.btn_focus.setFixedWidth(40)

        self.btn_insert_attachment.setStyleSheet(
            "QPushButton { border: 1px solid #d6d6d6; border-radius: 4px; background: #fafafa; padding: 0 8px; }"
            "QPushButton:hover { background: #f0f0f0; }"
            "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
        )

        self.btn_ai.setStyleSheet(
            "QPushButton { border: 1px solid #d6d6d6; border-radius: 4px; background: #fafafa; }"
            "QPushButton:hover { background: #f0f0f0; }"
            "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
        )

        self.btn_focus.setStyleSheet(
            "QPushButton { border: 1px solid #d6d6d6; border-radius: 4px; background: #fafafa; }"
            "QPushButton:hover { background: #f0f0f0; }"
            "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
        )

        toolbar_row.addWidget(self.btn_insert_attachment)
        toolbar_row.addWidget(self.btn_ai)
        toolbar_row.addWidget(self.btn_focus)
        self.btn_ai.clicked.connect(self.on_ai_clicked)
        self.btn_focus.clicked.connect(self.open_focus_editor)

        toolbar_row.addStretch()
        self.content_layout.addWidget(self.toolbar_widget)

        self.editor = QTextEdit()
        self.editor.setContextMenuPolicy(Qt.DefaultContextMenu)
        self.editor.setStyleSheet("""
            QTextEdit {
                border: none;
                background: white;
                color: black;
                font-size: 14px;
            }

            QMenu {
                background-color: #f2f2f2;
                color: black;
                border: 1px solid #c8c8c8;
            }

            QMenu::item {
                background-color: transparent;
                padding: 6px 24px 6px 24px;
            }

            QMenu::item:selected {
                background-color: #dcdcdc;
                color: black;
            }
        """)
        self.editor.setAlignment(Qt.AlignJustify)
        self.editor.setPlaceholderText("Scrivi qui... Usa [[FIG:ID]] per i riferimenti alle immagini.")
        self.editor.setMinimumHeight(300)
        self.editor.setAcceptRichText(True)
        self.editor.setTabChangesFocus(False)
        self.editor.installEventFilter(self)
        self.editor.viewport().installEventFilter(self)

        html = self.data.get("html", "")
        text = self.data.get("text", "")

        if html:
            self.editor.setHtml(html)
        else:
            self.editor.setPlainText(text)

        self.btn_bold.clicked.connect(self.toggle_bold)
        self.btn_italic.clicked.connect(self.toggle_italic)
        self.btn_underline.clicked.connect(self.toggle_underline)
        self.btn_superscript.clicked.connect(self.toggle_superscript)
        self.btn_subscript.clicked.connect(self.toggle_subscript)
        self.btn_bullets.clicked.connect(self.toggle_bullet_list)
        self.btn_numbers.clicked.connect(self.toggle_numbered_list)
        self.btn_justify.clicked.connect(self.apply_justify)
        self.btn_center.clicked.connect(self.apply_center)
        self.btn_insert_attachment.clicked.connect(self.insert_attachment_reference)
        self.font_size_combo.currentTextChanged.connect(self.apply_font_size)

        self.editor.textChanged.connect(self.on_text_changed)
        self.editor.cursorPositionChanged.connect(self.update_format_buttons)

        self.content_layout.addWidget(self.editor)

        self.shortcut_focus = QShortcut(QKeySequence("Ctrl+Shift+F"), self)
        self.shortcut_focus.setContext(Qt.WidgetWithChildrenShortcut)
        self.shortcut_focus.activated.connect(self.open_focus_editor)

        self.shortcut_superscript = QShortcut(QKeySequence("Ctrl+."), self)
        self.shortcut_superscript.setContext(Qt.WidgetWithChildrenShortcut)
        self.shortcut_superscript.activated.connect(self.toggle_superscript)

        self.shortcut_subscript = QShortcut(QKeySequence("Ctrl+,"), self)
        self.shortcut_subscript.setContext(Qt.WidgetWithChildrenShortcut)
        self.shortcut_subscript.activated.connect(self.toggle_subscript)

        self.shortcut_bullets = QShortcut(QKeySequence("Ctrl+Shift+7"), self)
        self.shortcut_bullets.setContext(Qt.WidgetWithChildrenShortcut)
        self.shortcut_bullets.activated.connect(self.toggle_bullet_list)

        self.shortcut_numbers = QShortcut(QKeySequence("Ctrl+Shift+8"), self)
        self.shortcut_numbers.setContext(Qt.WidgetWithChildrenShortcut)
        self.shortcut_numbers.activated.connect(self.toggle_numbered_list)

        self.shortcut_justify = QShortcut(QKeySequence("Ctrl+J"), self)
        self.shortcut_justify.setContext(Qt.WidgetWithChildrenShortcut)
        self.shortcut_justify.activated.connect(self.apply_justify)

        self.shortcut_center = QShortcut(QKeySequence("Ctrl+E"), self)
        self.shortcut_center.setContext(Qt.WidgetWithChildrenShortcut)
        self.shortcut_center.activated.connect(self.apply_center)

        self.shortcut_attachment = QShortcut(QKeySequence("Ctrl+Shift+L"), self)
        self.shortcut_attachment.setContext(Qt.WidgetWithChildrenShortcut)
        self.shortcut_attachment.activated.connect(self.insert_attachment_reference)

        self.update_summary()
        self.update_format_buttons()


    def open_focus_editor(self):
        dialog = QDialog(self.window())
        dialog.setWindowTitle("Editor Tecnico (designed by Ing. Luca Monardi)")

        settings = self.parent_editor.main_window.settings
        saved_geometry = settings.value("focus_mode/geometry")

        if saved_geometry:
            dialog.restoreGeometry(saved_geometry)
        else:
            dialog.resize(1000, 720)

        dialog.setWindowOpacity(0.0)

        self.focus_anim = QPropertyAnimation(dialog, b"windowOpacity")
        self.focus_anim.setDuration(160)
        self.focus_anim.setStartValue(0.0)
        self.focus_anim.setEndValue(1.0)

        layout = QVBoxLayout(dialog)

        # rimuove temporaneamente toolbar ed editor dal blocco
        self.content_layout.removeWidget(self.toolbar_widget)
        self.content_layout.removeWidget(self.editor)

        layout.addWidget(self.toolbar_widget)
        layout.addWidget(self.editor)
        buttons_row = QHBoxLayout()
        buttons_row.addStretch()

        btn_close = QPushButton("Chiudi")
        btn_close.setFixedWidth(120)
        btn_close.clicked.connect(dialog.accept)

        buttons_row.addWidget(btn_close)

        layout.addLayout(buttons_row)

        old_min_height = self.editor.minimumHeight()
        self.editor.setMinimumHeight(560)
        self.btn_focus.hide()

        try:
            self.editor.setFocus()
            self.focus_anim.start()
            dialog.exec()
        finally:
            settings.setValue("focus_mode/geometry", dialog.saveGeometry())

            layout.removeWidget(self.toolbar_widget)
            layout.removeWidget(self.editor)

            self.content_layout.insertWidget(0, self.toolbar_widget)
            self.content_layout.insertWidget(1, self.editor)

            self.editor.setMinimumHeight(old_min_height)
            self.editor.setFocus()

            self.btn_focus.show()

            self.update_summary()
            self.parent_editor.main_window.mark_dirty()    

    def on_text_changed(self):
        self.update_summary()
        self.parent_editor.main_window.mark_dirty()

    def update_format_buttons(self):
        fmt = self.editor.currentCharFormat()

        self.btn_bold.setStyleSheet(
            (
                "QPushButton { font-weight: bold; border: 1px solid #999; background: #e8e8e8; }"
                "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
            )
            if fmt.fontWeight() >= 700 else
            (
                "QPushButton { font-weight: bold; border: 1px solid #d6d6d6; background: #fafafa; }"
                "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
            )
        )

        self.btn_italic.setStyleSheet(
            (
                "QPushButton { font-style: italic; border: 1px solid #999; background: #e8e8e8; }"
                "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
            )
            if fmt.fontItalic() else
            (
                "QPushButton { font-style: italic; border: 1px solid #d6d6d6; background: #fafafa; }"
                "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
            )
        )

        self.btn_underline.setStyleSheet(
            (
                "QPushButton { text-decoration: underline; border: 1px solid #999; background: #e8e8e8; }"
                "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
            )
            if fmt.fontUnderline() else
            (
                "QPushButton { text-decoration: underline; border: 1px solid #d6d6d6; background: #fafafa; }"
                "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
            )
        )

        size = fmt.fontPointSize()
        if size > 0:
            rounded_size = str(int(round(size)))
            if rounded_size in [self.font_size_combo.itemText(i) for i in range(self.font_size_combo.count())]:
                self.font_size_combo.blockSignals(True)
                self.font_size_combo.setCurrentText(rounded_size)
                self.font_size_combo.blockSignals(False)

    def toggle_bold(self):
        fmt = self.editor.currentCharFormat()
        fmt.setFontWeight(400 if fmt.fontWeight() >= 700 else 700)
        self.editor.mergeCurrentCharFormat(fmt)
        self.editor.setFocus()
        self.parent_editor.main_window.mark_dirty()

    def toggle_italic(self):
        fmt = self.editor.currentCharFormat()
        fmt.setFontItalic(not fmt.fontItalic())
        self.editor.mergeCurrentCharFormat(fmt)
        self.editor.setFocus()
        self.parent_editor.main_window.mark_dirty()

    def toggle_underline(self):
        fmt = self.editor.currentCharFormat()
        fmt.setFontUnderline(not fmt.fontUnderline())
        self.editor.mergeCurrentCharFormat(fmt)
        self.editor.setFocus()
        self.parent_editor.main_window.mark_dirty()

    def toggle_superscript(self):
        fmt = self.editor.currentCharFormat()

        if fmt.verticalAlignment() == QTextCharFormat.AlignSuperScript:
            fmt.setVerticalAlignment(QTextCharFormat.AlignNormal)
        else:
            fmt.setVerticalAlignment(QTextCharFormat.AlignSuperScript)

        self.editor.mergeCurrentCharFormat(fmt)
        self.editor.setFocus()
        self.parent_editor.main_window.mark_dirty()


    def toggle_subscript(self):
        fmt = self.editor.currentCharFormat()

        if fmt.verticalAlignment() == QTextCharFormat.AlignSubScript:
            fmt.setVerticalAlignment(QTextCharFormat.AlignNormal)
        else:
            fmt.setVerticalAlignment(QTextCharFormat.AlignSubScript)

        self.editor.mergeCurrentCharFormat(fmt)
        self.editor.setFocus()
        self.parent_editor.main_window.mark_dirty()

    def apply_font_size(self, value):
        try:
            size = float(value)
        except ValueError:
            return

        fmt = self.editor.currentCharFormat()
        fmt.setFontPointSize(size)
        self.editor.mergeCurrentCharFormat(fmt)
        self.editor.setFocus()
        self.parent_editor.main_window.mark_dirty()

    def apply_justify(self):
        self.editor.setAlignment(Qt.AlignJustify)
        self.editor.setFocus()
        self.parent_editor.main_window.mark_dirty()

    def apply_center(self):
        self.editor.setAlignment(Qt.AlignCenter)
        self.editor.setFocus()
        self.parent_editor.main_window.mark_dirty()

    def toggle_bullet_list(self):
        cursor = self.editor.textCursor()
        cursor.beginEditBlock()
        cursor.createList(QTextListFormat.ListDisc)
        cursor.endEditBlock()
        self.editor.setTextCursor(cursor)
        self.editor.setFocus()
        self.parent_editor.main_window.mark_dirty()

    def toggle_numbered_list(self):
        cursor = self.editor.textCursor()
        cursor.beginEditBlock()
        cursor.createList(QTextListFormat.ListDecimal)
        cursor.endEditBlock()
        self.editor.setTextCursor(cursor)
        self.editor.setFocus()
        self.parent_editor.main_window.mark_dirty()

    def insert_attachment_reference(self):
        self.parent_editor.main_window.insert_attachment_reference_in_text(self.editor)

    def eventFilter(self, obj, event):
        if obj in (self.editor, self.editor.viewport()) and event.type() == event.Type.KeyPress:
            if event.matches(QKeySequence.Paste):
                clipboard = QApplication.clipboard()
                if not clipboard.image().isNull():
                    image_path = self.parent_editor.main_window.save_clipboard_image_to_assets()
                    if image_path:
                        self.parent_editor.insert_image_block_after(self, image_path)
                        return True

            if event.key() == Qt.Key_Tab:
                self.change_list_indent(+1)
                return True

            if event.key() == Qt.Key_Backtab:
                self.change_list_indent(-1)
                return True

        return super().eventFilter(obj, event)


    def change_list_indent(self, direction):
        cursor = self.editor.textCursor()
        current_list = cursor.currentList()

        if current_list is None:
            if direction > 0:
                cursor.insertText("    ")
            return

        fmt = current_list.format()
        current_indent = fmt.indent()
        new_indent = max(1, current_indent + direction)

        fmt.setIndent(new_indent)
        cursor.createList(fmt)

        self.editor.setTextCursor(cursor)
        self.editor.setFocus()
        self.parent_editor.main_window.mark_dirty()

    def build_summary(self):
        text = self.editor.toPlainText().strip()
        if text:
            return f"{self.block_type_label} — {text.splitlines()[0][:60]}"
        return self.block_type_label

    def export_data(self):
        data = {
            "type": "text",
            "text": self.editor.toPlainText(),
            "html": self.editor.toHtml(),
        }
        data.update(self.export_common_data())
        return data

    def on_ai_clicked(self):
        dialog = QDialog(self.window())
        dialog.setWindowTitle("AI blocco testo")
        dialog.resize(760, 520)

        layout = QVBoxLayout(dialog)
        main_window = self.window()
        ai_provider, saved_api_key, saved_model, ai_reasoning = main_window.get_active_ai_config()

        # --- selezione testo ---
        cursor = self.editor.textCursor()
        if cursor.hasSelection():
            original_text = cursor.selectedText()
        else:
            original_text = self.editor.toPlainText()

        # --- layout orizzontale testi ---
        texts_row = QHBoxLayout()

        # testo originale (solo lettura)
        original_edit = QTextEdit()
        original_edit.setPlainText(original_text)
        original_edit.setReadOnly(True)

        # testo modificato
        modified_edit = QTextEdit()

        texts_row.addWidget(original_edit)
        texts_row.addWidget(modified_edit)

        layout.addLayout(texts_row)
        
        # --- prompt ---
        layout.addWidget(QLabel("Prompt:"))

        prompt_edit = QTextEdit()
        prompt_edit.setMaximumHeight(80)
        layout.addWidget(prompt_edit)

        # --- preset prompt ---
        preset_row = QHBoxLayout()

        btn_improve = QPushButton("Migliora")
        btn_technical = QPushButton("Tecnico")
        btn_summarize = QPushButton("Sintetizza")
        btn_correct = QPushButton("Correggi")

        preset_row.addWidget(btn_improve)
        preset_row.addWidget(btn_technical)
        preset_row.addWidget(btn_summarize)
        preset_row.addWidget(btn_correct)
        preset_row.addStretch()

        layout.addLayout(preset_row)

        # collegamenti preset
        btn_improve.clicked.connect(
            lambda: prompt_edit.setPlainText("Migliora il testo mantenendo il significato.")
        )
        btn_technical.clicked.connect(
            lambda: prompt_edit.setPlainText("Rendi il testo più tecnico e formale.")
        )
        btn_summarize.clicked.connect(
            lambda: prompt_edit.setPlainText("Sintetizza il testo mantenendo i concetti chiave.")
        )
        btn_correct.clicked.connect(
            lambda: prompt_edit.setPlainText("Correggi errori grammaticali e sintattici mantenendo il contenuto tecnico.")
        )

        # --- azioni finali ---
        action_row = QHBoxLayout()

        btn_generate = QPushButton("Genera")
        btn_replace = QPushButton("Sostituisci")
        btn_cancel = QPushButton("Annulla")

        btn_cancel.clicked.connect(dialog.reject)
        def apply_text():
            new_text = modified_edit.toPlainText()
            cursor = self.editor.textCursor()

            if cursor.hasSelection():
                cursor.insertText(new_text)
            else:
                self.editor.setPlainText(new_text)

            dialog.accept()

        btn_replace.clicked.connect(apply_text)
        def generate_preview():
            prompt = prompt_edit.toPlainText().strip()

            if not saved_api_key:
                QMessageBox.warning(
                    dialog,
                    "API key mancante",
                    "Configura prima la API key del provider AI dal pulsante AI generale."
                )
                return

            if not prompt:
                QMessageBox.warning(
                    dialog,
                    "Prompt mancante",
                    "Scrivi prima un prompt oppure usa un comando predefinito."
                )
                return

            modified_edit.setPlainText("Generazione in corso...")
            btn_generate.setEnabled(False)

            self.ai_block_thread = QThread()
            self.ai_block_worker = GeminiWorker(
                api_key=saved_api_key,
                model_name=saved_model,
                prompt_text=prompt,
                mode="Rielabora testo",
                style="Libero",
                source_text=original_text,
                provider=ai_provider,
                reasoning=ai_reasoning,
            )

            receiver = AiResultReceiver(modified_edit, btn_generate, dialog)

            self.ai_block_worker.moveToThread(self.ai_block_thread)
            self.ai_block_thread.started.connect(self.ai_block_worker.run)

            self.ai_block_worker.finished.connect(receiver.show_result)
            self.ai_block_worker.finished.connect(self.ai_block_thread.quit)
            self.ai_block_worker.finished.connect(self.ai_block_worker.deleteLater)

            self.ai_block_worker.error.connect(receiver.show_error)
            self.ai_block_worker.error.connect(self.ai_block_thread.quit)
            self.ai_block_worker.error.connect(self.ai_block_worker.deleteLater)

            self.ai_block_thread.finished.connect(receiver.generation_finished)
            self.ai_block_thread.finished.connect(self.ai_block_thread.deleteLater)

            self.ai_block_thread.start()

        btn_generate.clicked.connect(generate_preview)

        action_row.addStretch()
        action_row.addWidget(btn_generate)
        action_row.addWidget(btn_replace)
        action_row.addWidget(btn_cancel)

        layout.addLayout(action_row)
        dialog.exec()

class ImageBlockWidget(BaseBlockWidget):
    block_type_label = "Immagine"

    def __init__(self, parent_editor, data=None):
        base = {"type": "image", "path": "", "caption": "", "collapsed": False}
        base.update(data or {})
        super().__init__(parent_editor, base)
        self.image_path = self.data.get("path", "")

        top_row = QHBoxLayout()
        top_row.setContentsMargins(0, 0, 0, 0)

        self.btn_load = QPushButton("Immagine…")
        self.btn_paste = QPushButton("Incolla immagine")
        self.btn_view = QPushButton("Visualizza")
        self.btn_copy_ref = QPushButton("Copia ID immagine")
        for btn in (self.btn_load, self.btn_paste, self.btn_view, self.btn_copy_ref):
            style_small_button(btn)
        self.btn_load.clicked.connect(self.load_image)
        self.btn_paste.clicked.connect(self.paste_image_from_clipboard)
        self.btn_view.clicked.connect(self.show_preview)
        self.btn_copy_ref.clicked.connect(self.copy_reference_token)

        top_row.addWidget(self.btn_load)
        top_row.addWidget(self.btn_paste)
        top_row.addWidget(self.btn_view)
        top_row.addWidget(self.btn_copy_ref)

        self.caption_label = QLabel("")
        self.caption_label.setStyleSheet("color: #666;")
        self.caption_label.setWordWrap(True)
        top_row.addWidget(self.caption_label, 1)
        self.content_layout.addLayout(top_row)

        self.reference_label = QLabel("")
        self.reference_label.setStyleSheet("color: #3b5b8a;")
        self.reference_label.setWordWrap(True)
        self.content_layout.addWidget(self.reference_label)

        self.update_summary()
        self.refresh_figure_info()

    def _display_file_name(self):
        if not self.image_path:
            return "Nessuna immagine"
        return os.path.basename(as_posix_path(self.image_path))

    def resolved_path(self):
        return self.parent_editor.main_window.resolve_runtime_path(self.image_path)

    def reference_token(self):
        return f"[[FIG:{self.data.get('id','')}]]"

    def copy_reference_token(self):
        QApplication.clipboard().setText(self.reference_token())
        self.parent_editor.main_window.statusBar().showMessage("Riferimento immagine copiato", 2000)

    def load_image(self):
        path, _ = QFileDialog.getOpenFileName(self, "Seleziona immagine", "", "Immagini (*.png *.jpg *.jpeg *.bmp *.gif *.tif *.tiff)")
        if not path:
            return
        self.image_path = self.parent_editor.main_window.ingest_image_into_project(path)
        self.update_summary()
        self.parent_editor.main_window.mark_dirty()
        self.parent_editor.main_window.update_figure_references()

    def paste_image_from_clipboard(self):
        clipboard = QApplication.clipboard()
        image = clipboard.image()

        if image.isNull():
            QMessageBox.warning(
                self,
                "Incolla immagine",
                "Negli appunti non è presente un'immagine valida."
            )
            return

        block_id = self.data.get("id") or new_block_id()
        self.data["id"] = block_id

        if self.parent_editor.main_window.project_dir:
            images_dir = self.parent_editor.main_window.assets_images_dir()
            file_path = os.path.join(images_dir, f"{block_id}.png")
            stored_path = f"assets/images/{block_id}.png"
        else:
            temp_dir = tempfile.gettempdir()
            file_path = os.path.join(temp_dir, f"{block_id}.png")
            stored_path = file_path

        if not image.save(file_path, "PNG"):
            QMessageBox.critical(
                self,
                "Errore",
                "Impossibile salvare l'immagine dagli appunti."
            )
            return

        self.image_path = stored_path
        self.update_summary()
        self.parent_editor.main_window.mark_dirty()
        self.parent_editor.main_window.update_figure_references()

    def show_preview(self):
        resolved = self.resolved_path()
        if not resolved:
            return
        dialog = ImagePreviewDialog(resolved, self.data.get("caption", ""), self)
        dialog.exec()
        self.data["caption"] = dialog.caption_value
        self.update_summary()
        self.parent_editor.main_window.mark_dirty()
        self.parent_editor.main_window.update_figure_references()

    def refresh_figure_info(self):
        fig_map = self.parent_editor.main_window.figure_reference_map
        block_id = self.data.get("id", "")
        info = fig_map.get(block_id)
        if info:
            self.reference_label.setText(f"{info['label']} — incolla nel testo: {self.reference_token()}")
        else:
            self.reference_label.setText(f"Incolla nel testo: {self.reference_token()}")

    def build_summary(self):
        caption = self.data.get("caption", "").strip()
        if caption:
            self.caption_label.setText(caption)
            return f"{self.block_type_label} — {caption}"
        self.caption_label.setText(self._display_file_name())
        return f"{self.block_type_label} — {self._display_file_name()}"

    def export_data(self):
        data = {"type": "image", "path": self.image_path, "caption": self.data.get("caption", "")}
        data.update(self.export_common_data())
        return data


class MultiImageBlockWidget(BaseBlockWidget):
    block_type_label = "Gruppo immagini"

    def __init__(self, parent_editor, data=None):
        base = {"type": "images", "title": "", "images": [], "collapsed": False}
        base.update(data or {})
        super().__init__(parent_editor, base)
        self.images = self.parent_editor.main_window.ensure_group_images_defaults(deepcopy(self.data.get("images", [])))

        title_row = QHBoxLayout()
        title_row.setContentsMargins(0, 0, 0, 0)
        title_row.addWidget(QLabel("Titolo blocco"))
        self.title_edit = QLineEdit()
        self.title_edit.setPlaceholderText("Es. diagramma di carico")
        self.title_edit.setText(base.get("title", ""))
        self.title_edit.textChanged.connect(self.on_title_changed)
        title_row.addWidget(self.title_edit, 1)
        self.content_layout.addLayout(title_row)

        top_row = QHBoxLayout()
        top_row.setContentsMargins(0, 0, 0, 0)
        self.btn_load = QPushButton("Aggiungi immagini…")
        self.btn_view = QPushButton("Visualizza")
        self.btn_copy_refs = QPushButton("Copia ID gruppo immagini")
        for btn in (self.btn_load, self.btn_view, self.btn_copy_refs):
            style_small_button(btn)
        self.btn_load.clicked.connect(self.load_images)
        self.btn_view.clicked.connect(self.show_preview)
        self.btn_copy_refs.clicked.connect(self.copy_reference_tokens)
        top_row.addWidget(self.btn_load)
        top_row.addWidget(self.btn_view)
        top_row.addWidget(self.btn_copy_refs)

        self.info_label = QLabel("")
        self.info_label.setStyleSheet("color: #666;")
        top_row.addWidget(self.info_label, 1)
        self.content_layout.addLayout(top_row)

        self.reference_label = QLabel("")
        self.reference_label.setStyleSheet("color: #3b5b8a;")
        self.reference_label.setWordWrap(True)
        self.content_layout.addWidget(self.reference_label)

        remove_row = QHBoxLayout()
        remove_row.setContentsMargins(0, 0, 0, 0)
        remove_row.addWidget(QLabel("Rimuovi immagine"))
        self.remove_image_combo = QComboBox()
        self.remove_image_combo.setMinimumWidth(220)
        self.btn_remove_image = QPushButton("Rimuovi")
        style_small_button(self.btn_remove_image)
        self.btn_remove_image.clicked.connect(self.remove_selected_image)
        remove_row.addWidget(self.remove_image_combo, 1)
        remove_row.addWidget(self.btn_remove_image)
        self.content_layout.addLayout(remove_row)

        self.update_info_label()
        self.update_remove_combo()
        self.update_summary()
        self.refresh_figure_info()

    def image_reference_tokens(self):
        return [f"[[FIG:{img.get('id','')}]]" for img in self.images if img.get("id")]

    def copy_reference_tokens(self):
        text = format_reference_tokens(self.image_reference_tokens())
        QApplication.clipboard().setText(text)
        self.parent_editor.main_window.statusBar().showMessage("Riferimenti gruppo immagini copiati", 2000)

    def load_images(self):
        paths, _ = QFileDialog.getOpenFileNames(self, "Seleziona immagini", "", "Immagini (*.png *.jpg *.jpeg *.bmp *.gif *.tif *.tiff)")
        if not paths:
            return
        for path in paths:
            stored_path = self.parent_editor.main_window.ingest_image_into_project(path)
            self.images.append({"id": new_block_id(), "path": stored_path, "caption": ""})
        self.update_info_label()
        self.update_remove_combo()
        self.update_summary()
        self.parent_editor.main_window.mark_dirty()
        self.parent_editor.main_window.update_figure_references()

    def update_info_label(self):
        count = len(self.images)
        if count == 0:
            self.info_label.setText("Nessuna immagine")
        elif count == 1:
            self.info_label.setText("1 immagine")
        else:
            self.info_label.setText(f"{count} immagini")

    def update_remove_combo(self):
        self.remove_image_combo.clear()
        for idx, img in enumerate(self.images, start=1):
            caption = img.get("caption", "").strip()
            label = caption if caption else os.path.basename(as_posix_path(img.get("path", ""))) or f"Immagine {idx}"
            self.remove_image_combo.addItem(f"{idx}. {label}", img.get("id"))

    def remove_selected_image(self):
        idx = self.remove_image_combo.currentIndex()
        if idx < 0 or idx >= len(self.images):
            return
        self.images.pop(idx)
        self.update_info_label()
        self.update_remove_combo()
        self.update_summary()
        self.parent_editor.main_window.mark_dirty()
        self.parent_editor.main_window.update_figure_references()

    def resolved_images(self):
        resolved = []
        for img in self.images:
            entry = deepcopy(img)
            entry["path"] = self.parent_editor.main_window.resolve_runtime_path(img.get("path", ""))
            resolved.append(entry)
        return resolved

    def show_preview(self):
        if not self.images:
            return
        resolved = self.resolved_images()
        dialog = MultiImagePreviewDialog(resolved, self)
        dialog.exec()
        updated = []
        for idx, original in enumerate(self.images):
            item = deepcopy(original)
            if idx < len(resolved):
                item["caption"] = resolved[idx].get("caption", "")
            updated.append(item)
        self.images = self.parent_editor.main_window.ensure_group_images_defaults(updated)
        self.update_remove_combo()
        self.parent_editor.main_window.mark_dirty()
        self.parent_editor.main_window.update_figure_references()
        self.update_summary()

    def on_title_changed(self):
        self.update_summary()
        self.parent_editor.main_window.mark_dirty()
        self.parent_editor.main_window.update_figure_references()

    def refresh_figure_info(self):
        parts = []
        for img in self.images:
            info = self.parent_editor.main_window.figure_reference_map.get(img.get("id"))
            if info:
                parts.append(info["label"])
        if parts:
            self.reference_label.setText(
                f"{format_reference_tokens(parts)} — incolla nel testo: {format_reference_tokens(self.image_reference_tokens())}"
            )
        else:
            self.reference_label.setText(
                f"Incolla nel testo: {format_reference_tokens(self.image_reference_tokens())}"
            )

    def build_summary(self):
        count = len(self.images)
        title = self.title_edit.text().strip()
        if title:
            return f"{self.block_type_label} — {title} ({count})"
        return f"{self.block_type_label} — {count} elementi"

    def export_data(self):
        data = {"type": "images", "title": self.title_edit.text(), "images": deepcopy(self.images)}
        data.update(self.export_common_data())
        return data

class EquationBlockWidget(BaseBlockWidget):
    block_type_label = "Equazione"

    def __init__(self, parent_editor, data=None):
        base = {
            "type": "equation",
            "latex": "",
            "caption": "",
            "numbering_mode": "none",
            "collapsed": False,
        }
        base.update(data or {})
        super().__init__(parent_editor, base)

        self.top_row = QHBoxLayout()
        self.top_row.addWidget(QLabel("Formula LaTeX"))

        self.btn_frac = QPushButton("a⁄b")
        self.btn_sup = QPushButton("xⁿ")
        self.btn_sub = QPushButton("xᵢ")
        self.btn_sqrt = QPushButton("√")
        self.btn_integral = QPushButton("∫")
        self.btn_sum = QPushButton("Σ")
        self.btn_overline = QPushButton("x̅")
        self.btn_brackets = QPushButton("( )")
        self.btn_greek = QPushButton("αβ")
        self.btn_focus = QPushButton("⛶")
        self.btn_ai = QPushButton("AI")
        self.btn_copy_ref = QPushButton("Copia riferimento")

        for btn in (
            self.btn_frac,
            self.btn_sup,
            self.btn_sub,
            self.btn_sqrt,
            self.btn_integral,
            self.btn_sum,
            self.btn_overline,
            self.btn_brackets,
            self.btn_greek,
            self.btn_ai,
            self.btn_copy_ref,
            self.btn_focus,
        ):
            style_small_button(btn)
            btn.setStyleSheet(
                "QPushButton { border: 1px solid #d6d6d6; border-radius: 4px; background: #fafafa; padding: 0 8px; }"
                "QPushButton:hover { background: #f0f0f0; }"
                "QToolTip { background-color: #f2f2f2; color: black; border: 1px solid #c8c8c8; padding: 4px; }"
            )
            self.top_row.addWidget(btn)

        self.btn_frac.setToolTip("Inserisci frazione")
        self.btn_sup.setToolTip("Inserisci apice")
        self.btn_sub.setToolTip("Inserisci pedice")
        self.btn_sqrt.setToolTip("Inserisci radice")
        self.btn_integral.setToolTip("Inserisci integrale")
        self.btn_sum.setToolTip("Inserisci sommatoria")
        self.btn_overline.setToolTip("Accenti e vettori")
        self.btn_brackets.setToolTip("Parentesi e matrici")
        self.btn_greek.setToolTip("Lettere greche, operatori e frecce")
        self.btn_focus.setToolTip("Editor equazione esteso")
        self.btn_ai.setToolTip("Assistente AI formule")
        self.btn_copy_ref.setToolTip("Copia riferimento equazione")

        self.top_row.addStretch()
        self.content_layout.addLayout(self.top_row)

        self.latex_edit = QTextEdit()
        self.latex_edit.setPlaceholderText(
            r"Esempio: \sigma = \frac{N}{A} + \frac{M}{W}"
        )
        self.latex_edit.setMinimumHeight(90)
        self.latex_edit.setPlainText(base.get("latex", ""))
        self.latex_edit.setStyleSheet("""
            QTextEdit {
                background: white;
                color: black;
                border: 1px solid #d0d0d0;
                font-family: Consolas, monospace;
                font-size: 14px;
            }
        """)
        self.content_layout.addWidget(self.latex_edit)

        self.preview_label = QLabel()
        self.preview_label.setAlignment(Qt.AlignCenter)
        self.preview_label.setMinimumHeight(120)
        self.preview_label.setStyleSheet("""
            QLabel {
                background: white;
                color: black;
                border: 1px solid #e0e0e0;
                padding: 10px;
            }
        """)
        self.content_layout.addWidget(self.preview_label)

        numbering_row = QHBoxLayout()
        numbering_row.setContentsMargins(0, 0, 0, 0)

        numbering_row.addWidget(QLabel("Numerazione"))

        self.numbering_mode_combo = QComboBox()
        self.numbering_mode_combo.addItem("Nessuna numerazione", "none")
        self.numbering_mode_combo.addItem("Solo numero", "number")
        self.numbering_mode_combo.addItem("Numero + descrizione", "number_caption")

        saved_numbering_mode = base.get("numbering_mode", "none")
        index = self.numbering_mode_combo.findData(saved_numbering_mode)
        if index < 0:
            index = 0
        self.numbering_mode_combo.setCurrentIndex(index)

        numbering_row.addWidget(self.numbering_mode_combo)

        self.caption_edit = QLineEdit()
        self.caption_edit.setPlaceholderText("Descrizione equazione")
        self.caption_edit.setText(base.get("caption", ""))
        self.caption_edit.setStyleSheet("background: white; color: black;")
        numbering_row.addWidget(self.caption_edit, 1)

        self.content_layout.addLayout(numbering_row)

        self.caption_edit.setVisible(
            self.numbering_mode_combo.currentData() == "number_caption"
        )

        self.reference_label = QLabel("")
        self.reference_label.setStyleSheet("color: #3b5b8a;")
        self.reference_label.setWordWrap(True)
        self.content_layout.addWidget(self.reference_label)

        self.btn_focus.clicked.connect(self.open_focus_editor)
        self.btn_ai.clicked.connect(self.open_equation_ai_dialog)
        self.btn_copy_ref.clicked.connect(self.copy_reference_token)

        self.btn_frac.clicked.connect(
            lambda: self.insert_latex_snippet(r"\frac{}{}", -3)
        )

        self.btn_sup.clicked.connect(
            lambda: self.insert_latex_snippet(r"^{}", -1)
        )

        self.btn_sub.clicked.connect(
            lambda: self.insert_latex_snippet(r"_{}", -1)
        )

        self.btn_sqrt.clicked.connect(
            lambda: self.insert_latex_snippet(r"\sqrt{}", -1)
        )

        self.btn_integral.clicked.connect(self.show_integral_menu)
        self.btn_sum.clicked.connect(self.show_sum_menu)
        self.btn_overline.clicked.connect(self.show_accent_menu)
        self.btn_brackets.clicked.connect(self.show_brackets_menu)
        self.btn_greek.clicked.connect(self.show_greek_menu)

        self.latex_edit.textChanged.connect(self.on_changed)
        self.caption_edit.textChanged.connect(self.on_changed)

        self.numbering_mode_combo.currentIndexChanged.connect(
            self.on_numbering_mode_changed
        )

        self.refresh_preview()
        self.update_summary()
        self.refresh_equation_info()

    def on_changed(self):
        self.refresh_preview()
        self.update_summary()
        self.parent_editor.main_window.mark_dirty()

    def on_numbering_mode_changed(self):
        mode = self.numbering_mode_combo.currentData()

        self.caption_edit.setVisible(
            mode == "number_caption"
        )

        if mode != "number_caption":
            self.caption_edit.clear()

        self.update_summary()
        self.parent_editor.main_window.mark_dirty()

    def insert_latex_snippet(self, snippet, cursor_offset=0):
        cursor = self.latex_edit.textCursor()

        if cursor.hasSelection():
            selected = cursor.selectedText()

            if "{}" in snippet:
                snippet = snippet.replace("{}", "{" + selected + "}", 1)
                cursor_offset = 0
            else:
                snippet = snippet + selected

        cursor.insertText(snippet)

        if cursor_offset != 0:
            cursor.movePosition(
                QTextCursor.MoveOperation.Left if cursor_offset < 0 else QTextCursor.MoveOperation.Right,
                QTextCursor.MoveMode.MoveAnchor,
                abs(cursor_offset)
            )

        self.latex_edit.setTextCursor(cursor)
        self.latex_edit.setFocus()
        self.refresh_preview()
        self.update_summary()
        self.parent_editor.main_window.mark_dirty()

    def show_greek_menu(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Simboli LaTeX")
        dialog.setStyleSheet("""
            QDialog {
                background: #f2f2f2;
                color: black;
            }
            QPushButton {
                border: 1px solid #d6d6d6;
                border-radius: 4px;
                background: #fafafa;
                padding: 4px 8px;
                color: black;
            }
            QPushButton:hover {
                background: #e8e8e8;
            }
        """)

        layout = QHBoxLayout(dialog)

        col_greek = QVBoxLayout()
        col_operators = QVBoxLayout()
        col_arrows = QVBoxLayout()

        col_greek.addWidget(QLabel("Lettere greche"))

        greek_items = [
            ("α  alpha", r"\alpha"),
            ("β  beta", r"\beta"),
            ("γ  gamma", r"\gamma"),
            ("δ  delta", r"\delta"),
            ("Δ  Delta", r"\Delta"),
            ("ε  epsilon", r"\varepsilon"),
            ("θ  theta", r"\theta"),
            ("λ  lambda", r"\lambda"),
            ("μ  mu", r"\mu"),
            ("ν  nu", r"\nu"),
            ("π  pi", r"\pi"),
            ("ρ  rho", r"\rho"),
            ("σ  sigma", r"\sigma"),
            ("τ  tau", r"\tau"),
            ("φ  phi", r"\varphi"),
            ("ω  omega", r"\omega"),
        ]

        for label, latex in greek_items:
            btn = QPushButton(label)
            btn.clicked.connect(
                lambda checked=False, value=latex: (
                    self.insert_latex_snippet(value, 0),
                    dialog.accept()
                )
            )
            col_greek.addWidget(btn)

        col_operators.addWidget(QLabel("Operatori"))

        operator_items = [
            ("±  plus/minus", r"\pm"),
            ("×  per", r"\times"),
            ("÷  diviso", r"\div"),
            ("≤  minore uguale", r"\leq"),
            ("≥  maggiore uguale", r"\geq"),
            ("≠  diverso", r"\neq"),
            ("≈  circa uguale", r"\approx"),
            ("∞  infinito", r"\infty"),
        ]

        for label, latex in operator_items:
            btn = QPushButton(label)
            btn.clicked.connect(
                lambda checked=False, value=latex: (
                    self.insert_latex_snippet(value, 0),
                    dialog.accept()
                )
            )
            col_operators.addWidget(btn)

        col_arrows.addWidget(QLabel("Frecce e relazioni"))

        arrow_items = [
            ("→  right arrow", r"\rightarrow"),
            ("⇒  implies", r"\Rightarrow"),
            ("↔  double arrow", r"\leftrightarrow"),
            ("⇔  iff", r"\Leftrightarrow"),
            ("∈  appartiene", r"\in"),
            ("∉  non appartiene", r"\notin"),
            ("⊂  sottoinsieme", r"\subset"),
            ("⊆  sottoinsieme eq", r"\subseteq"),
        ]

        for label, latex in arrow_items:
            btn = QPushButton(label)
            btn.clicked.connect(
                lambda checked=False, value=latex: (
                    self.insert_latex_snippet(value, 0),
                    dialog.accept()
                )
            )
            col_arrows.addWidget(btn)

        col_greek.addStretch()
        col_operators.addStretch()
        col_arrows.addStretch()

        layout.addLayout(col_greek)
        layout.addLayout(col_operators)
        layout.addLayout(col_arrows)

        dialog.exec()

    def show_brackets_menu(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Parentesi e matrici")
        dialog.setStyleSheet("""
            QDialog {
                background: #f2f2f2;
                color: black;
            }
            QPushButton {
                border: 1px solid #d6d6d6;
                border-radius: 4px;
                background: #fafafa;
                padding: 4px 8px;
                color: black;
            }
            QPushButton:hover {
                background: #e8e8e8;
            }
        """)

        layout = QVBoxLayout(dialog)

        items = [
            ("( )", r"\left(  \right)", -8),
            ("[ ]", r"\left[  \right]", -8),
            ("{ }", r"\left\{  \right\}", -9),
            ("| |", r"\left|  \right|", -8),
            ("‖ ‖", r"\left\|  \right\|", -9),
            ("Matrice 2x2", r"\begin{matrix} a & b \\ c & d \end{matrix}", 0),
            ("Matrice 3x3", r"\begin{matrix} a & b & c \\ d & e & f \\ g & h & i \end{matrix}", 0),
            ("Sistema / casi", r"\begin{cases} a & b \\ c & d \end{cases}", 0),
        ]

        for label, latex, offset in items:
            btn = QPushButton(label)
            btn.clicked.connect(
                lambda checked=False, value=latex, off=offset: (
                    self.insert_latex_snippet(value, off),
                    dialog.accept()
                )
            )
            layout.addWidget(btn)

        dialog.exec()

    def show_integral_menu(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Integrali")
        dialog.setStyleSheet("""
            QDialog {
                background: #f2f2f2;
                color: black;
            }
            QPushButton {
                border: 1px solid #d6d6d6;
                border-radius: 4px;
                background: #fafafa;
                padding: 4px 8px;
                color: black;
            }
            QPushButton:hover {
                background: #e8e8e8;
            }
        """)

        layout = QVBoxLayout(dialog)

        items = [
            ("Integrale indefinito", r"\int  \,dx", -5),
            ("Integrale definito", r"\int_{}^{}  \,dx", -8),
            ("Integrale doppio", r"\iint  \,dA", -5),
            ("Integrale triplo", r"\iiint  \,dV", -5),
            ("Integrale curvilineo", r"\int_C  \,ds", -5),
            ("Integrale di superficie", r"\iint_S  \,dS", -5),
            ("Integrale chiuso", r"\oint  \,ds", -5),
            ("Integrale chiuso su C", r"\oint_C  \,ds", -5),
        ]

        for label, latex, offset in items:
            btn = QPushButton(label)
            btn.clicked.connect(
                lambda checked=False, value=latex, off=offset: (
                    self.insert_latex_snippet(value, off),
                    dialog.accept()
                )
            )
            layout.addWidget(btn)

        dialog.exec()

    def show_sum_menu(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Sommatorie e produttorie")
        dialog.setStyleSheet("""
            QDialog {
                background: #f2f2f2;
                color: black;
            }
            QPushButton {
                border: 1px solid #d6d6d6;
                border-radius: 4px;
                background: #fafafa;
                padding: 4px 8px;
                color: black;
            }
            QPushButton:hover {
                background: #e8e8e8;
            }
        """)

        layout = QVBoxLayout(dialog)

        items = [
            ("Sommatoria semplice", r"\sum  ", -1),
            ("Sommatoria con indice", r"\sum_{}  ", -3),
            ("Sommatoria con limiti", r"\sum_{}^{}  ", -5),
            ("Sommatoria classica i=1..n", r"\sum_{i=1}^{n} x_i", 0),
            ("Produttoria semplice", r"\prod  ", -1),
            ("Produttoria con indice", r"\prod_{}  ", -3),
            ("Produttoria con limiti", r"\prod_{}^{}  ", -5),
            ("Produttoria classica i=1..n", r"\prod_{i=1}^{n} x_i", 0),
        ]

        for label, latex, offset in items:
            btn = QPushButton(label)
            btn.clicked.connect(
                lambda checked=False, value=latex, off=offset: (
                    self.insert_latex_snippet(value, off),
                    dialog.accept()
                )
            )
            layout.addWidget(btn)

        dialog.exec()

    def show_accent_menu(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Accenti e vettori")
        dialog.setStyleSheet("""
            QDialog {
                background: #f2f2f2;
                color: black;
            }
            QPushButton {
                border: 1px solid #d6d6d6;
                border-radius: 4px;
                background: #fafafa;
                padding: 4px 8px;
                color: black;
            }
            QPushButton:hover {
                background: #e8e8e8;
            }
        """)

        layout = QVBoxLayout(dialog)

        items = [
            ("Barra sopra", r"\bar{}", -1),
            ("Linea sopra", r"\overline{}", -1),
            ("Cappello", r"\hat{}", -1),
            ("Tilde", r"\tilde{}", -1),
            ("Punto sopra", r"\dot{}", -1),
            ("Doppio punto sopra", r"\ddot{}", -1),
            ("Vettore", r"\vec{}", -1),
            ("Vettore grassetto", r"\mathbf{}", -1),
            ("Versore", r"\hat{\mathbf{}}", -2),
        ]

        for label, latex, offset in items:
            btn = QPushButton(label)
            btn.clicked.connect(
                lambda checked=False, value=latex, off=offset: (
                    self.insert_latex_snippet(value, off),
                    dialog.accept()
                )
            )
            layout.addWidget(btn)

        dialog.exec()

    def open_focus_editor(self):
        dialog = QDialog(self.window())
        dialog.setWindowTitle("Editor equazione esteso")

        settings = self.parent_editor.main_window.settings
        saved_geometry = settings.value("equation_focus_mode/geometry")

        if saved_geometry:
            dialog.restoreGeometry(saved_geometry)
        else:
            dialog.resize(1100, 720)

        dialog.setWindowOpacity(0.0)

        self.focus_anim = QPropertyAnimation(dialog, b"windowOpacity")
        self.focus_anim.setDuration(160)
        self.focus_anim.setStartValue(0.0)
        self.focus_anim.setEndValue(1.0)

        layout = QVBoxLayout(dialog)

        # Sposta temporaneamente TUTTO il contenuto del blocco equazione,
        # inclusa toolbar, editor LaTeX, preview e didascalia.
        self.root_layout.removeWidget(self.content_widget)
        layout.addWidget(self.content_widget)

        buttons_row = QHBoxLayout()
        buttons_row.addStretch()

        btn_close = QPushButton("Chiudi")
        btn_close.setFixedWidth(120)
        style_small_button(btn_close)
        btn_close.clicked.connect(dialog.accept)

        buttons_row.addWidget(btn_close)
        layout.addLayout(buttons_row)

        old_latex_min_height = self.latex_edit.minimumHeight()
        old_preview_min_height = self.preview_label.minimumHeight()

        self.latex_edit.setMinimumHeight(320)
        self.preview_label.setMinimumHeight(220)
        self.btn_focus.hide()

        try:
            self.latex_edit.setFocus()
            self.focus_anim.start()
            dialog.exec()
        finally:
            settings.setValue("equation_focus_mode/geometry", dialog.saveGeometry())

            layout.removeWidget(self.content_widget)
            self.root_layout.insertWidget(1, self.content_widget)

            self.latex_edit.setMinimumHeight(old_latex_min_height)
            self.preview_label.setMinimumHeight(old_preview_min_height)

            self.btn_focus.show()
            self.refresh_preview()
            self.update_summary()
            self.parent_editor.main_window.mark_dirty()

    def open_equation_ai_dialog(self):
        dialog = QDialog(self.window())
        dialog.setWindowTitle("AI blocco equazione")
        dialog.resize(760, 560)

        layout = QVBoxLayout(dialog)

        cursor = self.latex_edit.textCursor()
        if cursor.hasSelection():
            source_text = cursor.selectedText()
        else:
            source_text = self.latex_edit.toPlainText()

        texts_row = QHBoxLayout()

        source_edit = QTextEdit()
        source_edit.setPlainText(source_text)
        source_edit.setReadOnly(True)

        result_edit = QTextEdit()
        result_edit.setPlaceholderText("Il LaTeX generato comparirà qui...")

        texts_row.addWidget(source_edit)
        texts_row.addWidget(result_edit)

        layout.addLayout(texts_row)

        layout.addWidget(QLabel("Prompt:"))

        prompt_edit = QTextEdit()
        prompt_edit.setMaximumHeight(90)
        prompt_edit.setPlaceholderText(
            "Es. sviluppa i passaggi del calcolo, semplifica la formula, riscrivi in forma compatta..."
        )
        layout.addWidget(prompt_edit)

        preset_row = QHBoxLayout()

        btn_develop = QPushButton("Sviluppa")
        btn_simplify = QPushButton("Semplifica")
        btn_latex = QPushButton("Correggi LaTeX")
        btn_structural = QPushButton("Tecnico")

        preset_row.addWidget(btn_develop)
        preset_row.addWidget(btn_simplify)
        preset_row.addWidget(btn_latex)
        preset_row.addWidget(btn_structural)
        preset_row.addStretch()

        layout.addLayout(preset_row)

        btn_develop.clicked.connect(
            lambda: prompt_edit.setPlainText(
                "Sviluppa i passaggi matematici della formula, una formula per riga."
            )
        )

        btn_simplify.clicked.connect(
            lambda: prompt_edit.setPlainText(
                "Semplifica la formula mantenendo il significato matematico."
            )
        )

        btn_latex.clicked.connect(
            lambda: prompt_edit.setPlainText(
                "Correggi la sintassi LaTeX mantenendo invariato il contenuto matematico."
            )
        )

        btn_structural.clicked.connect(
            lambda: prompt_edit.setPlainText(
                "Riscrivi la formula in forma tecnica adatta a una relazione strutturale."
            )
        )

        action_row = QHBoxLayout()

        btn_generate = QPushButton("Genera")
        btn_insert = QPushButton("Inserisci")
        btn_cancel = QPushButton("Annulla")

        action_row.addStretch()
        action_row.addWidget(btn_generate)
        action_row.addWidget(btn_insert)
        action_row.addWidget(btn_cancel)

        layout.addLayout(action_row)

        main_window = self.parent_editor.main_window
        receiver = AiResultReceiver(result_edit, btn_generate, dialog)

        def generate_latex():
            ai_provider, api_key, model_name, ai_reasoning = main_window.get_active_ai_config()
            prompt = prompt_edit.toPlainText().strip()
            current_source = source_edit.toPlainText().strip()

            if not api_key:
                QMessageBox.warning(
                    dialog,
                    "API key mancante",
                    "Configura prima la API key del provider AI."
                )
                return

            if not prompt:
                QMessageBox.warning(
                    dialog,
                    "Prompt mancante",
                    "Scrivi prima un prompt oppure usa un comando predefinito."
                )
                return

            full_prompt = f"""
    Sei un assistente tecnico per formule matematiche e strutturali.

    Rielabora il testo sorgente secondo le istruzioni dell'utente.
    Restituisci SOLO codice LaTeX valido per matplotlib mathtext.
    Non scrivere spiegazioni.
    Non usare markdown.
    Non usare blocchi ```.

    Se serve mostrare uno sviluppo di calcolo, scrivi una formula per riga.
    Non inventare dati numerici non presenti.
    Se mancano dati, usa simboli.

    Testo sorgente:
    {current_source}

    Istruzioni:
    {prompt}
    """

            result_edit.setPlainText("Generazione in corso...")
            btn_generate.setEnabled(False)

            self.equation_ai_thread = QThread()
            self.equation_ai_worker = GeminiWorker(
                api_key=api_key,
                model_name=model_name,
                prompt_text=full_prompt,
                mode="Genera testo",
                style="Libero",
                source_text="",
                provider=ai_provider,
                reasoning=ai_reasoning,
            )

            self.equation_ai_worker.moveToThread(self.equation_ai_thread)
            self.equation_ai_thread.started.connect(self.equation_ai_worker.run)

            self.equation_ai_worker.finished.connect(receiver.show_result)
            self.equation_ai_worker.finished.connect(self.equation_ai_thread.quit)
            self.equation_ai_worker.finished.connect(self.equation_ai_worker.deleteLater)

            self.equation_ai_worker.error.connect(receiver.show_error)
            self.equation_ai_worker.error.connect(self.equation_ai_thread.quit)
            self.equation_ai_worker.error.connect(self.equation_ai_worker.deleteLater)

            self.equation_ai_thread.finished.connect(receiver.generation_finished)
            self.equation_ai_thread.finished.connect(self.equation_ai_thread.deleteLater)

            self.equation_ai_thread.start()

        def insert_result():
            text = result_edit.toPlainText().strip()

            if not text:
                return

            cursor = self.latex_edit.textCursor()

            if cursor.hasSelection():
                cursor.insertText(text)
            else:
                self.latex_edit.setPlainText(text)

            self.refresh_preview()
            self.update_summary()
            self.parent_editor.main_window.mark_dirty()

            dialog.accept()

        btn_generate.clicked.connect(generate_latex)
        btn_insert.clicked.connect(insert_result)
        btn_cancel.clicked.connect(dialog.reject)

        dialog.exec()

    def reference_token(self):
        return f"[[EQ:{self.data.get('id', '')}]]"


    def copy_reference_token(self):
        token = self.reference_token()
        QApplication.clipboard().setText(token)
        self.parent_editor.main_window.statusBar().showMessage(
            "Riferimento equazione copiato",
            2000
        )

    def render_equation_pixmap(self):
        raw_text = self.latex_edit.toPlainText().strip()

        if not raw_text:
            return None, "Nessuna equazione"

        lines = [
            line.strip()
            for line in raw_text.splitlines()
            if line.strip()
        ]

        if not lines:
            return None, "Nessuna equazione"

        try:
            line_count = len(lines)
            fig_height = max(1.2, 0.65 * line_count)

            fig = Figure(figsize=(6.0, fig_height), dpi=180)
            fig.patch.set_alpha(0.0)

            canvas = FigureCanvasAgg(fig)
            ax = fig.add_axes([0, 0, 1, 1])
            ax.axis("off")

            for idx, line in enumerate(lines):
                latex = line
                if not (latex.startswith("$") and latex.endswith("$")):
                    latex = f"${latex}$"

                y = 1.0 - ((idx + 0.5) / line_count)

                ax.text(
                    0.5,
                    y,
                    latex,
                    fontsize=18,
                    ha="center",
                    va="center",
                )

            canvas.draw()

            width, height = canvas.get_width_height()
            buffer = canvas.buffer_rgba()

            image = QImage(
                buffer,
                width,
                height,
                QImage.Format_RGBA8888,
            ).copy()

            pixmap = QPixmap.fromImage(image)

            return pixmap, ""

        except Exception as exc:
            return None, f"Errore rendering formula:\n{exc}"

    def refresh_preview(self):
        pixmap, error = self.render_equation_pixmap()

        if pixmap is None:
            self.preview_label.setPixmap(QPixmap())
            self.preview_label.setText(error)
            return

        scaled = pixmap.scaled(
            760,
            150,
            Qt.KeepAspectRatio,
            Qt.SmoothTransformation,
        )

        self.preview_label.setText("")
        self.preview_label.setPixmap(scaled)

    def refresh_equation_info(self):
        block_id = self.data.get("id", "")

        if not block_id:
            self.reference_label.setText("")
            return

        self.reference_label.setText(
            f"Incolla nel testo: [[EQ:{block_id}]]"
        )

    def build_summary(self):
        latex = self.latex_edit.toPlainText().strip()
        if latex:
            return f"{self.block_type_label} — {latex[:60]}"
        return self.block_type_label

    def export_data(self):
        data = {
            "type": "equation",
            "latex": self.latex_edit.toPlainText(),
            "caption": self.caption_edit.text(),
            "numbering_mode": (
                self.numbering_mode_combo.currentData()
                or "none"
            ),
        }

        data.update(self.export_common_data())
        return data

class TableBlockWidget(BaseBlockWidget):
    block_type_label = "Tabella"

    def __init__(self, parent_editor, data=None):
        base = {
            "type": "table",
            "rows": 3,
            "cols": 3,
            "data": [["" for _ in range(3)] for _ in range(3)],
            "spans": [],
            "collapsed": False,
        }
        base.update(data or {})
        super().__init__(parent_editor, base)

        rows = max(1, base.get("rows", 3))
        cols = max(1, base.get("cols", 3))
        table_data = base.get("data", [])

        top_row = QHBoxLayout()
        top_row.setContentsMargins(0, 0, 0, 0)

        self.btn_add_row = QPushButton("+ Riga")
        self.btn_del_row = QPushButton("- Riga")
        self.btn_add_col = QPushButton("+ Colonna")
        self.btn_del_col = QPushButton("- Colonna")
        self.btn_merge = QPushButton("Accorpa celle")
        self.btn_unmerge = QPushButton("Separa celle")
        self.btn_fit_cols = QPushButton("Adatta colonne")
        self.btn_fit_rows = QPushButton("Adatta righe")
        self.btn_align_left = QPushButton("SX")
        self.btn_align_center = QPushButton("Centro")
        self.btn_align_right = QPushButton("DX")
        self.btn_align_justify = QPushButton("Giustifica")

        self.btn_valign_top = QPushButton("Alto")
        self.btn_valign_middle = QPushButton("Medio")
        self.btn_valign_bottom = QPushButton("Basso")

        for btn in (
            self.btn_add_row,
            self.btn_del_row,
            self.btn_add_col,
            self.btn_del_col,
            self.btn_merge,
            self.btn_unmerge,
            self.btn_align_left,
            self.btn_align_center,
            self.btn_align_right,
            self.btn_align_justify,
            self.btn_valign_top,
            self.btn_valign_middle,
            self.btn_valign_bottom,
            self.btn_fit_cols,
            self.btn_fit_rows,
        ):
            style_small_button(btn)

        self.btn_add_row.clicked.connect(self.add_row)
        self.btn_del_row.clicked.connect(self.remove_row)
        self.btn_add_col.clicked.connect(self.add_column)
        self.btn_del_col.clicked.connect(self.remove_column)
        self.btn_merge.clicked.connect(self.merge_selected_cells)
        self.btn_unmerge.clicked.connect(self.unmerge_selected_cells)
        self.btn_align_left.clicked.connect(lambda: self.apply_cell_alignment("left"))
        self.btn_align_center.clicked.connect(lambda: self.apply_cell_alignment("center"))
        self.btn_align_right.clicked.connect(lambda: self.apply_cell_alignment("right"))
        self.btn_align_justify.clicked.connect(lambda: self.apply_cell_alignment("justify"))

        self.btn_valign_top.clicked.connect(lambda: self.apply_cell_vertical_alignment("top"))
        self.btn_valign_middle.clicked.connect(lambda: self.apply_cell_vertical_alignment("middle"))
        self.btn_valign_bottom.clicked.connect(lambda: self.apply_cell_vertical_alignment("bottom"))
        self.btn_fit_cols.clicked.connect(self.fit_columns)
        self.btn_fit_rows.clicked.connect(self.fit_rows)

        top_row.addWidget(self.btn_add_row)
        top_row.addWidget(self.btn_del_row)
        top_row.addWidget(self.btn_add_col)
        top_row.addWidget(self.btn_del_col)
        top_row.addWidget(self.btn_merge)
        top_row.addWidget(self.btn_unmerge)
        top_row.addWidget(self.btn_fit_cols)
        top_row.addWidget(self.btn_fit_rows)
        top_row.addWidget(self.btn_align_left)
        top_row.addWidget(self.btn_align_center)
        top_row.addWidget(self.btn_align_right)
        top_row.addWidget(self.btn_align_justify)
        top_row.addWidget(self.btn_valign_top)
        top_row.addWidget(self.btn_valign_middle)
        top_row.addWidget(self.btn_valign_bottom)
        top_row.addStretch()
        self.content_layout.addLayout(top_row)

        self.table = QTableWidget(rows, cols)
        self.table.setStyleSheet("""
            QTableWidget {
                background-color: white;
                alternate-background-color: white;
                gridline-color: #d0d0d0;
                selection-background-color: #cfe8ff;
                selection-color: black;
            }

            QTableWidget::item {
                background-color: white;
                color: black;
            }

            QHeaderView::section {
                background-color: #f2f2f2;
                color: black;
                border: 1px solid #d0d0d0;
                padding: 3px;
            }
        """)
        self.table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.table.setSizeAdjustPolicy(QAbstractScrollArea.AdjustIgnored)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Interactive)
        self.table.setMinimumHeight(190)

        # leggibilità testo lungo
        self.table.setWordWrap(True)
        self.table.setTextElideMode(Qt.ElideNone)
        self.table.verticalHeader().setDefaultSectionSize(32)
        self.table.verticalHeader().setSectionResizeMode(QHeaderView.Interactive)

        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Interactive)
        self.table.horizontalHeader().setStretchLastSection(False)

        # selezione per accorpamento celle
        self.table.setSelectionMode(QTableWidget.ContiguousSelection)
        self.table.setSelectionBehavior(QTableWidget.SelectItems)

        self.table.verticalHeader().setVisible(True)
        self.table.verticalHeader().setMinimumWidth(24)
        self.table.horizontalHeader().sectionDoubleClicked.connect(self.fit_single_column)
        self.table.verticalHeader().sectionDoubleClicked.connect(self.fit_single_row)
        self.table.itemChanged.connect(self.on_item_changed)

        self.content_layout.addWidget(self.table)

        for r in range(rows):
            for c in range(cols):
                value = ""
                if r < len(table_data) and c < len(table_data[r]):
                    value = table_data[r][c]

                item = QTableWidgetItem(value)
                item.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                self.table.setItem(r, c, item)

        self.cell_formats = base.get("cell_formats", {})
        self.spans = base.get("spans", [])
        
        self.apply_spans_to_table()
        self.apply_cell_formats_to_table()

        column_widths = base.get("column_widths", [])
        for c, width in enumerate(column_widths):
            if c < self.table.columnCount():
                self.table.setColumnWidth(c, int(width))

        row_heights = base.get("row_heights", [])
        for r, height in enumerate(row_heights):
            if r < self.table.rowCount():
                self.table.setRowHeight(r, int(height))

        self.update_summary()

    def on_item_changed(self, *_):
        self.table.resizeRowsToContents()
        self.update_summary()
        self.parent_editor.main_window.mark_dirty()

    def add_row(self):
        self.table.insertRow(self.table.rowCount())
        self.update_summary()
        self.parent_editor.main_window.mark_dirty()

    def remove_row(self):
        if self.table.rowCount() > 1:
            self.table.removeRow(self.table.rowCount() - 1)
            self.update_summary()
            self.parent_editor.main_window.mark_dirty()

    def add_column(self):
        self.table.insertColumn(self.table.columnCount())
        self.update_summary()
        self.parent_editor.main_window.mark_dirty()

    def remove_column(self):
        if self.table.columnCount() > 1:
            self.table.removeColumn(self.table.columnCount() - 1)
            self.update_summary()
            self.parent_editor.main_window.mark_dirty()

    def build_summary(self):
        return f"{self.block_type_label} — {self.table.rowCount()}x{self.table.columnCount()}"

    def apply_spans_to_table(self):
        self.table.clearSpans()

        for span in self.spans:
            r1 = span.get("r1", 0)
            c1 = span.get("c1", 0)
            r2 = span.get("r2", r1)
            c2 = span.get("c2", c1)

            row_span = r2 - r1 + 1
            col_span = c2 - c1 + 1

            if row_span > 1 or col_span > 1:
                self.table.setSpan(r1, c1, row_span, col_span)


    def merge_selected_cells(self):
        ranges = self.table.selectedRanges()

        if not ranges:
            return

        selected = ranges[0]

        r1 = selected.topRow()
        r2 = selected.bottomRow()
        c1 = selected.leftColumn()
        c2 = selected.rightColumn()

        if r1 == r2 and c1 == c2:
            return

        self.spans.append({
            "r1": r1,
            "c1": c1,
            "r2": r2,
            "c2": c2,
        })

        self.apply_spans_to_table()
        self.parent_editor.main_window.mark_dirty()


    def unmerge_selected_cells(self):
        row = self.table.currentRow()
        col = self.table.currentColumn()

        if row < 0 or col < 0:
            return

        kept = []

        for span in self.spans:
            r1 = span.get("r1", 0)
            c1 = span.get("c1", 0)
            r2 = span.get("r2", r1)
            c2 = span.get("c2", c1)

            inside = r1 <= row <= r2 and c1 <= col <= c2

            if not inside:
                kept.append(span)

        self.spans = kept
        self.apply_spans_to_table()
        self.parent_editor.main_window.mark_dirty()

    def _selected_cell_positions(self):
        positions = []

        for rng in self.table.selectedRanges():
            for r in range(rng.topRow(), rng.bottomRow() + 1):
                for c in range(rng.leftColumn(), rng.rightColumn() + 1):
                    positions.append((r, c))

        if not positions:
            row = self.table.currentRow()
            col = self.table.currentColumn()
            if row >= 0 and col >= 0:
                positions.append((row, col))

        return positions


    def _cell_key(self, row, col):
        return f"{row},{col}"


    def apply_cell_alignment(self, alignment):
        for row, col in self._selected_cell_positions():
            key = self._cell_key(row, col)
            fmt = self.cell_formats.get(key, {})
            fmt["halign"] = alignment
            self.cell_formats[key] = fmt

        self.apply_cell_formats_to_table()
        self.parent_editor.main_window.mark_dirty()


    def apply_cell_vertical_alignment(self, alignment):
        for row, col in self._selected_cell_positions():
            key = self._cell_key(row, col)
            fmt = self.cell_formats.get(key, {})
            fmt["valign"] = alignment
            self.cell_formats[key] = fmt

        self.apply_cell_formats_to_table()
        self.parent_editor.main_window.mark_dirty()


    def apply_cell_formats_to_table(self):
        for r in range(self.table.rowCount()):
            for c in range(self.table.columnCount()):
                item = self.table.item(r, c)
                if item is None:
                    continue

                fmt = self.cell_formats.get(self._cell_key(r, c), {})
                halign = fmt.get("halign", "left")
                valign = fmt.get("valign", "middle")

                if halign in ("center", "justify"):
                    hflag = Qt.AlignHCenter
                elif halign == "right":
                    hflag = Qt.AlignRight
                else:
                    hflag = Qt.AlignLeft

                if valign == "top":
                    vflag = Qt.AlignTop
                elif valign == "bottom":
                    vflag = Qt.AlignBottom
                else:
                    vflag = Qt.AlignVCenter

                item.setTextAlignment(hflag | vflag)

        self.table.resizeRowsToContents()

    def export_data(self):
        rows = self.table.rowCount()
        cols = self.table.columnCount()

        data = []
        for r in range(rows):
            row_data = []
            for c in range(cols):
                item = self.table.item(r, c)
                row_data.append(item.text() if item else "")
            data.append(row_data)

        column_widths = []
        for c in range(cols):
            column_widths.append(self.table.columnWidth(c))

        row_heights = []
        for r in range(rows):
            row_heights.append(self.table.rowHeight(r))

        result = {
            "type": "table",
            "rows": rows,
            "cols": cols,
            "data": data,
            "spans": deepcopy(self.spans),
            "cell_formats": deepcopy(self.cell_formats),
            "column_widths": column_widths,
            "row_heights": row_heights,
        }

        result.update(self.export_common_data())
        return result

    def fit_columns(self):
        self.table.resizeColumnsToContents()
        self.table.resizeRowsToContents()
    
    def fit_rows(self):
        self.table.resizeRowsToContents()
        self.parent_editor.main_window.mark_dirty()

    def fit_single_column(self, column):
        self.table.resizeColumnToContents(column)
        self.parent_editor.main_window.mark_dirty()

    def fit_single_row(self, row):
        self.table.resizeRowToContents(row)
        self.parent_editor.main_window.mark_dirty()

class BlocksNavigatorWidget(QListWidget):
    orderChanged = Signal(list)

    def __init__(self, parent=None):
        super().__init__(parent)
        self._drag_source_row = -1

        self.setSelectionMode(QAbstractItemView.SingleSelection)
        self.setDragEnabled(True)
        self.setAcceptDrops(True)
        self.setDropIndicatorShown(True)
        self.setDefaultDropAction(Qt.MoveAction)
        self.setDragDropMode(QAbstractItemView.DragDrop)
        self.setDragDropOverwriteMode(False)

    def startDrag(self, supported_actions):
        self._drag_source_row = self.currentRow()
        super().startDrag(supported_actions)

    def dropEvent(self, event):
        source_row = self._drag_source_row

        if source_row < 0 or source_row >= self.count():
            event.ignore()
            return

        pos = event.position().toPoint()
        target_row = self.indexAt(pos).row()

        if target_row < 0:
            target_row = self.count()
        else:
            rect = self.visualItemRect(self.item(target_row))
            if pos.y() > rect.center().y():
                target_row += 1

        if target_row > source_row:
            target_row -= 1

        if target_row == source_row:
            event.ignore()
            return

        item = self.takeItem(source_row)
        self.insertItem(target_row, item)
        self.setCurrentItem(item)

        ordered_ids = []

        for row in range(self.count()):
            nav_item = self.item(row)
            block_id = nav_item.data(Qt.UserRole)

            if block_id:
                ordered_ids.append(block_id)

        self.orderChanged.emit(ordered_ids)
        event.acceptProposedAction()

class BlockEditorWidget(QWidget):
    def __init__(self, main_window, parent=None):
        super().__init__(parent)
        self.main_window = main_window
        self.current_blocks_widgets = []

        root_layout = QVBoxLayout(self)
        root_layout.setContentsMargins(0, 0, 0, 0)
        root_layout.setSpacing(8)

        top_row = QHBoxLayout()
        self.btn_add_text = QPushButton("+ Testo")
        self.btn_add_image = QPushButton("+ Immagine")
        self.btn_add_images = QPushButton("+ GRUPPO DI IMMAGINI")
        self.btn_add_equation = QPushButton("+ Equazione")
        self.btn_add_table = QPushButton("+ Tabella")
        for btn in (
            self.btn_add_text,
            self.btn_add_image,
            self.btn_add_images,
            self.btn_add_equation,
            self.btn_add_table
        ):
            btn.setFixedHeight(30)
            btn.setStyleSheet(
                "QPushButton { border: 1px solid #d6d6d6; border-radius: 5px; background: #fafafa; padding: 0 10px; }"
                "QPushButton:hover { background: #f0f0f0; }"
            )
        self.btn_expand_all = QPushButton("Espandi tutti")
        self.btn_collapse_all = QPushButton("Collassa tutti")
        self.btn_preview_node = QPushButton("Anteprima pagina")
        for btn in (self.btn_expand_all, self.btn_collapse_all, self.btn_preview_node):
            btn.setFixedHeight(30)
            btn.setStyleSheet(
                "QPushButton { border: 1px solid #d6d6d6; border-radius: 5px; background: #fafafa; padding: 0 10px; }"
                "QPushButton:hover { background: #f0f0f0; }"
            )

        top_row.addWidget(self.btn_add_text)
        top_row.addWidget(self.btn_add_image)
        top_row.addWidget(self.btn_add_images)
        top_row.addWidget(self.btn_add_equation)
        top_row.addWidget(self.btn_add_table)
        top_row.addWidget(self.btn_expand_all)
        top_row.addWidget(self.btn_collapse_all)
        top_row.addWidget(self.btn_preview_node)
        top_row.addStretch()
        root_layout.addLayout(top_row)

        self.btn_add_text.clicked.connect(self.add_text_block)
        self.btn_add_image.clicked.connect(self.add_image_block)
        self.btn_add_images.clicked.connect(self.add_multi_image_block)
        self.btn_add_equation.clicked.connect(self.add_equation_block)
        self.btn_add_table.clicked.connect(self.add_table_block)
        self.btn_expand_all.clicked.connect(self.expand_all_blocks)
        self.btn_collapse_all.clicked.connect(self.collapse_all_blocks)
        self.btn_preview_node.clicked.connect(self.main_window.show_node_preview)

        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setFrameShape(QFrame.NoFrame)
        root_layout.addWidget(self.scroll_area)

        self.container = QWidget()
        self.container.setStyleSheet("background: transparent;")
        self.blocks_layout = QVBoxLayout(self.container)
        self.blocks_layout.setContentsMargins(0, 0, 0, 0)
        self.blocks_layout.setSpacing(10)
        self.blocks_layout.addStretch()
        self.scroll_area.setWidget(self.container)

    def set_enabled_for_node(self, enabled):
        for btn in (
            self.btn_add_text,
            self.btn_add_image,
            self.btn_add_images,
            self.btn_add_equation,
            self.btn_add_table
        ):
            btn.setEnabled(enabled)
        self.scroll_area.setEnabled(enabled)

    def clear_blocks(self):
        for widget in list(self.current_blocks_widgets):
            self.blocks_layout.removeWidget(widget)
            widget.deleteLater()
        self.current_blocks_widgets.clear()

    def load_blocks(self, blocks_data):
        self.clear_blocks()
        for block_data in blocks_data or []:
            t = block_data.get("type")
            if t == "text":
                self.add_text_block(block_data)
            elif t == "image":
                self.add_image_block(block_data)
            elif t == "images":
                self.add_multi_image_block(block_data)
            elif t == "equation":
                self.add_equation_block(block_data)
            elif t == "table":
                self.add_table_block(block_data)
        self.main_window.update_figure_references()

    def export_blocks(self):
        return [w.export_data() for w in self.current_blocks_widgets]

    def iter_widgets(self):
        return list(self.current_blocks_widgets)

    def duplicate_block(self, widget):
        try:
            index = self.current_blocks_widgets.index(widget)
        except ValueError:
            return
        data = deepcopy(widget.export_data())
        data = self.main_window.clone_block_data(data)
        block_type = data.get("type")
        if block_type == "text":
            new_widget = TextBlockWidget(self, data)
        elif block_type == "image":
            new_widget = ImageBlockWidget(self, data)
        elif block_type == "images":
            new_widget = MultiImageBlockWidget(self, data)
        elif block_type == "equation":
            new_widget = EquationBlockWidget(self, data)
        elif block_type == "table":
            new_widget = TableBlockWidget(self, data)
        else:
            return
        self.current_blocks_widgets.insert(index + 1, new_widget)
        self.blocks_layout.insertWidget(index + 1, new_widget)
        self.main_window.mark_dirty()
        self.main_window.refresh_blocks_navigator()
        self.main_window.update_figure_references()

    def add_text_block(self, data=None):
        w = TextBlockWidget(self, data)
        self.current_blocks_widgets.append(w)
        self.blocks_layout.insertWidget(self.blocks_layout.count() - 1, w)
        self.main_window.mark_dirty()
        self.main_window.refresh_blocks_navigator()

    def add_image_block(self, data=None):
        w = ImageBlockWidget(self, data)
        self.current_blocks_widgets.append(w)
        self.blocks_layout.insertWidget(self.blocks_layout.count() - 1, w)
        self.main_window.mark_dirty()
        self.main_window.refresh_blocks_navigator()

    def insert_image_block_after(self, reference_widget, image_path):
        try:
            index = self.current_blocks_widgets.index(reference_widget)
        except ValueError:
            self.add_image_block({
                "type": "image",
                "path": image_path,
                "caption": "",
                "collapsed": False,
            })
            return

        data = {
            "type": "image",
            "path": image_path,
            "caption": "",
            "collapsed": False,
        }

        w = ImageBlockWidget(self, data)
        self.current_blocks_widgets.insert(index + 1, w)
        self.blocks_layout.insertWidget(index + 1, w)

        self.main_window.mark_dirty()
        self.main_window.update_figure_references()

    def add_multi_image_block(self, data=None):
        w = MultiImageBlockWidget(self, data)
        self.current_blocks_widgets.append(w)
        self.blocks_layout.insertWidget(self.blocks_layout.count() - 1, w)
        self.main_window.mark_dirty()
        self.main_window.refresh_blocks_navigator()

    def add_equation_block(self, data=None):
        w = EquationBlockWidget(self, data)
        self.current_blocks_widgets.append(w)
        self.blocks_layout.insertWidget(self.blocks_layout.count() - 1, w)
        self.main_window.mark_dirty()

    def add_table_block(self, data=None):
        w = TableBlockWidget(self, data)
        self.current_blocks_widgets.append(w)
        self.blocks_layout.insertWidget(self.blocks_layout.count() - 1, w)
        self.main_window.mark_dirty()
        self.main_window.refresh_blocks_navigator()

    def expand_all_blocks(self):
        for w in self.current_blocks_widgets:
            w.set_collapsed_state(False)
        self.main_window.mark_dirty()

    def collapse_all_blocks(self):
        for w in self.current_blocks_widgets:
            w.set_collapsed_state(True)
        self.main_window.mark_dirty()

    def move_block(self, widget, direction):
        try:
            current_index = self.current_blocks_widgets.index(widget)
        except ValueError:
            return
        new_index = current_index + direction
        if new_index < 0 or new_index >= len(self.current_blocks_widgets):
            return
        self.current_blocks_widgets.pop(current_index)
        self.current_blocks_widgets.insert(new_index, widget)
        self.blocks_layout.removeWidget(widget)
        self.blocks_layout.insertWidget(new_index, widget)
        self.main_window.mark_dirty()
        self.main_window.refresh_blocks_navigator()
        self.main_window.update_figure_references()

    def remove_block(self, widget):
        try:
            self.current_blocks_widgets.remove(widget)
        except ValueError:
            return
        self.blocks_layout.removeWidget(widget)
        widget.deleteLater()
        self.main_window.mark_dirty()
        self.main_window.refresh_blocks_navigator()
        self.main_window.update_figure_references()

class DynamicTocEndBreak(ActionFlowable):
    def apply(self, doc):
        toc_end_page = doc.page

        blank_pages_after_toc = 1 if toc_end_page % 2 == 1 else 2
        content_start_page = toc_end_page + blank_pages_after_toc + 1

        doc.toc_end_page = toc_end_page
        doc.content_start_page = content_start_page

        blank_pages = set(getattr(doc, "pdf_blank_pages", set()))
        blank_pages.add(2)

        for page in range(toc_end_page + 1, content_start_page):
            blank_pages.add(page)

        doc.pdf_blank_pages = blank_pages

        for _ in range(blank_pages_after_toc + 1):
            doc.handle_pageBreak()

class TechnicalDocTemplate(BaseDocTemplate):
    def __init__(self, filename, owner, **kwargs):
        self.owner = owner
        self.previous_chapter_starts = {}
        self.current_build_chapter_starts = {}
        super().__init__(filename, **kwargs)
        frame = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id="normal")
        template = PageTemplate(id="main", frames=[frame], onPage=self._on_page)
        self.addPageTemplates([template])

    def beforeDocument(self):
        self.previous_chapter_starts = getattr(self, "current_build_chapter_starts", {}) or {}
        self.current_build_chapter_starts = {}

        # Stato PDF dinamico: va azzerato a ogni passaggio di multiBuild.
        # Altrimenti le pagine bianche calcolate nel passaggio precedente
        # vengono riusate mentre il sommario è ancora in costruzione.
        self.pdf_blank_pages = {2}
        self.content_start_page = None
        self.toc_end_page = None

    def afterFlowable(self, flowable):
        if hasattr(flowable, "toc_level") and hasattr(flowable, "toc_text"):
            level = flowable.toc_level
            text = flowable.toc_text
            content_start_page = getattr(self, "content_start_page", None)

            if content_start_page is None:
                content_start_page = self.page

            display_page = self.page - content_start_page + 1
            self.notify("TOCEntry", (level, text, display_page))
            if level == 0:
                self.current_build_chapter_starts[self.page] = self.owner._extract_base_title(text)

    def _chapter_for_page(self, physical_page: int) -> str:
        chapter_map = self.previous_chapter_starts or self.current_build_chapter_starts
        if not chapter_map:
            return ""
        starts = sorted(chapter_map.keys())
        current = ""
        for page in starts:
            if page <= physical_page:
                current = chapter_map[page]
            else:
                break
        return current

    def _draw_layout_block(self, canvas, cfg, x_left, x_right, y_top, max_height_cm=2.0, align="center"):
        if not cfg:
            return 0

        block_type = cfg.get("type", "none")
        value = cfg.get("value", "")

        # ===== TESTO =====
        if block_type == "text":
            html = cfg.get("html", "")
            lines = []

            if html:
                qdoc = QTextDocument()
                qdoc.setHtml(html)

                block = qdoc.begin()
                while block.isValid():
                    fragments = []
                    it = block.begin()

                    while not it.atEnd():
                        fragment = it.fragment()
                        if fragment.isValid():
                            txt = fragment.text()
                            if txt:
                                char_fmt = fragment.charFormat()

                                bold = char_fmt.fontWeight() >= 700
                                italic = char_fmt.fontItalic()
                                underline = char_fmt.fontUnderline()

                                size = char_fmt.fontPointSize()
                                if size <= 0:
                                    size = 9

                                if bold and italic:
                                    font_name = "Helvetica-BoldOblique"
                                elif bold:
                                    font_name = "Helvetica-Bold"
                                elif italic:
                                    font_name = "Helvetica-Oblique"
                                else:
                                    font_name = "Helvetica"

                                fragments.append({
                                    "text": txt,
                                    "font": font_name,
                                    "size": size,
                                    "underline": underline,
                                })

                        it += 1

                    if fragments:
                        lines.append(fragments)

                    block = block.next()

            else:
                value = str(value).strip()
                if value:
                    for line in value.splitlines():
                        line = line.rstrip()
                        if line.strip():
                            lines.append([{
                                "text": line,
                                "font": "Helvetica",
                                "size": 9,
                                "underline": False,
                            }])

            if not lines:
                return 0

            used_h = 0
            list_indent = 0

            for line_fragments in lines:
                line_height = max(f["size"] for f in line_fragments) + 2
                y = y_top - used_h

                total_w = 0
                for f in line_fragments:
                    total_w += canvas.stringWidth(f["text"], f["font"], f["size"])

                # rileva indentazione da spazi iniziali
                first_text = line_fragments[0]["text"]
                leading_spaces = len(first_text) - len(first_text.lstrip(" "))
                indent_level = leading_spaces // 4  # ogni 4 spazi = livello

                indent_px = indent_level * 10
                x_left_indented = x_left + indent_px

                if align == "right":
                    x = x_right - total_w
                elif align == "left":
                    x = x_left_indented
                else:
                    x = (x_left + x_right - total_w) / 2

                for f in line_fragments:
                    canvas.setFont(f["font"], f["size"])
                    canvas.drawString(x, y, f["text"])

                    text_w = canvas.stringWidth(f["text"], f["font"], f["size"])
                    if f["underline"]:
                        canvas.line(x, y - 1.5, x + text_w, y - 1.5)

                    x += text_w

                used_h += line_height

            return used_h

        # ===== IMMAGINE =====
        elif block_type == "image" and value and os.path.exists(value):
            try:
                pix = QPixmap(value)
                if pix.isNull():
                    return 0

                max_h = max_height_cm * cm
                iw = pix.width()
                ih = pix.height()
                if iw <= 0 or ih <= 0:
                    return 0

                scale = min(max_h / ih, 1.0)
                draw_w = iw * scale
                draw_h = ih * scale

                if align == "right":
                    x = x_right - draw_w
                elif align == "left":
                    x = x_left
                else:
                    x = (x_left + x_right - draw_w) / 2

                canvas.drawImage(
                    value,
                    x,
                    y_top - draw_h + 4,
                    width=draw_w,
                    height=draw_h,
                    preserveAspectRatio=True,
                    mask='auto'
                )
                return draw_h

            except Exception:
                return 0

        return 0

    def _layout_text_lines(self, cfg):
        if not cfg:
            return []
        if cfg.get("type") != "text":
            return []
        value = str(cfg.get("value", "")).strip()
        if not value:
            return []
        return [line.rstrip() for line in value.splitlines() if line.strip()]

    def _layout_block_height_points(self, cfg, max_height_cm=2.0):
        if not cfg:
            return 0

        block_type = cfg.get("type", "none")
        value = cfg.get("value", "")

        if block_type == "text":
            lines = self._layout_text_lines(cfg)
            if not lines:
                return 0
            line_height = 10
            return len(lines) * line_height

        if block_type == "image" and value and os.path.exists(value):
            try:
                pix = QPixmap(value)
                if pix.isNull():
                    return 0

                max_h = max_height_cm * cm
                iw = pix.width()
                ih = pix.height()
                if iw <= 0 or ih <= 0:
                    return 0

                scale = min(max_h / ih, 1.0)
                draw_h = ih * scale
                return draw_h
            except Exception:
                return 0

        return 0

    def _on_page(self, canvas, doc):
        physical_page = canvas.getPageNumber()

        blank_pages = getattr(doc, "pdf_blank_pages", set())

        if physical_page in blank_pages:
            return

        content_start_page = getattr(doc, "content_start_page", None)

        page_w, page_h = A4
        left = doc.leftMargin
        right = page_w - doc.rightMargin

        is_cover = physical_page == 1
        is_content = (
            content_start_page is not None
            and physical_page >= content_start_page
        )

        canvas.setStrokeColor(colors.black)
        canvas.setLineWidth(0.8)
        canvas.setFont("Helvetica", 9)

        if is_cover:
            section_name = "cover"
        else:
            section_name = "content"

        header_cfg = self.owner.layout_config.get(
            section_name, {}
        ).get(
            "header",
            {"type": "none", "value": "", "align": "right"}
        )

        header_align = header_cfg.get("align", "right")
        custom_h = self._layout_block_height_points(header_cfg, max_height_cm=2.0)

        if is_cover:
            y_header_custom_top = page_h - doc.topMargin + 18 + custom_h
        else:
            y_header_custom_top = page_h - doc.topMargin + 28 + custom_h

        used_h = self._draw_layout_block(
            canvas,
            header_cfg,
            left,
            right,
            y_header_custom_top,
            max_height_cm=2.0,
            align=header_align
        )

        if is_content:
            chapter_y = y_header_custom_top - used_h - 6 if used_h > 0 else page_h - doc.topMargin + 14
            header_text = self._chapter_for_page(physical_page)
            canvas.setFont("Helvetica", 9)
            canvas.drawRightString(right, chapter_y, header_text)
            line_y = chapter_y - 4
        else:
            line_y = y_header_custom_top - used_h - 6 if used_h > 0 else page_h - doc.topMargin + 0.6 * cm

        canvas.line(left, line_y, right, line_y)

        y_footer_line = doc.bottomMargin - 0.35 * cm
        y_footer_custom_top = y_footer_line - 14
        y_footer_page = y_footer_line - 30

        canvas.line(left, y_footer_line, right, y_footer_line)

        footer_cfg = self.owner.layout_config.get(
            section_name, {}
        ).get(
            "footer",
            {"type": "none", "value": "", "align": "center"}
        )

        footer_align = footer_cfg.get("align", "center")

        self._draw_layout_block(
            canvas,
            footer_cfg,
            left,
            right,
            y_footer_custom_top,
            max_height_cm=2.0,
            align=footer_align
        )

        if is_content:
            page_no = physical_page - content_start_page + 1
            canvas.drawCentredString(page_w / 2, y_footer_page, f"- {page_no} -")

class LayoutDialog(QDialog):
    def __init__(self, layout_config, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Layout pagina")
        self.resize(600, 500)

        self.layout_config = layout_config

        root = QVBoxLayout(self)

        self.tabs = QTabWidget()
        root.addWidget(self.tabs)

        self.tabs.addTab(self.build_tab("cover"), "Copertina")
        self.tabs.addTab(self.build_tab("content"), "Contenuto")

        btn_row = QHBoxLayout()
        btn_ok = QPushButton("OK")
        btn_cancel = QPushButton("Annulla")
        btn_ok.clicked.connect(self.accept)
        btn_cancel.clicked.connect(self.reject)
        btn_row.addStretch()
        btn_row.addWidget(btn_ok)
        btn_row.addWidget(btn_cancel)
        root.addLayout(btn_row)

    def build_tab(self, section):
        widget = QWidget()
        widget.setStyleSheet("QWidget { background: palette(window); }")
        layout = QVBoxLayout(widget)

        layout.addWidget(QLabel("HEADER"))
        layout.addLayout(self.build_block(section, "header"))

        layout.addSpacing(10)

        layout.addWidget(QLabel("FOOTER"))
        layout.addLayout(self.build_block(section, "footer"))

        layout.addStretch()
        return widget

    def build_block(self, section, part):
        container = QVBoxLayout()

        combo = QComboBox()
        combo.addItems(["none", "text", "image"])
        combo.setStyleSheet("background: white;")      

        align_combo = QComboBox()
        align_combo.addItems(["left", "center", "right"])
        align_combo.setStyleSheet("background: white;")

        text_edit = QTextEdit()
        text_edit.setStyleSheet("background: white;")
        text_edit.setPlaceholderText("Inserisci testo...")
        text_edit.setMinimumHeight(80)
        text_edit.setAcceptRichText(True)
        text_edit.viewport().setAutoFillBackground(True)

        btn_image = QPushButton("Seleziona immagine")
        image_label = QLabel("")
        image_label.setStyleSheet("color: #666;")
        image_label.setWordWrap(True)
        format_row = QHBoxLayout()

        btn_bold = QPushButton("B")
        btn_italic = QPushButton("I")
        btn_underline = QPushButton("U")

        font_size_combo = QComboBox()
        font_size_combo.addItems(["8", "9", "10", "11", "12", "14", "16"])
        font_size_combo.setCurrentText("9")
        font_size_combo.setFixedWidth(70)

        for btn in (btn_bold, btn_italic, btn_underline):
            style_icon_button(btn)
            format_row.addWidget(btn)

        format_row.addWidget(QLabel("Dim."))
        format_row.addWidget(font_size_combo)
        format_row.addStretch()

        default_align = "right" if part == "header" else "center"

        current_cfg = self.layout_config.get(section, {}).get(
            part,
            {"type": "none", "value": "", "align": default_align}
        )
        current_type = current_cfg.get("type", "none")
        current_value = current_cfg.get("value", "")
        current_align = current_cfg.get("align", default_align)

        combo.setCurrentText(current_type)
        align_combo.setCurrentText(current_align)

        if current_type == "text":
            current_html = current_cfg.get("html", "")
            if current_html:
                text_edit.setHtml(current_html)
            else:
                text_edit.setPlainText(current_value)
        elif current_type == "image":
            image_label.setText(current_value if current_value else "Nessuna immagine selezionata")

        def update_ui():
            selected = combo.currentText()
            align = align_combo.currentText()

            if align == "left":
                text_edit.setAlignment(Qt.AlignLeft)
            elif align == "right":
                text_edit.setAlignment(Qt.AlignRight)
            else:
                text_edit.setAlignment(Qt.AlignCenter)

            text_edit.setReadOnly(selected != "text")
            btn_image.setEnabled(selected == "image")
            align_combo.setEnabled(selected in ("text", "image"))

            if selected == "none":
                self.layout_config[section][part] = {
                    "type": "none",
                    "value": "",
                    "align": align
                }
                image_label.setText("")

            elif selected == "text":
                self.layout_config[section][part] = {
                    "type": "text",
                    "value": text_edit.toPlainText(),
                    "html": text_edit.toHtml(),
                    "align": align
                }
                image_label.setText("")

            elif selected == "image":
                existing = self.layout_config[section][part].get("value", "")
                self.layout_config[section][part] = {
                    "type": "image",
                    "value": existing,
                    "align": align
                }
                image_label.setText(existing if existing else "Nessuna immagine selezionata")

        def load_image():
            path, _ = QFileDialog.getOpenFileName(
                self,
                "Seleziona immagine",
                "",
                "Immagini (*.png *.jpg *.jpeg *.bmp *.gif *.tif *.tiff)"
            )
            if path:
                self.layout_config[section][part] = {
                    "type": "image",
                    "value": path,
                    "align": align_combo.currentText()
                }
                image_label.setText(path)

        def on_text_changed():
            if combo.currentText() == "text":
                self.layout_config[section][part] = {
                    "type": "text",
                    "value": text_edit.toPlainText(),
                    "html": text_edit.toHtml(),
                    "align": align_combo.currentText()
                }

        def apply_char_format(fmt):
            text_edit.mergeCurrentCharFormat(fmt)
            text_edit.setFocus()
            on_text_changed()

        def toggle_bold():
            fmt = text_edit.currentCharFormat()
            fmt.setFontWeight(400 if fmt.fontWeight() >= 700 else 700)
            apply_char_format(fmt)

        def toggle_italic():
            fmt = text_edit.currentCharFormat()
            fmt.setFontItalic(not fmt.fontItalic())
            apply_char_format(fmt)

        def toggle_underline():
            fmt = text_edit.currentCharFormat()
            fmt.setFontUnderline(not fmt.fontUnderline())
            apply_char_format(fmt)

        def apply_font_size(value):
            try:
                size = float(value)
            except ValueError:
                return
            fmt = text_edit.currentCharFormat()
            fmt.setFontPointSize(size)
            apply_char_format(fmt)

        combo.currentTextChanged.connect(update_ui)
        align_combo.currentTextChanged.connect(update_ui)
        text_edit.textChanged.connect(on_text_changed)
        btn_image.clicked.connect(load_image)

        btn_bold.clicked.connect(toggle_bold)
        btn_italic.clicked.connect(toggle_italic)
        btn_underline.clicked.connect(toggle_underline)
        font_size_combo.currentTextChanged.connect(apply_font_size)

        container.addWidget(QLabel("Tipo"))
        container.addWidget(combo)

        container.addWidget(QLabel("Allineamento"))
        container.addWidget(align_combo)

        container.addLayout(format_row)
        container.addWidget(text_edit)
        container.addWidget(btn_image)
        container.addWidget(image_label)

        update_ui()
        return container

class NewRelationDialog(QDialog):
    def __init__(self, workspace_dir, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Nuova relazione")
        self.resize(520, 260)

        self.workspace_dir = workspace_dir

        layout = QVBoxLayout(self)
        info_label = QLabel(
            "Seleziona un valore esistente oppure "
            "digita un nuovo valore:\n\n"
            "1. Cliente\n"
            "2. Commessa\n"
            "3. Nome relazione\n\n"
            "Il nome della relazione non deve "
            "essere già presente nella commessa."
        )

        info_label.setWordWrap(True)

        info_label.setStyleSheet(
            "color: #404040; padding-bottom: 6px;"
        )

        layout.addWidget(info_label)

        # --- Cliente ---
        layout.addWidget(QLabel("Cliente"))
        self.client_combo = QComboBox()
        self.client_combo.setEditable(True)
        self.client_combo.setInsertPolicy(QComboBox.NoInsert)
        self.client_combo.setPlaceholderText("Seleziona o scrivi nuovo cliente")
        layout.addWidget(self.client_combo)

        # --- Commessa ---
        layout.addWidget(QLabel("Commessa"))
        self.project_combo = QComboBox()
        self.project_combo.setEditable(True)
        self.project_combo.setInsertPolicy(QComboBox.NoInsert)
        self.project_combo.setPlaceholderText("Seleziona o scrivi nuova commessa")
        layout.addWidget(self.project_combo)
        self.client_combo.setStyleSheet("QComboBox { padding: 2px; }")
        self.project_combo.setStyleSheet("QComboBox { padding: 2px; }")

        # --- Relazione ---
        layout.addWidget(QLabel("Nome relazione"))
        self.relation_combo = QComboBox()
        self.relation_combo.setEditable(True)
        self.relation_combo.setInsertPolicy(QComboBox.NoInsert)
        self.relation_combo.setPlaceholderText("Seleziona o scrivi nuova relazione")
        layout.addWidget(self.relation_combo)

        # --- Bottoni ---
        btn_row = QHBoxLayout()
        btn_ok = QPushButton("Crea")
        btn_cancel = QPushButton("Annulla")

        btn_ok.clicked.connect(self.accept)
        btn_cancel.clicked.connect(self.reject)

        btn_row.addStretch()
        btn_row.addWidget(btn_ok)
        btn_row.addWidget(btn_cancel)

        layout.addStretch()
        layout.addLayout(btn_row)

        # --- Signal ---
        self.client_combo.currentTextChanged.connect(self.refresh_projects)
        self.project_combo.currentTextChanged.connect(self.refresh_relations)

        # --- Init ---
        self.refresh_clients()

    # -------------------------

    def refresh_clients(self):
        self.client_combo.clear()

        if not self.workspace_dir or not os.path.exists(self.workspace_dir):
            return

        work_archive_dir = os.path.join(
            self.workspace_dir,
            "ARCHIVIO LAVORI"
        )

        if not os.path.exists(work_archive_dir):
            return

        clients = [
            name for name in os.listdir(work_archive_dir)
            if os.path.isdir(
                os.path.join(work_archive_dir, name)
            )
        ]

        clients = sorted(clients)
        self.client_combo.addItems(clients)

        self.client_combo.setCurrentIndex(-1)
        self.client_combo.clearEditText()

        self.refresh_projects()

    # -------------------------

    def refresh_projects(self, *_):
        self.project_combo.clear()

        cliente = self.client_combo.currentText().strip()
        if not cliente:
            return

        client_dir = os.path.join(
            self.workspace_dir,
            "ARCHIVIO LAVORI",
            cliente
        )

        if not os.path.exists(client_dir):
            return

        projects = [
            name for name in os.listdir(client_dir)
            if os.path.isdir(os.path.join(client_dir, name))
        ]

        projects = sorted(projects)
        self.project_combo.addItems(projects)

        self.refresh_relations()
    # -------------------------

    def refresh_relations(self, *_):
        self.relation_combo.clear()

        cliente = self.client_combo.currentText().strip()
        commessa = self.project_combo.currentText().strip()

        if not cliente or not commessa:
            return

        relations_dir = os.path.join(
            self.workspace_dir,
            "ARCHIVIO LAVORI",
            cliente,
            commessa,
            "Relazioni tecniche"
        )

        if not os.path.exists(relations_dir):
            return

        relations = [
            name for name in os.listdir(relations_dir)
            if os.path.isdir(
                os.path.join(relations_dir, name)
            )
        ]

        self.relation_combo.addItems(
            sorted(relations)
        )

    def values(self):
        return {
            "cliente": self.client_combo.currentText().strip(),
            "commessa": self.project_combo.currentText().strip(),
            "relazione": self.relation_combo.currentText().strip(),
        }

    # -------------------------

    def accept(self):
        data = self.values()

        if not data["cliente"] or not data["commessa"] or not data["relazione"]:
            QMessageBox.warning(
                self,
                "Dati mancanti",
                "Compila cliente, commessa e nome relazione."
            )
            return

        checks = [
            ("Cliente", data["cliente"]),
            ("Commessa", data["commessa"]),
            ("Relazione", data["relazione"]),
        ]

        for label, value in checks:
            ok, message = validate_folder_name(value)
            if not ok:
                QMessageBox.warning(
                    self,
                    "Nome non valido",
                    f"{label}: {message}"
                )
                return

        super().accept()

class SaveAsRelationDialog(NewRelationDialog):
    def __init__(self, workspace_dir, parent=None):
        super().__init__(workspace_dir, parent)
        self.setWindowTitle("Salva con Nome")

class OpenRelationDialog(QDialog):
    def __init__(self, workspace_dir, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Apri relazione")
        self.resize(520, 260)

        self.workspace_dir = workspace_dir

        layout = QVBoxLayout(self)

        layout.addWidget(QLabel("Cliente"))
        self.client_combo = QComboBox()
        layout.addWidget(self.client_combo)

        layout.addWidget(QLabel("Commessa"))
        self.project_combo = QComboBox()
        layout.addWidget(self.project_combo)

        layout.addWidget(QLabel("Relazione"))
        self.relation_combo = QComboBox()
        layout.addWidget(self.relation_combo)

        btn_row = QHBoxLayout()
        btn_ok = QPushButton("Apri")
        btn_cancel = QPushButton("Annulla")
        btn_ok.clicked.connect(self.accept)
        btn_cancel.clicked.connect(self.reject)

        btn_row.addStretch()
        btn_row.addWidget(btn_ok)
        btn_row.addWidget(btn_cancel)

        layout.addStretch()
        layout.addLayout(btn_row)

        self.client_combo.currentTextChanged.connect(self.refresh_projects)
        self.project_combo.currentTextChanged.connect(self.refresh_relations)

        self.refresh_clients()

    def refresh_clients(self):
        self.client_combo.clear()

        if not self.workspace_dir or not os.path.exists(self.workspace_dir):
            return

        work_archive_dir = os.path.join(
            self.workspace_dir,
            "ARCHIVIO LAVORI"
        )

        if not os.path.exists(work_archive_dir):
            return

        clients = [
            name for name in os.listdir(work_archive_dir)
            if os.path.isdir(
                os.path.join(work_archive_dir, name)
            )
        ]

        self.client_combo.addItems(sorted(clients))

        self.client_combo.setCurrentIndex(-1)
        self.client_combo.clearEditText()

        self.refresh_projects()

    def refresh_projects(self):
        self.project_combo.clear()
        self.relation_combo.clear()

        cliente = self.client_combo.currentText().strip()

        if not cliente:
            return

        client_dir = os.path.join(
            self.workspace_dir,
            "ARCHIVIO LAVORI",
            cliente
        )

        if not os.path.exists(client_dir):
            return

        projects = [
            name for name in os.listdir(client_dir)
            if os.path.isdir(
                os.path.join(client_dir, name)
            )
        ]

        self.project_combo.addItems(sorted(projects))

    def refresh_relations(self):
        self.relation_combo.clear()

        cliente = self.client_combo.currentText().strip()
        commessa = self.project_combo.currentText().strip()

        if not cliente or not commessa:
            return

        relations_dir = os.path.join(
            self.workspace_dir,
            "ARCHIVIO LAVORI",
            cliente,
            commessa,
            "Relazioni tecniche"
        )

        if not os.path.exists(relations_dir):
            return

        relations = [
            name for name in os.listdir(relations_dir)
            if os.path.isdir(
                os.path.join(relations_dir, name)
            )
        ]

        self.relation_combo.addItems(sorted(relations))

    def selected_document_path(self):
        cliente = self.client_combo.currentText().strip()
        commessa = self.project_combo.currentText().strip()
        relazione = self.relation_combo.currentText().strip()

        if not cliente or not commessa or not relazione:
            return ""

        return os.path.join(
            self.workspace_dir,
            "ARCHIVIO LAVORI",
            cliente,
            commessa,
            "Relazioni tecniche",
            relazione,
            "document.json"
        )

    def accept(self):
        document_path = self.selected_document_path()

        if not document_path or not os.path.exists(document_path):
            QMessageBox.warning(
                self,
                "Relazione non valida",
                "Seleziona una relazione valida."
            )
            return

        super().accept()

class AttachmentsDialog(QDialog):
    def __init__(self, allegati_dir, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Gestisci allegati")
        self.resize(720, 420)

        self.allegati_dir = allegati_dir
        os.makedirs(self.allegati_dir, exist_ok=True)
        self.setAcceptDrops(True)

        layout = QVBoxLayout(self)

        btn_row = QHBoxLayout()

        self.btn_add = QPushButton("+ Aggiungi file")
        self.btn_open = QPushButton("Apri")
        self.btn_remove = QPushButton("Rimuovi")
        self.btn_refresh = QPushButton("Aggiorna")

        btn_row.addWidget(self.btn_add)
        btn_row.addWidget(self.btn_open)
        btn_row.addWidget(self.btn_remove)
        btn_row.addWidget(self.btn_refresh)
        btn_row.addStretch()

        layout.addLayout(btn_row)

        self.table = QTableWidget(0, 2)
        self.table.setHorizontalHeaderLabels(["Nome file", "Dimensione"])
        self.table.setSelectionBehavior(QTableWidget.SelectRows)
        self.table.setSelectionMode(QTableWidget.SingleSelection)
        self.table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)

        layout.addWidget(self.table)

        close_row = QHBoxLayout()
        btn_close = QPushButton("Chiudi")
        btn_close.clicked.connect(self.accept)
        close_row.addStretch()
        close_row.addWidget(btn_close)
        layout.addLayout(close_row)

        self.btn_add.clicked.connect(self.add_files)
        self.btn_open.clicked.connect(self.open_selected)
        self.btn_remove.clicked.connect(self.remove_selected)
        self.btn_refresh.clicked.connect(self.refresh)
        self.table.doubleClicked.connect(self.open_selected)

        self.refresh()

    def format_size(self, size):
        if size < 1024:
            return f"{size} B"
        if size < 1024 * 1024:
            return f"{size / 1024:.1f} KB"
        return f"{size / (1024 * 1024):.1f} MB"

    def selected_file_path(self):
        row = self.table.currentRow()
        if row < 0:
            return ""
        item = self.table.item(row, 0)
        if item is None:
            return ""
        return os.path.join(self.allegati_dir, item.text())

    def refresh(self):
        self.table.setRowCount(0)

        files = []
        for name in os.listdir(self.allegati_dir):
            path = os.path.join(self.allegati_dir, name)
            if os.path.isfile(path):
                files.append((name, os.path.getsize(path)))

        for row, (name, size) in enumerate(sorted(files)):
            self.table.insertRow(row)
            self.table.setItem(row, 0, QTableWidgetItem(name))
            self.table.setItem(row, 1, QTableWidgetItem(self.format_size(size)))

    def copy_files_to_attachments(self, paths):
        for source in paths:
            filename = os.path.basename(source)
            dest = os.path.join(self.allegati_dir, filename)

            if os.path.exists(dest):
                QMessageBox.warning(
                    self,
                    "File già presente",
                    f"Il file esiste già negli allegati:\n{filename}"
                )
                continue

            try:
                shutil.copy2(source, dest)
            except Exception as exc:
                QMessageBox.critical(
                    self,
                    "Errore copia",
                    f"Impossibile copiare il file:\n{filename}\n\n{exc}"
                )

        self.refresh()

    def add_files(self):
        paths, _ = QFileDialog.getOpenFileNames(
            self,
            "Aggiungi allegati",
            "",
            "Tutti i file (*.*)"
        )

        if not paths:
            return

        self.copy_files_to_attachments(paths)

    def open_selected(self):
        path = self.selected_file_path()

        if not path or not os.path.exists(path):
            QMessageBox.warning(
                self,
                "Nessun file",
                "Seleziona un allegato da aprire."
            )
            return

        os.startfile(path)

    def remove_selected(self):
        path = self.selected_file_path()

        if not path or not os.path.exists(path):
            QMessageBox.warning(
                self,
                "Nessun file",
                "Seleziona un allegato da rimuovere."
            )
            return

        filename = os.path.basename(path)

        choice = QMessageBox.question(
            self,
            "Rimuovi allegato",
            f"Vuoi rimuovere questo allegato?\n\n{filename}",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )

        if choice != QMessageBox.Yes:
            return

        try:
            os.remove(path)
        except Exception as exc:
            QMessageBox.critical(
                self,
                "Errore",
                f"Impossibile rimuovere il file:\n{exc}"
            )

        self.refresh()

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()
    
    def dropEvent(self, event):
        if not event.mimeData().hasUrls():
            event.ignore()
            return

        paths = []

        for url in event.mimeData().urls():
            path = url.toLocalFile()
            if path and os.path.isfile(path):
                paths.append(path)

        if not paths:
            event.ignore()
            return

        self.copy_files_to_attachments(paths)
        event.acceptProposedAction()

class GeminiWorker(QObject):
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, api_key, model_name, prompt_text, mode, style, source_text="",
                 provider="gemini", reasoning="base"):
        super().__init__()
        self.api_key = api_key
        self.model_name = model_name
        self.prompt_text = prompt_text
        self.mode = mode
        self.style = style
        self.source_text = source_text
        self.provider = provider
        self.reasoning = reasoning

    def run(self):
        try:
            style_instruction = {
                "Tecnico formale": "Usa un linguaggio tecnico, formale e adatto a una relazione professionale.",
                "Sintetico": "Produci un testo sintetico, chiaro e senza ripetizioni.",
                "Descrittivo": "Produci un testo descrittivo, completo e scorrevole.",
                "Correttivo": "Correggi forma, grammatica e chiarezza, mantenendo il contenuto tecnico invariato.",
                "Libero": "Segui esclusivamente le istruzioni dell'utente.",
            }.get(
                self.style,
                "Usa un linguaggio tecnico, formale e adatto a una relazione professionale."
            )

            if self.mode == "Rielabora testo":
                full_prompt = f"""
                    Sei un assistente tecnico per la redazione di relazioni tecniche.

                    Rielabora il testo sorgente secondo le istruzioni dell'utente.
                    Mantieni invariato il significato tecnico.
                    Non inventare dati numerici, riferimenti normativi o informazioni non presenti.
                    Non aggiungere conclusioni non richieste.
                    Restituisci solo il testo rielaborato, senza premesse.

                    Stile richiesto:
                    {style_instruction}

                    Testo sorgente:
                    {self.source_text}

                    Istruzioni dell'utente:
                    {self.prompt_text}
                    """
            else:
                full_prompt = f"""
                    Sei un assistente tecnico per la redazione di relazioni tecniche.

                    Scrivi un testo professionale, chiaro, coerente e utilizzabile in un documento tecnico.
                    Non inventare dati numerici, riferimenti normativi o informazioni non richieste.
                    Restituisci solo il testo generato, senza premesse.

                    Stile richiesto:
                    {style_instruction}

                    Richiesta dell'utente:
                    {self.prompt_text}
                    """

            if self.provider == "openai":
                result = self._run_openai(full_prompt)
            elif self.provider == "anthropic":
                result = self._run_anthropic(full_prompt)
            else:
                result = self._run_gemini(full_prompt)

            if not result:
                self.error.emit(f"{self._provider_label()} ha restituito una risposta vuota.")
                return

            self.finished.emit(result)

        except Exception as exc:
            self.error.emit(str(exc))

    def _provider_label(self):
        return {
            "openai": "OpenAI",
            "anthropic": "Claude",
        }.get(self.provider, "Gemini")

    def _run_gemini(self, full_prompt):
        from google import genai

        client = genai.Client(api_key=self.api_key)

        if self.reasoning == "avanzato":
            from google.genai import types
            config = types.GenerateContentConfig(
                thinking_config=types.ThinkingConfig(thinking_budget=8192)
            )
            response = client.models.generate_content(
                model=self.model_name,
                contents=full_prompt,
                config=config,
            )
        else:
            response = client.models.generate_content(
                model=self.model_name,
                contents=full_prompt,
            )

        result = ""
        if hasattr(response, "text") and response.text:
            result = response.text.strip()
        return result

    def _run_openai(self, full_prompt):
        from openai import OpenAI

        client = OpenAI(api_key=self.api_key)

        kwargs = {
            "model": self.model_name,
            "input": full_prompt,
        }
        if self.reasoning == "avanzato":
            kwargs["reasoning"] = {"effort": "high"}

        response = client.responses.create(**kwargs)

        text = getattr(response, "output_text", "") or ""
        return text.strip()

    def _run_anthropic(self, full_prompt):
        from anthropic import Anthropic

        client = Anthropic(api_key=self.api_key)

        kwargs = {
            "model": self.model_name,
            "max_tokens": 4096,
            "messages": [{"role": "user", "content": full_prompt}],
        }
        if self.reasoning == "avanzato":
            kwargs["max_tokens"] = 8192
            kwargs["thinking"] = {"type": "enabled", "budget_tokens": 4096}

        response = client.messages.create(**kwargs)

        parts = []
        for block in getattr(response, "content", []) or []:
            if getattr(block, "type", None) == "text":
                parts.append(getattr(block, "text", "") or "")
        return "".join(parts).strip()

class AiResultReceiver(QObject):
    def __init__(self, result_edit, button, parent=None):
        super().__init__(parent)
        self.result_edit = result_edit
        self.button = button

    @Slot(str)
    def show_result(self, text):
        self.result_edit.setPlainText(text)

    @Slot(str)
    def show_error(self, message):
        msg = str(message)

        if "429" in msg or "RESOURCE_EXHAUSTED" in msg or "quota" in msg.lower():
            self.result_edit.setPlainText(
                "Quota esaurita.\n\n"
                "Hai superato il limite disponibile per il modello selezionato.\n"
                "Riprova più tardi oppure usa un modello più leggero."
            )
            return

        if "503" in msg or "UNAVAILABLE" in msg:
            self.result_edit.setPlainText(
                "Servizio temporaneamente non disponibile.\n\n"
                "Il modello è sovraccarico o non raggiungibile in questo momento.\n"
                "Riprova più tardi."
            )
            return

        self.result_edit.setPlainText(f"Errore AI:\n\n{msg}")

    @Slot()
    def generation_finished(self):
        self.button.setEnabled(True)

class LicenseManager:
    def __init__(self, settings):
        self.settings = settings

    def get_license_key(self):
        return self.settings.value("license/key", "")

    def set_license_key(self, license_key):
        self.settings.setValue("license/key", license_key)
        self.settings.sync()

    def get_machine_fingerprint(self):
        return platform.node()

    def get_server_url(self):
        return self.settings.value("license/server_url", "http://51.45.32.64")

    def set_server_url(self, url):
        self.settings.setValue("license/server_url", url)
        self.settings.sync()

    def get_license_status(self):
        license_key = self.get_license_key().strip()
        if not license_key:
            return "NON CONFIGURATA"
        ultimo_esito = self.settings.value("license/ultimo_esito", "")
        if ultimo_esito == "ok":
            return "ATTIVA"
        elif ultimo_esito:
            return f"NON VALIDA ({ultimo_esito})"
        return "DA VERIFICARE"

    def verifica_licenza_server(self):
        """Chiama l'API licenze_verifica.php e salva sempre l'esito nel file .ini."""
        import urllib.request
        import json as _json
        from datetime import datetime

        chiave = self.get_license_key().strip()

        def salva_esito(
            esito,
            messaggio="",
            scadenza=None,
            tipo_licenza="",
            tipo_utente="",
            id_utente=""
        ):

            self.settings.setValue("license/ultimo_esito", esito)

            self.settings.setValue(
                "license/ultima_verifica",
                datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            )

            if scadenza is not None:
                self.settings.setValue("license/scadenza", scadenza)

            self.settings.setValue(
                "license/tipo_licenza",
                tipo_licenza
            )

            self.settings.setValue(
                "license/tipo_utente",
                tipo_utente
            )

            self.settings.setValue(
                "license/id_utente",
                id_utente
            )

            self.settings.sync()

            return {
                "esito": esito,
                "messaggio": messaggio,
                "scadenza": scadenza,
                "tipo_licenza": tipo_licenza,
                "tipo_utente": tipo_utente
            }

        if not chiave:
            return salva_esito(
                "errore",
                "Chiave licenza non configurata.",
                ""
            )

        server_url = self.get_server_url().rstrip("/")
        url = f"{server_url}/api/licenze_verifica.php"

        payload = _json.dumps({
            "chiave": chiave,
            "fingerprint": self.get_machine_fingerprint(),
            "nome_macchina": platform.node(),
            "nome_software": "EDITOR TECNICO"
        }).encode("utf-8")

        try:

            req = urllib.request.Request(
                url,
                data=payload,
                headers={"Content-Type": "application/json"},
                method="POST"
            )

            with urllib.request.urlopen(req, timeout=10) as resp:
                data = _json.loads(
                    resp.read().decode("utf-8")
                )

        except Exception as e:

            return salva_esito(
                "errore",
                f"Impossibile contattare il server: {e}",
                ""
            )

        esito = data.get("esito", "errore")
        messaggio = data.get("messaggio", "")
        scadenza = data.get("scadenza", "")

        return salva_esito(
            esito,
            messaggio,
            scadenza,
            data.get("tipo_licenza", ""),
            data.get("tipo_utente", ""),
            data.get("id_utente", "")
        )

    def get_ultima_verifica(self):
        return self.settings.value("license/ultima_verifica", "Mai eseguita")

    def get_scadenza(self):
        return self.settings.value("license/scadenza", "Non disponibile")

    def is_license_locally_valid(self):
        license_key = self.get_license_key().strip()
        if not license_key:
            return False

        ultimo_esito = self.settings.value("license/ultimo_esito", "")
        if ultimo_esito != "ok":
            return False

        scadenza = self.get_scadenza()

        if not scadenza or scadenza == "Non disponibile":
            return False

        try:
            from datetime import datetime, date

            scadenza_str = str(scadenza).strip()[:10]
            scadenza_date = datetime.strptime(scadenza_str, "%Y-%m-%d").date()

            return scadenza_date >= date.today()

        except Exception:
            return False


    def license_block_reason(self):
        if not self.get_license_key().strip():
            return "Chiave licenza mancante."

        ultimo_esito = self.settings.value("license/ultimo_esito", "")
        if ultimo_esito != "ok":
            return "Ultima verifica licenza non valida."

        scadenza = self.get_scadenza()
        if not scadenza or scadenza == "Non disponibile":
            return "Scadenza licenza non disponibile."

        try:
            from datetime import datetime, date

            scadenza_str = str(scadenza).strip()[:10]
            scadenza_date = datetime.strptime(scadenza_str, "%Y-%m-%d").date()

            if scadenza_date < date.today():
                return "Licenza scaduta."

        except Exception:
            return "Formato scadenza licenza non valido."

        return ""

class MainWindow(QMainWindow):    
    def open_manuali_web(self):
        # I manuali non sono più gestiti dal desktop: vengono pubblicati sul sito.
        # Questa voce apre il sito nel browser predefinito.
        # Per ora punta alla landing page (index.html); aggiornare MANUALI_URL_PATH
        # con il percorso della pagina manuali dedicata quando sarà disponibile.
        MANUALI_URL_PATH = "index.html"  # landing page del sito (provvisorio)

        base_url = self.license_manager.get_server_url().rstrip("/")
        manuali_url = f"{base_url}/{MANUALI_URL_PATH.lstrip('/')}"

        try:
            opened = QDesktopServices.openUrl(QUrl(manuali_url))
            if not opened:
                raise RuntimeError("Apertura del browser non riuscita.")

        except Exception as exc:
            QMessageBox.critical(
                self,
                "Errore manuali",
                f"Impossibile aprire la pagina dei manuali:\n\n{exc}"
            )

    def general_settings_path(self):
        base_dir = QStandardPaths.writableLocation(QStandardPaths.AppDataLocation)

        if not base_dir:
            base_dir = os.path.join(os.path.expanduser("~"), ".EditorTecnico")

        os.makedirs(base_dir, exist_ok=True)

        return os.path.join(base_dir, "settings.ini")

    def __init__(self):
        self.current_project_path = None
        super().__init__()
        self.setWindowTitle("Editor Tecnico (designed by Ing. Luca Monardi)")
        self.resize(1180, 760)
        self.current_item = None
        self.project_dir = None
        self.document_path = None
        self.project_dirty = False
        self.settings = QSettings(self.general_settings_path(), QSettings.IniFormat)
        self.license_manager = LicenseManager(self.settings)
        saved_workspace_dir = self.settings.value("workspace/workspace_dir", "")

        if self.license_manager.is_license_locally_valid():
            self.workspace_dir = saved_workspace_dir
            self.workspace_ready = bool(
                self.workspace_dir and os.path.exists(self.workspace_dir)
            )
        else:
            self.workspace_dir = ""
            self.workspace_ready = False

        self.workspace_dependent_actions = []
        if not self.license_manager.is_license_locally_valid():
            QTimer.singleShot(
                300,
                lambda: QMessageBox.warning(
                    self,
                    "Licenza non valida",
                    "Licenza non valida, mancante o scaduta.\n\n"
                    "Il workspace è stato disabilitato.\n"
                    "Verificare la licenza per continuare a utilizzare il software."
                )
            )
        self.layout_config = {
            "cover": {
                "header": {"type": "none", "value": "", "align": "right"},
                "footer": {"type": "none", "value": "", "align": "center"}
            },
            "content": {
                "header": {"type": "none", "value": "", "align": "right"},
                "footer": {"type": "none", "value": "", "align": "center"}
            }
        }

        self.cover_config = {
            "blocks": []
        }

        self.build_top_toolbar()
        self.build_status_bar()
        self.build_ui()
        self.shortcut_add_text_block = QShortcut(QKeySequence("Ctrl+T"), self)
        self.shortcut_add_text_block.setContext(Qt.ApplicationShortcut)
        self.shortcut_add_text_block.activated.connect(self.add_text_block_from_shortcut)
        self.update_workspace_enabled_state()
        self.populate_empty_structure()
        self.hyphenator = pyphen.Pyphen(lang="it_IT")
        self.autosave_timer = QTimer(self)
        self.autosave_timer.setInterval(30000)  # ogni 30 secondi
        self.autosave_timer.timeout.connect(self.auto_save_project)
        self.autosave_timer.start()

    def add_text_block_from_shortcut(self):
        if self.current_item is None:
            self.statusBar().showMessage("Seleziona un nodo prima di inserire un blocco testo", 2500)
            return

        self.block_editor.add_text_block()
        self.statusBar().showMessage("Blocco testo inserito", 2000)

    def open_layout_dialog(self):
        dlg = LayoutDialog(self.layout_config, self)
        if dlg.exec():
            self.mark_dirty()

    def open_cover_dialog(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Copertina documento")
        dialog.resize(1250, 760)
        dialog.setMinimumWidth(1180)

        layout = QVBoxLayout(dialog)

        class CoverEditorAdapter:
            def __init__(self, main_window, blocks_layout):
                self.main_window = main_window
                self.blocks_layout = blocks_layout
                self.current_blocks_widgets = []

            def prepare_cover_text_block(self, widget):
                if hasattr(widget, "btn_ai"):
                    widget.btn_ai.hide()

                if hasattr(widget, "btn_insert_attachment"):
                    widget.btn_insert_attachment.hide()

                if hasattr(widget, "btn_bullets"):
                    widget.btn_bullets.hide()

                if hasattr(widget, "btn_numbers"):
                    widget.btn_numbers.hide()

                widget.editor.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
                widget.editor.setMaximumHeight(90)
                widget.editor.setMinimumHeight(70)
                widget.content_layout.setSpacing(2)
                widget.toolbar_widget.setContentsMargins(0, 0, 0, 0)
                widget.editor.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
                widget.editor.setLineWrapMode(QTextEdit.WidgetWidth)

                widget.toolbar_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

                return widget


            def prepare_cover_image_block(self, widget):
                if hasattr(widget, "btn_copy_ref"):
                    widget.btn_copy_ref.hide()

                if hasattr(widget, "reference_label"):
                    widget.reference_label.hide()

                widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

                scale_row = QHBoxLayout()

                scale_label = QLabel("Scala immagine (%)")
                scale_spin = QSpinBox()
                scale_spin.setRange(10, 200)
                scale_spin.setSingleStep(5)
                scale_spin.setValue(int(widget.data.get("scale_percent", 100) or 100))
                scale_spin.setFixedWidth(80)

                scale_row.addWidget(scale_label)
                scale_row.addWidget(scale_spin)
                scale_row.addStretch()

                widget.content_layout.addLayout(scale_row)

                def on_cover_image_scale_changed(value):
                    widget.data["scale_percent"] = int(value)
                    self.main_window.mark_dirty()

                scale_spin.valueChanged.connect(on_cover_image_scale_changed)
                widget.adjustSize()
                h = widget.sizeHint().height()
                widget.setMinimumHeight(h)
                widget.setMaximumHeight(h)

                def store_cover_image_from_path(source_path):
                    if not source_path:
                        return

                    if not self.main_window.project_dir:
                        QMessageBox.warning(
                            widget,
                            "Cartella progetto mancante",
                            "Salva prima la relazione, poi inserisci l'immagine di copertina."
                        )
                        return

                    block_id = widget.data.get("id") or new_block_id()
                    widget.data["id"] = block_id

                    images_dir = os.path.join(
                        self.main_window.project_dir,
                        "assets",
                        "images",
                        "cover"
                    )
                    os.makedirs(images_dir, exist_ok=True)

                    ext = os.path.splitext(source_path)[1] or ".png"
                    filename = f"{block_id}{ext}"
                    dest_path = os.path.join(images_dir, filename)

                    shutil.copy2(source_path, dest_path)

                    widget.image_path = f"assets/images/cover/{filename}"
                    widget.data["path"] = widget.image_path

                    widget.update_summary()
                    self.main_window.mark_dirty()

                try:
                    widget.btn_load.clicked.disconnect()
                except (TypeError, RuntimeError):
                    pass

                def load_cover_image():
                    path, _ = QFileDialog.getOpenFileName(
                        widget,
                        "Seleziona immagine",
                        "",
                        "Immagini (*.png *.jpg *.jpeg *.bmp *.gif *.tif *.tiff)"
                    )

                    if not path:
                        return

                    store_cover_image_from_path(path)

                widget.btn_load.clicked.connect(load_cover_image)

                try:
                    widget.btn_paste.clicked.disconnect()
                except (TypeError, RuntimeError):
                    pass

                def paste_cover_image():
                    clipboard = QApplication.clipboard()
                    image = clipboard.image()

                    if image.isNull():
                        QMessageBox.warning(
                            widget,
                            "Incolla immagine",
                            "Negli appunti non è presente un'immagine valida."
                        )
                        return

                    if not self.main_window.project_dir:
                        QMessageBox.warning(
                            widget,
                            "Cartella progetto mancante",
                            "Salva prima la relazione, poi incolla l'immagine di copertina."
                        )
                        return

                    block_id = widget.data.get("id") or new_block_id()
                    widget.data["id"] = block_id

                    images_dir = os.path.join(
                        self.main_window.project_dir,
                        "assets",
                        "images",
                        "cover"
                    )
                    os.makedirs(images_dir, exist_ok=True)

                    filename = f"{block_id}.png"
                    file_path = os.path.join(images_dir, filename)

                    if not image.save(file_path, "PNG"):
                        QMessageBox.critical(
                            widget,
                            "Errore",
                            "Impossibile salvare l'immagine dagli appunti."
                        )
                        return

                    widget.image_path = f"assets/images/cover/{filename}"
                    widget.data["path"] = widget.image_path

                    widget.update_summary()
                    self.main_window.mark_dirty()

                widget.btn_paste.clicked.connect(paste_cover_image)

                try:
                    widget.btn_view.clicked.disconnect()
                except (TypeError, RuntimeError):
                    pass

                def show_cover_image_preview():
                    image_path = widget.image_path or widget.data.get("path", "")

                    if not image_path:
                        QMessageBox.warning(
                            widget,
                            "Immagine mancante",
                            "Nessuna immagine selezionata."
                        )
                        return

                    resolved = self.main_window.resolve_runtime_path(image_path)

                    if not os.path.exists(resolved):
                        QMessageBox.warning(
                            widget,
                            "Immagine non trovata",
                            resolved
                        )
                        return

                    dlg = QDialog(widget)
                    dlg.setWindowTitle("Anteprima immagine")
                    dlg.resize(900, 700)

                    lay = QVBoxLayout(dlg)

                    scroll = QScrollArea()
                    scroll.setWidgetResizable(True)

                    lbl = QLabel()
                    lbl.setAlignment(Qt.AlignCenter)

                    pix = QPixmap(resolved)

                    scale_percent = int(widget.data.get("scale_percent", 100) or 100)

                    base_w = pix.width()
                    base_h = pix.height()

                    scaled_w = max(1, int(base_w * scale_percent / 100))
                    scaled_h = max(1, int(base_h * scale_percent / 100))

                    scaled_pix = pix.scaled(
                        scaled_w,
                        scaled_h,
                        Qt.KeepAspectRatio,
                        Qt.SmoothTransformation
                    )

                    lbl.setPixmap(scaled_pix)

                    scroll.setWidget(lbl)
                    lay.addWidget(scroll)

                    dlg.exec()

                widget.btn_view.clicked.connect(show_cover_image_preview)

                return widget

            def add_text_block(self, data=None):
                w = self.prepare_cover_text_block(
                    TextBlockWidget(self, data)
                )
                self.current_blocks_widgets.append(w)
                self.blocks_layout.insertWidget(
                    self.blocks_layout.count() - 1,
                    w
                )

            def add_image_block(self, data=None):
                w = self.prepare_cover_image_block(
                    ImageBlockWidget(self, data)
                )
                self.current_blocks_widgets.append(w)
                self.blocks_layout.insertWidget(
                    self.blocks_layout.count() - 1,
                    w
                )

            def add_table_block(self, data=None):
                w = TableBlockWidget(self, data)

                w.setMaximumHeight(220)
                w.setMinimumHeight(160)

                if hasattr(w, "table"):
                    w.table.setMinimumHeight(90)
                    w.table.setMaximumHeight(140)

                self.current_blocks_widgets.append(w)
                self.blocks_layout.insertWidget(
                    self.blocks_layout.count() - 1,
                    w
                )

            def remove_block(self, widget):
                if widget in self.current_blocks_widgets:
                    self.current_blocks_widgets.remove(widget)
                    self.blocks_layout.removeWidget(widget)
                    widget.deleteLater()
                    self.main_window.mark_dirty()

            def duplicate_block(self, widget):
                if widget not in self.current_blocks_widgets:
                    return

                index = self.current_blocks_widgets.index(widget)
                data = self.main_window.clone_block_data(widget.export_data())
                block_type = data.get("type")

                if block_type == "text":
                    new_widget = self.prepare_cover_text_block(
                        TextBlockWidget(self, data)
                    )
                elif block_type == "image":
                    new_widget = self.prepare_cover_image_block(
                        ImageBlockWidget(self, data)
                    )
                elif block_type == "table":
                    new_widget = TableBlockWidget(self, data)
                else:
                    return

                self.current_blocks_widgets.insert(index + 1, new_widget)
                self.blocks_layout.insertWidget(index + 1, new_widget)
                self.main_window.mark_dirty()

            def move_block(self, widget, direction):
                if widget not in self.current_blocks_widgets:
                    return

                old_index = self.current_blocks_widgets.index(widget)
                new_index = old_index + direction

                if new_index < 0 or new_index >= len(self.current_blocks_widgets):
                    return

                self.current_blocks_widgets.pop(old_index)
                self.current_blocks_widgets.insert(new_index, widget)

                self.blocks_layout.removeWidget(widget)
                self.blocks_layout.insertWidget(new_index, widget)

                self.main_window.mark_dirty()

            def insert_image_block_after(self, reference_widget, image_path):
                if reference_widget not in self.current_blocks_widgets:
                    self.add_image_block({
                        "type": "image",
                        "path": image_path,
                        "caption": "",
                        "collapsed": False,
                    })
                    return

                index = self.current_blocks_widgets.index(reference_widget)

                data = {
                    "type": "image",
                    "path": image_path,
                    "caption": "",
                    "collapsed": False,
                }

                w = self.prepare_cover_image_block(
                    ImageBlockWidget(self, data)
                )
                self.current_blocks_widgets.insert(index + 1, w)
                self.blocks_layout.insertWidget(index + 1, w)
                self.main_window.mark_dirty()

            def export_blocks(self):
                exported = []

                for widget in self.current_blocks_widgets:
                    data = widget.export_data()

                    if data.get("type") == "image":
                        data["path"] = as_posix_path(
                            getattr(widget, "image_path", "") or data.get("path", "")
                        )
                        data["scale_percent"] = int(widget.data.get("scale_percent", 100) or 100)

                    exported.append(data)

                return exported

        toolbar_row = QHBoxLayout()

        btn_add_text = QPushButton("+ Testo")
        btn_add_image = QPushButton("+ Immagine")
        btn_add_table = QPushButton("+ Tabella")
        btn_preview_cover = QPushButton("Anteprima copertina")

        for btn in (
            btn_add_text,
            btn_add_image,
            btn_add_table,
            btn_preview_cover,
        ):
            style_small_button(btn)

        toolbar_row.addWidget(btn_add_text)
        toolbar_row.addWidget(btn_add_image)
        toolbar_row.addWidget(btn_add_table)
        toolbar_row.addWidget(btn_preview_cover)
        toolbar_row.addStretch()

        layout.addLayout(toolbar_row)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)

        cover_container = QWidget()
        cover_layout = QVBoxLayout(cover_container)
        cover_layout.setContentsMargins(0, 0, 0, 0)
        cover_layout.setSpacing(10)
        cover_layout.setAlignment(Qt.AlignTop)
        cover_layout.addStretch()

        scroll.setWidget(cover_container)
        layout.addWidget(scroll)

        cover_editor = CoverEditorAdapter(self, cover_layout)

        for block in self.cover_config.get("blocks", []):
            block_type = block.get("type")

            if block_type == "text":
                cover_editor.add_text_block(block)
            elif block_type == "image":
                cover_editor.add_image_block(block)
            elif block_type == "table":
                cover_editor.add_table_block(block)

        btn_add_text.clicked.connect(lambda: cover_editor.add_text_block())
        btn_add_image.clicked.connect(lambda: cover_editor.add_image_block())
        btn_add_table.clicked.connect(lambda: cover_editor.add_table_block())

        def show_cover_preview():
            blocks = cover_editor.export_blocks()

            dlg = QDialog(dialog)
            dlg.setWindowTitle("Anteprima copertina")
            dlg.resize(900, 1000)

            root = QVBoxLayout(dlg)

            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
            scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)

            container = QWidget()
            container_layout = QVBoxLayout(container)
            container_layout.setAlignment(Qt.AlignTop)

            page = QFrame()
            page.setFrameShape(QFrame.Box)
            page.setStyleSheet("QFrame { background: white; border: 1px solid #999; }")
            page.setMinimumWidth(720)

            page_layout = QVBoxLayout(page)
            page_layout.setContentsMargins(60, 60, 60, 60)
            page_layout.setSpacing(18)

            for block in blocks:
                btype = block.get("type")

                if btype == "text":
                    text_preview = QTextEdit()
                    text_preview.setReadOnly(True)
                    text_preview.setFrameShape(QFrame.NoFrame)
                    text_preview.setStyleSheet("QTextEdit { background: white; color: black; border: none; }")

                    html = block.get("html", "")
                    if html:
                        text_preview.setHtml(html)
                    else:
                        text_preview.setPlainText(block.get("text", ""))

                    text_preview.setMinimumHeight(80)
                    text_preview.setMaximumHeight(160)
                    page_layout.addWidget(text_preview)

                elif btype == "image":
                    image_path = self.resolve_runtime_path(block.get("path", ""))

                    if image_path and os.path.exists(image_path):
                        pix = QPixmap(image_path)

                        if not pix.isNull():
                            scale_percent = int(block.get("scale_percent", 100) or 100)

                            scaled_w = max(1, int(pix.width() * scale_percent / 100))
                            scaled_h = max(1, int(pix.height() * scale_percent / 100))

                            scaled_pix = pix.scaled(
                                scaled_w,
                                scaled_h,
                                Qt.KeepAspectRatio,
                                Qt.SmoothTransformation
                            )

                            img_lbl = QLabel()
                            img_lbl.setAlignment(Qt.AlignCenter)
                            img_lbl.setPixmap(scaled_pix)

                            page_layout.addWidget(img_lbl)

                elif btype == "table":
                    rows = int(block.get("rows", 0) or 0)
                    cols = int(block.get("cols", 0) or 0)
                    data = block.get("data", [])

                    if rows > 0 and cols > 0:
                        table = QTableWidget(rows, cols)
                        table.setEditTriggers(QTableWidget.NoEditTriggers)
                        table.horizontalHeader().setVisible(False)
                        table.verticalHeader().setVisible(False)
                        table.setStyleSheet(
                            "QTableWidget { background: white; color: black; gridline-color: #999; }"
                            "QTableWidget::item { background: white; color: black; }"
                        )

                        for r in range(rows):
                            for c in range(cols):
                                value = ""
                                if r < len(data) and c < len(data[r]):
                                    value = data[r][c]
                                table.setItem(r, c, QTableWidgetItem(value))

                        table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
                        table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)

                        table.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
                        table.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)

                        table.resizeRowsToContents()

                        table.setMinimumWidth(600)
                        table.setMaximumWidth(600)
                        table.setMaximumHeight(220)

                        table_row = QHBoxLayout()
                        table_row.addStretch()
                        table_row.addWidget(table)
                        table_row.addStretch()

                        page_layout.addLayout(table_row)

            container_layout.addWidget(page)
            scroll.setWidget(container)
            root.addWidget(scroll)

            dlg.exec()

        btn_preview_cover.clicked.connect(show_cover_preview)

        btn_row = QHBoxLayout()
        btn_save = QPushButton("Salva")
        btn_cancel = QPushButton("Annulla")

        btn_row.addStretch()
        btn_row.addWidget(btn_save)
        btn_row.addWidget(btn_cancel)

        layout.addLayout(btn_row)

        def save_cover():
            self.cover_config["blocks"] = cover_editor.export_blocks()
            self.mark_dirty()
            dialog.accept()

        btn_save.clicked.connect(save_cover)
        btn_cancel.clicked.connect(dialog.reject)

        dialog.exec()

    def show_node_preview(self):
        # Anteprima a schermo (sola lettura) del nodo selezionato: lavora su una
        # copia dei blocchi e non scrive nulla nel progetto.
        self.save_current_item_content()
        self.update_figure_references()
        self.update_equation_references()

        if self.current_item is None:
            QMessageBox.warning(
                self,
                "Nessun nodo selezionato",
                "Seleziona prima un capitolo o paragrafo nella struttura."
            )
            return

        blocks = self.ensure_blocks_defaults_list(
            deepcopy(self.current_item.data(0, CONTENT_ROLE) or [])
        )

        dialog = QDialog(self)
        dialog.setWindowTitle("Anteprima pagina")
        available = self.screen().availableGeometry()
        dialog.resize(min(900, available.width()), min(1000, available.height()))

        root = QVBoxLayout(dialog)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)

        container = QWidget()
        container_layout = QVBoxLayout(container)
        container_layout.setAlignment(Qt.AlignTop)

        page = QFrame()
        page.setFrameShape(QFrame.Box)
        page.setStyleSheet("QFrame { background: white; border: 1px solid #999; }")
        page.setMinimumWidth(720)

        page_layout = QVBoxLayout(page)
        page_layout.setContentsMargins(60, 60, 60, 60)
        page_layout.setSpacing(18)

        def add_image_to_page(image_path, scale_percent, caption_text):
            resolved = self.resolve_runtime_path(image_path or "")

            if resolved and os.path.exists(resolved):
                pix = QPixmap(resolved)

                if not pix.isNull():
                    scale_percent = int(scale_percent or 100)

                    scaled_w = max(1, int(pix.width() * scale_percent / 100))
                    scaled_h = max(1, int(pix.height() * scale_percent / 100))

                    scaled_pix = pix.scaled(
                        scaled_w,
                        scaled_h,
                        Qt.KeepAspectRatio,
                        Qt.SmoothTransformation
                    )

                    img_lbl = QLabel()
                    img_lbl.setAlignment(Qt.AlignCenter)
                    img_lbl.setPixmap(scaled_pix)

                    page_layout.addWidget(img_lbl)

            if caption_text:
                caption_lbl = QLabel(caption_text)
                caption_lbl.setAlignment(Qt.AlignCenter)
                caption_lbl.setWordWrap(True)
                caption_lbl.setStyleSheet("QLabel { font-style: italic; font-size: 9pt; color: black; }")
                page_layout.addWidget(caption_lbl)

        def build_image_caption(block_id, caption):
            info = self.figure_reference_map.get(block_id)
            caption = (caption or "").strip()

            if info:
                label = info.get("label", "")
                if caption:
                    return f"{label} — {caption}"
                return label

            return caption

        for block in blocks:
            btype = block.get("type")

            if btype == "text":
                text_preview = QTextEdit()
                text_preview.setReadOnly(True)
                text_preview.setFrameShape(QFrame.NoFrame)
                text_preview.setStyleSheet("QTextEdit { background: white; color: black; border: none; }")
                text_preview.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
                text_preview.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
                text_preview.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

                html = block.get("html", "")
                if html:
                    rendered_html = self.render_text_references(html)
                    text_preview.setHtml(rendered_html)
                else:
                    text_preview.setPlainText(self.render_text_references(block.get("text", "")))

                # Adatta l'altezza al contenuto reale, così i blocchi restano compatti.
                def _fit_text_height(_=None, te=text_preview):
                    doc = te.document()
                    te.setFixedHeight(int(doc.size().height() + 2 * doc.documentMargin()))

                text_preview.document().documentLayout().documentSizeChanged.connect(_fit_text_height)
                _fit_text_height()
                page_layout.addWidget(text_preview)

            elif btype == "image":
                caption_text = build_image_caption(block.get("id"), block.get("caption", ""))
                add_image_to_page(block.get("path", ""), block.get("scale_percent", 100), caption_text)

            elif btype == "images":
                title = (block.get("title", "") or "").strip()
                if title:
                    title_lbl = QLabel(title)
                    title_lbl.setStyleSheet("QLabel { font-weight: bold; color: black; }")
                    page_layout.addWidget(title_lbl)

                for img in self.ensure_group_images_defaults(block.get("images", [])):
                    caption_text = build_image_caption(img.get("id"), img.get("caption", ""))
                    add_image_to_page(img.get("path", ""), img.get("scale_percent", 100), caption_text)

            elif btype == "equation":
                latex = block.get("latex", "")
                caption = (block.get("caption", "") or "").strip()
                numbering_mode = block.get("numbering_mode", "none")

                eq_info = self.equation_reference_map.get(block.get("id"))
                eq_label = eq_info["label"] if eq_info else ""

                path = self.render_equation_to_png(latex, block.get("id"))

                if path and os.path.exists(path):
                    eq_lbl = QLabel()
                    eq_lbl.setAlignment(Qt.AlignCenter)
                    eq_lbl.setPixmap(QPixmap(path))

                    if numbering_mode == "number" and eq_label:
                        number_lbl = QLabel(eq_label)
                        number_lbl.setStyleSheet("QLabel { color: black; }")

                        eq_row = QHBoxLayout()
                        eq_row.addWidget(eq_lbl, 1)
                        eq_row.addWidget(number_lbl, 0, Qt.AlignRight)

                        page_layout.addLayout(eq_row)
                    else:
                        page_layout.addWidget(eq_lbl)

                        if numbering_mode == "number_caption" and eq_label and caption:
                            caption_lbl = QLabel(f"{eq_label} – {caption}")
                            caption_lbl.setAlignment(Qt.AlignCenter)
                            caption_lbl.setWordWrap(True)
                            caption_lbl.setStyleSheet("QLabel { font-style: italic; color: black; }")
                            page_layout.addWidget(caption_lbl)
                else:
                    error_lbl = QLabel(f"[Equazione non renderizzabile: {block.get('latex', '')}]")
                    error_lbl.setStyleSheet("QLabel { color: black; }")
                    page_layout.addWidget(error_lbl)

            elif btype == "table":
                rows = int(block.get("rows", 0) or 0)
                cols = int(block.get("cols", 0) or 0)
                data = block.get("data", [])

                if rows > 0 and cols > 0:
                    table = QTableWidget(rows, cols)
                    table.setEditTriggers(QTableWidget.NoEditTriggers)
                    table.horizontalHeader().setVisible(False)
                    table.verticalHeader().setVisible(False)
                    table.setStyleSheet(
                        "QTableWidget { background: white; color: black; gridline-color: #999; }"
                        "QTableWidget::item { background: white; color: black; }"
                    )

                    cell_formats = block.get("cell_formats", {})

                    for r in range(rows):
                        for c in range(cols):
                            value = ""
                            if r < len(data) and c < len(data[r]):
                                value = data[r][c]

                            item = QTableWidgetItem(value)

                            fmt = cell_formats.get(f"{r},{c}", {})
                            halign = fmt.get("halign", "left")
                            valign = fmt.get("valign", "middle")

                            if halign in ("center", "justify"):
                                hflag = Qt.AlignHCenter
                            elif halign == "right":
                                hflag = Qt.AlignRight
                            else:
                                hflag = Qt.AlignLeft

                            if valign == "top":
                                vflag = Qt.AlignTop
                            elif valign == "bottom":
                                vflag = Qt.AlignBottom
                            else:
                                vflag = Qt.AlignVCenter

                            item.setTextAlignment(hflag | vflag)
                            table.setItem(r, c, item)

                    for span in block.get("spans", []):
                        r1 = span.get("r1", 0)
                        c1 = span.get("c1", 0)
                        r2 = span.get("r2", r1)
                        c2 = span.get("c2", c1)

                        if (r2 - r1 + 1) > 1 or (c2 - c1 + 1) > 1:
                            table.setSpan(r1, c1, r2 - r1 + 1, c2 - c1 + 1)

                    table.resizeRowsToContents()
                    page_layout.addWidget(table)

        container_layout.addWidget(page)
        scroll.setWidget(container)
        root.addWidget(scroll)

        dialog.exec()

    def get_gemini_api_key(self):
        return self.settings.value("ai/gemini_api_key", "")

    def get_gemini_model(self):
        return self.settings.value("ai/gemini_model", "gemini-2.5-flash")
    
    def set_gemini_settings(self, api_key, model):
        self.settings.setValue("ai/gemini_api_key", api_key)
        self.settings.setValue("ai/gemini_model", model)
        self.settings.sync()

    # --- Multi-provider AI (Gemini / OpenAI / Claude) ---
    AI_PROVIDERS = ("gemini", "openai", "anthropic")
    AI_PROVIDER_LABELS = {
        "gemini": "Google Gemini",
        "openai": "OpenAI",
        "anthropic": "Anthropic Claude",
    }
    AI_MODEL_SUGGESTIONS = {
        "gemini": ["gemini-2.5-flash", "gemini-2.5-pro"],
        "openai": ["gpt-5", "gpt-5-mini", "o4-mini"],
        "anthropic": ["claude-sonnet-4-5", "claude-opus-4-1", "claude-haiku-4-5"],
    }
    AI_DEFAULT_MODEL = {
        "gemini": "gemini-2.5-flash",
        "openai": "gpt-5-mini",
        "anthropic": "claude-sonnet-4-5",
    }

    def get_active_provider(self):
        provider = self.settings.value("ai/active_provider", "gemini")
        return provider if provider in self.AI_PROVIDERS else "gemini"

    def set_active_provider(self, provider):
        if provider in self.AI_PROVIDERS:
            self.settings.setValue("ai/active_provider", provider)
            self.settings.sync()

    def get_provider_api_key(self, provider):
        # Gemini riusa la chiave storica gia' salvata
        if provider == "gemini":
            return self.settings.value("ai/gemini_api_key", "")
        return self.settings.value(f"ai/{provider}_api_key", "")

    def get_provider_model(self, provider):
        if provider == "gemini":
            return self.settings.value("ai/gemini_model", self.AI_DEFAULT_MODEL["gemini"])
        return self.settings.value(
            f"ai/{provider}_model",
            self.AI_DEFAULT_MODEL.get(provider, ""),
        )

    def get_provider_reasoning(self, provider):
        value = self.settings.value(f"ai/{provider}_reasoning", "base")
        return value if value in ("base", "avanzato") else "base"

    def set_provider_settings(self, provider, api_key, model, reasoning):
        if provider == "gemini":
            self.settings.setValue("ai/gemini_api_key", api_key)
            self.settings.setValue("ai/gemini_model", model)
        else:
            self.settings.setValue(f"ai/{provider}_api_key", api_key)
            self.settings.setValue(f"ai/{provider}_model", model)
        self.settings.setValue(f"ai/{provider}_reasoning", reasoning)
        self.settings.sync()

    def get_active_ai_config(self):
        provider = self.get_active_provider()
        return (
            provider,
            self.get_provider_api_key(provider),
            self.get_provider_model(provider),
            self.get_provider_reasoning(provider),
        )

    def open_api_connections_dialog(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Connessioni API")
        dialog.resize(440, 260)

        layout = QVBoxLayout(dialog)

        form = QFormLayout()

        provider_combo = QComboBox()
        for key in self.AI_PROVIDERS:
            provider_combo.addItem(self.AI_PROVIDER_LABELS[key], key)

        api_key_edit = QLineEdit()
        api_key_edit.setEchoMode(QLineEdit.Password)

        model_combo = QComboBox()
        model_combo.setEditable(True)

        reasoning_combo = QComboBox()
        reasoning_combo.addItem("Base", "base")
        reasoning_combo.addItem("Avanzato", "avanzato")

        form.addRow("Provider:", provider_combo)
        form.addRow("API key:", api_key_edit)
        form.addRow("Modello:", model_combo)
        form.addRow("Ragionamento:", reasoning_combo)

        layout.addLayout(form)
        layout.addStretch()

        def load_provider_fields(provider):
            model_combo.blockSignals(True)
            model_combo.clear()
            model_combo.addItems(self.AI_MODEL_SUGGESTIONS.get(provider, []))
            model_combo.setCurrentText(self.get_provider_model(provider))
            model_combo.blockSignals(False)

            api_key_edit.setText(self.get_provider_api_key(provider))

            reasoning = self.get_provider_reasoning(provider)
            idx = reasoning_combo.findData(reasoning)
            reasoning_combo.setCurrentIndex(idx if idx >= 0 else 0)

        def on_provider_changed():
            load_provider_fields(provider_combo.currentData())

        provider_combo.currentIndexChanged.connect(on_provider_changed)

        # inizializza sul provider attivo
        active = self.get_active_provider()
        active_idx = provider_combo.findData(active)
        provider_combo.setCurrentIndex(active_idx if active_idx >= 0 else 0)
        load_provider_fields(provider_combo.currentData())

        buttons_layout = QHBoxLayout()
        buttons_layout.addStretch()

        save_btn = QPushButton("Salva")
        close_btn = QPushButton("Chiudi")

        buttons_layout.addWidget(save_btn)
        buttons_layout.addWidget(close_btn)

        layout.addLayout(buttons_layout)

        def save_settings():
            provider = provider_combo.currentData()
            api_key = api_key_edit.text().strip()
            model = model_combo.currentText().strip()
            reasoning = reasoning_combo.currentData()

            if not api_key:
                QMessageBox.warning(
                    dialog,
                    "API key",
                    f"Inserisci una API key per {self.AI_PROVIDER_LABELS[provider]}."
                )
                return

            self.set_provider_settings(provider, api_key, model, reasoning)
            self.set_active_provider(provider)

            QMessageBox.information(
                dialog,
                "Connessioni API",
                f"Impostazioni salvate. Provider attivo: {self.AI_PROVIDER_LABELS[provider]}."
            )

        save_btn.clicked.connect(save_settings)
        close_btn.clicked.connect(dialog.accept)

        dialog.exec()

    def open_license_dialog(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Licenza")
        dialog.resize(460, 280)

        layout = QVBoxLayout(dialog)

        form = QFormLayout()

        license_edit = QLineEdit()
        license_edit.setPlaceholderText("Es. XXXX-XXXX-XXXX-XXXX")
        license_edit.setText(self.license_manager.get_license_key())

        fingerprint_edit = QLineEdit()
        fingerprint_edit.setReadOnly(True)
        fingerprint_edit.setText(self.license_manager.get_machine_fingerprint())

        status_label = QLabel(
            self.license_manager.get_license_status()
        )

        last_check_label = QLabel(
            self.license_manager.get_ultima_verifica()
        )

        expiration_label = QLabel(
            self.license_manager.get_scadenza()
        )

        license_type_label = QLabel("N.D.")
        user_type_label = QLabel("N.D.")

        license_type_label = QLabel(
            self.settings.value("license/tipo_licenza", "N.D.")
        )

        user_type_label = QLabel(
            self.settings.value("license/tipo_utente", "N.D.")
        )

        user_id_label = QLabel(
            str(self.settings.value("license/id_utente", "N.D."))
        )

        form.addRow("Chiave licenza:", license_edit)
        form.addRow("Fingerprint:", fingerprint_edit)
        form.addRow("Stato:", status_label)
        form.addRow("Ultima verifica:", last_check_label)
        form.addRow("Scadenza:", expiration_label)
        form.addRow("Tipo licenza:", license_type_label)
        form.addRow("Tipo utente:", user_type_label)
        form.addRow("ID utente:", user_id_label)

        layout.addLayout(form)
        layout.addStretch()

        buttons_layout = QHBoxLayout()
        buttons_layout.addStretch()

        btn_verify  = QPushButton("Verifica licenza")
        btn_close   = QPushButton("Chiudi")

        def verify_license():
            license_key = license_edit.text().strip()
            self.license_manager.set_license_key(license_key)

            if not license_key:

                self.settings.setValue("license/ultimo_esito", "errore")
                self.settings.setValue("license/scadenza", "")
                self.settings.setValue("license/tipo_licenza", "")
                self.settings.setValue("license/tipo_utente", "")

                self.settings.sync()

                self.workspace_dir = ""
                self.workspace_ready = False

                self.update_workspace_enabled_state()

                status_label.setText(
                    self.license_manager.get_license_status()
                )

                last_check_label.setText(
                    self.license_manager.get_ultima_verifica()
                )

                expiration_label.setText("Non disponibile")

                license_type_label.setText("N.D.")
                user_type_label.setText("N.D.")

                QMessageBox.warning(
                    dialog,
                    "Licenza",
                    "Chiave licenza non configurata."
                )

                return

            btn_verify.setText("Verifica in corso...")
            btn_verify.setEnabled(False)

            QApplication.processEvents()

            risultato = self.license_manager.verifica_licenza_server()

            btn_verify.setText("Verifica licenza")
            btn_verify.setEnabled(True)

            esito = risultato.get("esito", "errore")

            messaggio = risultato.get(
                "messaggio",
                "Errore sconosciuto."
            )

            status_label.setText(
                self.license_manager.get_license_status()
            )

            last_check_label.setText(
                self.license_manager.get_ultima_verifica()
            )

            license_type_label.setText(
                self.settings.value(
                    "license/tipo_licenza",
                    "N.D."
                )
            )

            user_type_label.setText(
                self.settings.value(
                    "license/tipo_utente",
                    "N.D."
                )
            )

            if esito == "ok":

                saved_workspace = self.settings.value(
                    "workspace/workspace_dir",
                    ""
                )

                if saved_workspace and os.path.exists(saved_workspace):
                    self.workspace_dir = saved_workspace

                self.update_workspace_enabled_state()

                scadenza = risultato.get("scadenza", "")

                if scadenza == "permanente":
                    expiration_label.setText("Permanente")
                else:
                    expiration_label.setText(
                        scadenza[:10] if scadenza else "N.D."
                    )

                avviso = risultato.get(
                    "giorni_scadenza",
                    -1
                )

                msg = f"✅ {messaggio}"

                if avviso >= 0 and avviso <= 30:
                    msg += (
                        f"\n\nAttenzione: "
                        f"la licenza scade tra {avviso} giorni."
                    )

                QMessageBox.information(
                    dialog,
                    "Verifica licenza",
                    msg
                )

            else:

                self.workspace_dir = ""
                self.workspace_ready = False

                self.update_workspace_enabled_state()

                expiration_label.setText("Non disponibile")

                QMessageBox.warning(
                    dialog,
                    "Verifica licenza",
                    f"❌ {messaggio}"
                )

        btn_verify.clicked.connect(verify_license)
        btn_close.clicked.connect(dialog.accept)

        buttons_layout.addWidget(btn_verify)
        buttons_layout.addWidget(btn_close)
        layout.addLayout(buttons_layout)

        dialog.exec()

    def open_ai_dialog(self):
        dialog = QDialog(self)
        dialog.setWindowTitle(
            f"Assistente AI — {self.AI_PROVIDER_LABELS.get(self.get_active_provider(), 'Gemini')}"
        )
        dialog.resize(850, 700)

        layout = QVBoxLayout(dialog)

        # =========================
        # MODALITÀ
        # =========================
        mode_row = QHBoxLayout()

        mode_row.addWidget(QLabel("Modalità"))

        mode_combo = QComboBox()
        mode_combo.addItems([
            "Genera testo",
            "Rielabora testo",
        ])
        mode_row.addWidget(mode_combo)

        mode_row.addWidget(QLabel("Stile"))

        style_combo = QComboBox()
        style_combo.addItems([
            "Tecnico formale",
            "Sintetico",
            "Descrittivo",
            "Correttivo",
            "Libero",
        ])
        mode_row.addWidget(style_combo)

        mode_row.addStretch()
        layout.addLayout(mode_row)

        # =========================
        # TESTO SORGENTE
        # =========================
        source_label = QLabel("Testo sorgente")
        layout.addWidget(source_label)

        source_edit = QTextEdit()
        source_edit.setPlaceholderText(
            "Incolla qui il testo della relazione da rielaborare..."
        )
        source_edit.setMinimumHeight(130)
        layout.addWidget(source_edit)

        # =========================
        # PROMPT / ISTRUZIONI
        # =========================
        prompt_label = QLabel("Richiesta")
        layout.addWidget(prompt_label)

        prompt_edit = QTextEdit()
        prompt_edit.setPlaceholderText(
            "Descrivi il testo da generare oppure indica come rielaborare il testo sorgente..."
        )
        prompt_edit.setMinimumHeight(120)
        layout.addWidget(prompt_edit)

        # =========================
        # RISULTATO
        # =========================
        result_label = QLabel("Risultato")
        layout.addWidget(result_label)

        result_edit = QTextEdit()
        result_edit.setPlaceholderText("Il testo generato comparirà qui...")
        result_edit.setMinimumHeight(180)
        layout.addWidget(result_edit)

        # =========================
        # PULSANTI
        # =========================
        btn_row = QHBoxLayout()

        btn_generate = QPushButton("Genera")
        btn_copy_result = QPushButton("Copia risultato")
        btn_clear = QPushButton("Svuota")
        btn_close = QPushButton("Chiudi")

        btn_row.addWidget(btn_generate)
        btn_row.addWidget(btn_copy_result)
        btn_row.addWidget(btn_clear)
        btn_row.addStretch()
        btn_row.addWidget(btn_close)

        layout.addLayout(btn_row)

        # =========================
        # COMPORTAMENTO UI BASE
        # =========================
        def update_mode_ui():
            is_rework = mode_combo.currentText() == "Rielabora testo"

            source_label.setVisible(is_rework)
            source_edit.setVisible(is_rework)

            if is_rework:
                prompt_label.setText("Istruzioni")
                prompt_edit.setPlaceholderText(
                    "Es. rendilo più tecnico, sintetico, formale, senza alterare il contenuto..."
                )
            else:
                prompt_label.setText("Richiesta")
                prompt_edit.setPlaceholderText(
                    "Es. Redigi un paragrafo tecnico sullo stato di conservazione delle strutture..."
                )

        def copy_result_to_clipboard():
            text = result_edit.toPlainText().strip()
            if not text:
                QMessageBox.warning(dialog, "AI", "Non c'è nessun risultato da copiare.")
                return

            QApplication.clipboard().setText(text)
            self.statusBar().showMessage("Risultato AI copiato", 2000)

        def clear_fields():
            source_edit.clear()
            prompt_edit.clear()
            result_edit.clear()

        receiver = AiResultReceiver(result_edit, btn_generate, dialog)

        def generate_with_gemini():
            ai_provider, api_key, model, ai_reasoning = self.get_active_ai_config()
            mode = mode_combo.currentText().strip()
            prompt = prompt_edit.toPlainText().strip()

            if not api_key:
                QMessageBox.warning(
                    dialog,
                    "API key mancante",
                    "Inserisci e salva una API key del provider AI prima di generare."
                )
                return

            if mode == "Rielabora testo":
                source_text = source_edit.toPlainText().strip()

                if not source_text:
                    QMessageBox.warning(
                        dialog,
                        "Testo mancante",
                        "Inserisci il testo sorgente da rielaborare."
                    )
                    return

            if not prompt:
                QMessageBox.warning(
                    dialog,
                    "Richiesta mancante",
                    "Scrivi prima una richiesta."
                )
                return

            result_edit.setPlainText("Generazione in corso...")
            btn_generate.setEnabled(False)

            self.ai_thread = QThread()
            self.ai_worker = GeminiWorker(
                api_key=api_key,
                model_name=model,
                prompt_text=prompt,
                mode=mode,
                style=style_combo.currentText().strip(),
                source_text=source_edit.toPlainText().strip(),
                provider=ai_provider,
                reasoning=ai_reasoning,
            )
            self.ai_worker.moveToThread(self.ai_thread)

            self.ai_thread.started.connect(self.ai_worker.run)

            self.ai_worker.finished.connect(receiver.show_result)
            self.ai_worker.finished.connect(self.ai_thread.quit)
            self.ai_worker.finished.connect(self.ai_worker.deleteLater)

            self.ai_worker.error.connect(receiver.show_error)
            self.ai_worker.error.connect(self.ai_thread.quit)
            self.ai_worker.error.connect(self.ai_worker.deleteLater)

            self.ai_thread.finished.connect(receiver.generation_finished)
            self.ai_thread.finished.connect(self.ai_thread.deleteLater)

            self.ai_thread.start()

        mode_combo.currentTextChanged.connect(update_mode_ui)
        btn_copy_result.clicked.connect(copy_result_to_clipboard)
        btn_clear.clicked.connect(clear_fields)
        btn_close.clicked.connect(dialog.reject)

        btn_generate.clicked.connect(generate_with_gemini)

        update_mode_ui()

        dialog.exec()

    def ensure_system_archive_structure(self):
        if not self.workspace_dir:
            return

        system_archive_dir = os.path.join(
            self.workspace_dir,
            "ARCHIVIO DI SISTEMA"
        )

        template_dir = os.path.join(
            system_archive_dir,
            "TEMPLATE"
        )

        technical_relations_template_dir = os.path.join(
            template_dir,
            "RELAZIONI TECNICHE"
        )

        work_archive_dir = os.path.join(
            self.workspace_dir,
            "ARCHIVIO LAVORI"
        )

        for folder in (
            system_archive_dir,
            template_dir,
            technical_relations_template_dir,
            work_archive_dir,
        ):
            os.makedirs(folder, exist_ok=True)

    def select_workspace(self):
        if not self.license_manager.is_license_locally_valid():
            QMessageBox.warning(
                self,
                "Licenza non valida",
                self.license_manager.license_block_reason()
            )
            return
        
        folder = QFileDialog.getExistingDirectory(
            self,
            "Seleziona workspace",
            self.workspace_dir or ""
        )

        if not folder:
            return

        self.workspace_dir = folder
        self.settings.setValue("workspace/workspace_dir", folder)
        self.settings.sync()

        self.ensure_system_archive_structure()

        self.update_workspace_enabled_state()

        self.statusBar().showMessage(
            f"Workspace attivo: {self.workspace_dir}",
            4000
        )

    def open_allegati_folder(self):
        allegati_path = self.current_allegati_dir()

        if not allegati_path:
            QMessageBox.warning(
                self,
                "Nessuna relazione",
                "Apri o crea prima una relazione."
            )
            return

        os.makedirs(allegati_path, exist_ok=True)

        try:
            os.startfile(allegati_path)
        except Exception:
            QMessageBox.critical(
                self,
                "Errore",
                "Impossibile aprire la cartella Allegati."
            )

    def manage_attachments(self):
        allegati_path = self.current_allegati_dir()

        if not allegati_path:
            QMessageBox.warning(
                self,
                "Nessuna relazione",
                "Apri o crea prima una relazione."
            )
            return

        dialog = AttachmentsDialog(allegati_path, self)
        dialog.exec()

    def insert_attachment_reference_in_text(self, text_editor):
        allegati_path = self.current_allegati_dir()

        if not allegati_path or not os.path.exists(allegati_path):
            QMessageBox.warning(
                self,
                "Nessun allegato",
                "Non ci sono allegati disponibili."
            )
            return

        files = [
            name for name in os.listdir(allegati_path)
            if os.path.isfile(os.path.join(allegati_path, name))
        ]

        if not files:
            QMessageBox.information(
                self,
                "Nessun allegato",
                "La cartella Allegati è vuota."
            )
            return

        file_name, ok = QInputDialog.getItem(
            self,
            "Inserisci allegato",
            "Seleziona allegato:",
            sorted(files),
            0,
            False
        )

        if not ok or not file_name:
            return

        cursor = text_editor.textCursor()
        cursor.insertText(f"[[ALL:{file_name}]]")        

    def current_allegati_dir(self):
        if not self.project_dir:
            return None

        allegati_dir = os.path.join(self.project_dir, "Allegati")
        os.makedirs(allegati_dir, exist_ok=True)

        return allegati_dir

    def update_workspace_enabled_state(self):

        license_ok = self.license_manager.is_license_locally_valid()

        self.workspace_ready = bool(
            license_ok and
            self.workspace_dir and
            os.path.exists(self.workspace_dir)
        )

        for action in self.workspace_dependent_actions:
            action.setEnabled(self.workspace_ready)

        if hasattr(self, "workspace_label"):

            self.workspace_label.setEnabled(license_ok)

            if self.workspace_ready:
                self.workspace_label.setText(
                    f"Workspace: {self.workspace_dir}  (doppio click per cambiare)"
                )
            else:

                if not license_ok:
                    self.workspace_label.setText(
                        "Workspace disabilitato — licenza non valida"
                    )
                else:
                    self.workspace_label.setText(
                        "Workspace non selezionato — doppio click per scegliere"
                    )

        license_key = (
            self.license_manager.get_license_key().strip()
        )

        if not license_key:
            license_text = "Licenza: non configurata"
        else:
            license_text = f"Licenza: {license_key}"

        if self.workspace_ready:
            self.statusBar().showMessage(
                f"Workspace attivo: {self.workspace_dir}   |   {license_text}",
                4000
            )
        else:

            if not license_ok:
                self.statusBar().showMessage(
                    f"Licenza non valida — workspace disabilitato   |   {license_text}",
                    4000
                )
            else:
                self.statusBar().showMessage(
                    f"Seleziona un workspace per iniziare   |   {license_text}",
                    4000
                )
        if hasattr(self, "act_beta_note"):
            tipo_licenza = self.settings.value("license/tipo_licenza", "")
            self.act_beta_note.setEnabled(tipo_licenza == "beta")
            
    def update_project_info_labels(self):
        if not self.project_dir:
            self.client_label.setText("Cliente: -")
            self.project_label.setText("Commessa: -")
            self.relation_label.setText("Relazione: -")
            return

        path = os.path.normpath(self.project_dir)
        parts = path.split(os.sep)

        try:
            relazione = parts[-1]
            relazioni = parts[-2]   # "Relazioni tecniche"
            commessa = parts[-3]
            cliente = parts[-4]

            self.client_label.setText(f"Cliente: {cliente}")
            self.project_label.setText(f"Commessa: {commessa}")
            self.relation_label.setText(f"Relazione: {relazione}")

        except Exception:
            self.client_label.setText("Cliente: -")
            self.project_label.setText("Commessa: -")
            self.relation_label.setText("Relazione: -")

    def open_beta_tester_note_dialog(self):
        import urllib.request
        import json as _json

        dialog = QDialog(self)
        dialog.setWindowTitle("Nota beta tester")
        dialog.resize(620, 460)

        layout = QVBoxLayout(dialog)

        # --- Riepilogo dati utente (sola lettura) ---
        info_form = QFormLayout()
        id_utente_val = str(self.settings.value("license/id_utente", "N.D."))
        fingerprint_val = self.license_manager.get_machine_fingerprint()
        tipo_utente_val = self.settings.value("license/tipo_utente", "N.D.")
        tipo_licenza_val = self.settings.value("license/tipo_licenza", "N.D.")

        info_form.addRow("ID utente:", QLabel(id_utente_val))
        info_form.addRow("Fingerprint:", QLabel(fingerprint_val))
        info_form.addRow("Tipo utente:", QLabel(tipo_utente_val))
        info_form.addRow("Tipo licenza:", QLabel(tipo_licenza_val))
        layout.addLayout(info_form)

        # --- Nota ---
        note_edit = QTextEdit()
        note_edit.setPlaceholderText("Descrivi problema, suggerimento o comportamento osservato...")
        layout.addWidget(QLabel("Nota"))
        layout.addWidget(note_edit)

        # --- Pulsanti ---
        buttons = QHBoxLayout()
        btn_send = QPushButton("Invia nota ad ADMIN")
        btn_cancel = QPushButton("Annulla")
        buttons.addStretch()
        buttons.addWidget(btn_send)
        buttons.addWidget(btn_cancel)
        layout.addLayout(buttons)

        def invia_nota():
            note = note_edit.toPlainText().strip()
            if not note:
                QMessageBox.warning(dialog, "Nota mancante", "Scrivi prima una nota da inviare.")
                return

            chiave = self.license_manager.get_license_key().strip()
            if not chiave:
                QMessageBox.warning(dialog, "Licenza mancante", "Chiave licenza non configurata.")
                return

            server_url = self.license_manager.get_server_url().rstrip("/")
            url = f"{server_url}/api/feedback_beta.php"

            payload = _json.dumps({
                "chiave": chiave,
                "fingerprint": self.license_manager.get_machine_fingerprint(),
                "nome_macchina": platform.node(),
                "versione_sw": "EDITOR TECNICO",
                "sistema": platform.platform(),
                "nota": note
            }).encode("utf-8")

            btn_send.setEnabled(False)
            try:
                req = urllib.request.Request(
                    url,
                    data=payload,
                    headers={"Content-Type": "application/json"},
                    method="POST"
                )
                with urllib.request.urlopen(req, timeout=10) as resp:
                    data = _json.loads(resp.read().decode("utf-8"))
            except Exception as e:
                btn_send.setEnabled(True)
                QMessageBox.warning(
                    dialog,
                    "Invio non riuscito",
                    f"Impossibile contattare il server: {e}"
                )
                return

            esito = data.get("esito", "errore")
            messaggio = data.get("messaggio", "")

            if esito == "ok":
                QMessageBox.information(dialog, "Nota inviata", messaggio or "Nota inviata correttamente.")
                dialog.accept()
            else:
                btn_send.setEnabled(True)
                QMessageBox.warning(dialog, "Invio non riuscito", messaggio or "Errore durante l'invio.")

        btn_send.clicked.connect(invia_nota)
        btn_cancel.clicked.connect(dialog.reject)
        dialog.exec()

    def open_shortcuts_dialog(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Shortcut da tastiera")
        dialog.resize(520, 420)

        layout = QVBoxLayout(dialog)

        text = QTextEdit()
        text.setReadOnly(True)
        text.setHtml("""
        <h2>Shortcut da tastiera</h2>

        <h3>Gestione relazioni</h3>
        <p><b>Ctrl+N</b> — Nuova relazione</p>
        <p><b>Ctrl+O</b> — Apri relazione</p>
        <p><b>Ctrl+S</b> — Salva relazione</p>
        <p><b>Ctrl+Shift+S</b> — Salva con nome</p>
        <p><b>Ctrl+Alt+S</b> — Salva come template</p>

        <h3>Editor testo</h3>
        <p><b>Ctrl+Shift+F</b> — Apri editor esteso</p>
        <p><b>Ctrl+B</b> — Grassetto</p>
        <p><b>Ctrl+I</b> — Corsivo</p>
        <p><b>Ctrl+U</b> — Sottolineato</p>
        <p><b>Ctrl+.</b> — Apice</p>
        <p><b>Ctrl+,</b> — Pedice</p>
        <p><b>Ctrl+Shift+7</b> — Elenco puntato</p>
        <p><b>Ctrl+Shift+8</b> — Elenco numerato</p>
        <p><b>Ctrl+J</b> — Giustifica testo</p>
        <p><b>Ctrl+E</b> — Centra testo</p>
        <p><b>Ctrl+T</b> — Inserisci blocco testo</p>
        <p><b>Click su Navigatore Blocchi</b> — Vai al blocco selezionato</p>
        <p><b>Doppio click su Navigatore Blocchi</b> — Vai al blocco e attiva il campo principale</p>
        <p><b>Drag & Drop su Navigatore Blocchi</b> — Riordina i blocchi del nodo selezionato</p>
        <p><b>Ctrl+Shift+L</b> — Inserisci allegato</p>
        <p><b>TAB</b> — Aumenta livello elenco</p>
        <p><b>Shift+TAB</b> — Riduce livello elenco</p>

        <h3>Navigazione immagini</h3>
        <p><b>Freccia SINISTRA</b> — Immagine precedente</p>
        <p><b>Freccia DESTRA</b> — Immagine successiva</p>

        <h3>Strumenti</h3>
        <p><b>Ctrl+Shift+A</b> — Assistente AI</p>

        <h3>Applicazione</h3>
        <p><b>Ctrl+Q</b> — Chiudi applicazione</p>
        """)
        layout.addWidget(text)

        def export_shortcuts_pdf():
            file_path, _ = QFileDialog.getSaveFileName(
                dialog,
                "Esporta shortcut in PDF",
                "shortcut_editor_tecnico.pdf",
                "PDF (*.pdf)"
            )

            if not file_path:
                return

            if not file_path.lower().endswith(".pdf"):
                file_path += ".pdf"

            try:
                doc = BaseDocTemplate(
                    file_path,
                    pagesize=A4,
                    rightMargin=2 * cm,
                    leftMargin=2 * cm,
                    topMargin=2 * cm,
                    bottomMargin=2 * cm,
                )

                frame = Frame(
                    doc.leftMargin,
                    doc.bottomMargin,
                    doc.width,
                    doc.height,
                    id="normal"
                )

                doc.addPageTemplates([PageTemplate(id="main", frames=[frame])])

                styles = getSampleStyleSheet()

                title_style = ParagraphStyle(
                    "ShortcutTitle",
                    parent=styles["Title"],
                    alignment=TA_CENTER,
                    fontSize=18,
                    spaceAfter=18,
                )

                section_style = ParagraphStyle(
                    "ShortcutSection",
                    parent=styles["Heading2"],
                    fontSize=13,
                    spaceBefore=10,
                    spaceAfter=6,
                )

                body_style = ParagraphStyle(
                    "ShortcutBody",
                    parent=styles["BodyText"],
                    fontSize=10,
                    leading=13,
                    spaceAfter=4,
                )

                story = []

                story.append(Paragraph("Shortcut da tastiera", title_style))

                shortcuts = [
                    ("Gestione relazioni", [
                        ("Ctrl+N", "Nuova relazione"),
                        ("Ctrl+O", "Apri relazione"),
                        ("Ctrl+S", "Salva relazione"),
                        ("Ctrl+Shift+S", "Salva con nome"),
                        ("Ctrl+Alt+S", "Salva come template"),
                    ]),
                    ("Editor testo", [
                        ("Ctrl+Shift+F", "Apri editor esteso"),
                        ("Ctrl+B", "Grassetto"),
                        ("Ctrl+I", "Corsivo"),
                        ("Ctrl+U", "Sottolineato"),
                        ("Ctrl+.", "Apice"),
                        ("Ctrl+,", "Pedice"),
                        ("Ctrl+Shift+7", "Elenco puntato"),
                        ("Ctrl+Shift+8", "Elenco numerato"),
                        ("Ctrl+J", "Giustifica testo"),
                        ("Ctrl+E", "Centra testo"),
                        ("Ctrl+T", "Inserisci blocco testo nel nodo selezionato"),
                        ("Click su Navigatore Blocchi", "Vai al blocco selezionato"),
                        ("Doppio click su Navigatore Blocchi", "Vai al blocco e attiva il campo principale"),
                        ("Drag & Drop su Navigatore Blocchi", "Riordina i blocchi del nodo selezionato"),
                        ("Ctrl+Shift+L", "Inserisci allegato"),
                        ("TAB", "Aumenta livello elenco"),
                        ("Shift+TAB", "Riduce livello elenco"),
                    ]),
                    ("Strumenti", [
                        ("Ctrl+Shift+A", "Assistente AI"),
                    ]),
                    ("Applicazione", [
                        ("Ctrl+Q", "Chiudi applicazione"),
                    ]),
                ]

                for section_title, rows in shortcuts:
                    story.append(Paragraph(section_title, section_style))

                    for shortcut, description in rows:
                        story.append(
                            Paragraph(
                                f"<b>{shortcut}</b> - {description}",
                                body_style
                            )
                        )

                doc.build(story)

                QMessageBox.information(
                    dialog,
                    "PDF generato",
                    f"Elenco shortcut esportato correttamente:\n\n{file_path}"
                )

            except Exception as exc:
                QMessageBox.critical(
                    dialog,
                    "Errore PDF",
                    f"Impossibile generare il PDF:\n\n{exc}"
                )


        btn_row = QHBoxLayout()

        btn_print = QPushButton("Stampa PDF")
        btn_close = QPushButton("Chiudi")

        btn_print.clicked.connect(export_shortcuts_pdf)
        btn_close.clicked.connect(dialog.accept)

        btn_row.addStretch()
        btn_row.addWidget(btn_print)
        btn_row.addWidget(btn_close)
        layout.addLayout(btn_row)

        dialog.exec()

    def open_latex_guide_dialog(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Guida LaTeX")
        dialog.resize(900, 650)

        dialog.setStyleSheet("""
            QDialog {
                background: #f2f2f2;
                color: black;
            }

            QTextEdit {
                background: white;
                color: black;
                border: 1px solid #cfcfcf;
                padding: 8px;
                font-family: Consolas;
                font-size: 13px;
            }

            QPushButton {
                background: #fafafa;
                border: 1px solid #d0d0d0;
                border-radius: 4px;
                padding: 6px 12px;
                color: black;
            }

            QPushButton:hover {
                background: #ececec;
            }
        """)

        layout = QVBoxLayout(dialog)

        title = QLabel("Guida rapida LaTeX")
        title.setStyleSheet("""
            font-size: 18px;
            font-weight: bold;
            color: black;
        """)

        layout.addWidget(title)

        guide_text = QTextEdit()
        guide_text.setReadOnly(True)

        guide_text.setPlainText(
    r"""
    FRAZIONI
    ----------------------------
    \frac{a}{b}

    APICI E PEDICI
    ----------------------------
    x^2
    x_i

    RADICI
    ----------------------------
    \sqrt{x}
    \sqrt[n]{x}

    SOMMATORIE
    ----------------------------
    \sum_{i=1}^{n}

    INTEGRALI
    ----------------------------
    \int_a^b
    \iint
    \oint

    ACCENTI E VETTORI
    ----------------------------
    \bar{x}
    \vec{F}
    \hat{x}

    PARENTESI
    ----------------------------
    \left( \frac{a}{b} \right)

    MATRICI
    ----------------------------
    \begin{bmatrix}
    a & b \\
    c & d
    \end{bmatrix}

    SIMBOLI GRECI
    ----------------------------
    \alpha
    \beta
    \gamma
    \sigma
    \Delta

    OPERATORI
    ----------------------------
    \times
    \pm
    \leq
    \geq
    \neq
    \approx

    FRECCE
    ----------------------------
    \rightarrow
    \Rightarrow
    \leftrightarrow

    MULTIRIGA
    ----------------------------
    M = \frac{qL^2}{8}

    M = \frac{12.5 \cdot 5.2^2}{8}

    M = 42.25 \, kNm
    """
        )

        layout.addWidget(guide_text)

        btn_print = QPushButton("Stampa PDF")
        btn_close = QPushButton("Chiudi")

        btn_close.clicked.connect(dialog.accept)

        def export_latex_guide_pdf():
            file_path, _ = QFileDialog.getSaveFileName(
                dialog,
                "Esporta guida LaTeX in PDF",
                "guida_latex_editor_tecnico.pdf",
                "PDF (*.pdf)"
            )

            if not file_path:
                return

            if not file_path.lower().endswith(".pdf"):
                file_path += ".pdf"

            try:
                doc = BaseDocTemplate(
                    file_path,
                    pagesize=A4,
                    rightMargin=2 * cm,
                    leftMargin=2 * cm,
                    topMargin=2 * cm,
                    bottomMargin=2 * cm,
                )

                frame = Frame(
                    doc.leftMargin,
                    doc.bottomMargin,
                    doc.width,
                    doc.height,
                    id="normal"
                )

                doc.addPageTemplates([
                    PageTemplate(id="main", frames=[frame])
                ])

                styles = getSampleStyleSheet()

                title_style = ParagraphStyle(
                    "LatexGuideTitle",
                    parent=styles["Title"],
                    alignment=TA_CENTER,
                    fontSize=18,
                    spaceAfter=18,
                )

                body_style = ParagraphStyle(
                    "LatexGuideBody",
                    parent=styles["BodyText"],
                    fontName="Courier",
                    fontSize=9,
                    leading=12,
                    spaceAfter=4,
                )

                story = []
                story.append(Paragraph("Guida rapida LaTeX", title_style))

                for line in guide_text.toPlainText().splitlines():
                    safe_line = (
                        line.replace("&", "&amp;")
                            .replace("<", "&lt;")
                            .replace(">", "&gt;")
                    )

                    if not safe_line.strip():
                        story.append(Spacer(1, 6))
                    else:
                        story.append(Paragraph(safe_line, body_style))

                doc.build(story)

                QMessageBox.information(
                    dialog,
                    "PDF generato",
                    f"Guida LaTeX esportata correttamente:\n\n{file_path}"
                )

            except Exception as exc:
                QMessageBox.critical(
                    dialog,
                    "Errore PDF",
                    f"Impossibile generare la guida LaTeX:\n\n{exc}"
                )

        btn_print.clicked.connect(export_latex_guide_pdf)

        btn_row = QHBoxLayout()
        btn_row.addStretch()
        btn_row.addWidget(btn_print)
        btn_row.addWidget(btn_close)

        layout.addLayout(btn_row)

        dialog.exec()

    def build_top_toolbar(self):
        toolbar = QToolBar("Barra principale")
        toolbar.setMovable(False)
        self.addToolBar(toolbar)
        act_quit = QAction("Esci", self)
        act_quit.setShortcut("Ctrl+Q")
        act_quit.triggered.connect(self.close)
        self.addAction(act_quit)

        # --- Azioni gestione relazioni ---
        act_new = QAction("Nuova Relazione", self)
        act_open = QAction("Apri Relazione", self)
        act_save = QAction("Salva Relazione", self)
        act_save_as = QAction("Salva con Nome", self)
        act_save_copy = QAction(
            "Salva in Archivio di Sistema",
            self
        )

        act_new.setShortcut("Ctrl+N")
        act_open.setShortcut("Ctrl+O")
        act_save.setShortcut("Ctrl+S")
        act_save_as.setShortcut("Ctrl+Shift+S")
        act_save_copy.setShortcut("Ctrl+Alt+S")

        act_new.triggered.connect(self.new_project)
        act_open.triggered.connect(self.open_project)
        act_save.triggered.connect(self.save_project)
        act_save_as.triggered.connect(self.save_project_as_new_relation)
        act_save_copy.triggered.connect(self.save_project_copy_to_system_archive)

        relation_menu = QMenu("GESTIONE RELAZIONI", self)
        relation_menu.addAction(act_new)
        relation_menu.addAction(act_open)
        relation_menu.addSeparator()
        relation_menu.addAction(act_save)
        relation_menu.addAction(act_save_as)

        btn_relation_menu = QPushButton("GESTIONE RELAZIONI")
        btn_relation_menu.setMenu(relation_menu)
        toolbar.addWidget(btn_relation_menu)

        toolbar.addSeparator()

        # --- Altre azioni toolbar ---
        act_verify = QAction("Verifica integrità", self)
        act_export = QAction("Esporta", self)
        act_ai = QAction("Assistente AI", self)
        act_api_connections = QAction("Connessioni API", self)
        act_shortcuts = QAction("Shortcut da tastiera", self)
        act_latex_guide = QAction("Guida LaTeX", self)
        act_beta_note = QAction("Invia nota beta tester", self)
        self.act_beta_note = act_beta_note  # salva riferimento
        tipo_licenza = self.settings.value("license/tipo_licenza", "")
        tipo_utente = self.settings.value("license/tipo_utente", "")

        act_beta_note.setEnabled(
            tipo_licenza == "beta"
        )
        act_ai.setShortcut("Ctrl+Shift+A")
        act_allegati = QAction("Apri Allegati", self)
        act_manage_allegati = QAction("Gestisci Allegati", self)
        act_layout = QAction("Layout pagina", self)

        tools_menu = QMenu("STRUMENTI", self)
        tools_menu.addAction(act_ai)
        tools_menu.addSeparator()
        tools_menu.addAction(act_beta_note)
        tools_menu.addAction(act_shortcuts)
        tools_menu.addAction(act_latex_guide)

        btn_tools_menu = QPushButton("STRUMENTI")
        btn_tools_menu.setMenu(tools_menu)
        toolbar.addWidget(btn_tools_menu)

        # --- Menu impostazioni ---
        act_settings_workspace = QAction("Workspace", self)
        act_settings_workspace.triggered.connect(self.select_workspace)

        act_license = QAction("Licenza", self)
        act_license.triggered.connect(self.open_license_dialog)

        settings_menu = QMenu("IMPOSTAZIONI", self)
        settings_menu.addAction(act_settings_workspace)
        settings_menu.addAction(act_api_connections)
        settings_menu.addSeparator()
        settings_menu.addAction(act_license)

        btn_settings_menu = QPushButton("IMPOSTAZIONI")
        btn_settings_menu.setMenu(settings_menu)
        toolbar.addWidget(btn_settings_menu)

        # --- Menu archivio di sistema ---
        system_archive_menu = QMenu("ARCHIVIO DI SISTEMA", self)

        act_open_archive_document = QAction(
            "Apri documento archivio",
            self
        )

        act_create_archive_document = QAction(
            "Crea documento archivio",
            self
        )
        
        act_open_archive_document.triggered.connect(
            self.open_system_archive_document
        )

        act_create_archive_document.triggered.connect(
            self.create_system_archive_document
        )

        system_archive_menu.addAction(act_open_archive_document)
        system_archive_menu.addAction(act_create_archive_document)
        system_archive_menu.addSeparator()
        system_archive_menu.addAction(act_save_copy)

        btn_system_archive_menu = QPushButton("ARCHIVIO DI SISTEMA")
        btn_system_archive_menu.setMenu(system_archive_menu)
        toolbar.addWidget(btn_system_archive_menu)

        # --- Pulsante MANUALI (apre la pagina manuali sul sito) ---
        btn_manuali = QPushButton("MANUALI")
        btn_manuali.clicked.connect(self.open_manuali_web)
        toolbar.addWidget(btn_manuali)

        toolbar.addSeparator()

        act_verify.triggered.connect(self.verify_project_integrity)
        act_export.triggered.connect(self.export_document)
        act_ai.triggered.connect(self.open_ai_dialog)
        act_api_connections.triggered.connect(self.open_api_connections_dialog)
        act_shortcuts.triggered.connect(self.open_shortcuts_dialog)
        act_latex_guide.triggered.connect(self.open_latex_guide_dialog)
        act_beta_note.triggered.connect(self.open_beta_tester_note_dialog)
        act_allegati.triggered.connect(self.open_allegati_folder)
        act_manage_allegati.triggered.connect(self.manage_attachments)
        act_layout.triggered.connect(self.open_layout_dialog)

        self.workspace_dependent_actions = [
            act_new,
            act_open,
            act_save,
            act_save_as,
            act_save_copy,
            act_ai,
        ]

        self.update_workspace_enabled_state()

    def build_status_bar(self):
        status = QStatusBar()
        self.setStatusBar(status)
        status.showMessage("Pronto")

    def build_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        outer_layout = QVBoxLayout(central)
        outer_layout.setContentsMargins(10, 4, 10, 10)
        outer_layout.setSpacing(0)

        # --- Workspace ---
        self.workspace_label = QLabel()
        self.workspace_label.setToolTip("Doppio click per selezionare il workspace")
        self.workspace_label.mouseDoubleClickEvent = lambda event: self.select_workspace()
        self.workspace_label.setFixedHeight(24)
        self.workspace_label.setStyleSheet(
            "QLabel { color: #555; background: #f7f7f7; border: 1px solid #ddd; padding-left: 6px; }"
        )
        outer_layout.addWidget(self.workspace_label)

        # --- Cliente / Commessa / Relazione ---
        info_row = QHBoxLayout()
        info_row.setContentsMargins(0, 0, 0, 0)
        info_row.setSpacing(6)

        self.client_label = QLabel("Cliente: -")
        self.project_label = QLabel("Commessa: -")
        self.relation_label = QLabel("Relazione: -")

        for lbl in (self.client_label, self.project_label, self.relation_label):
            lbl.setFixedHeight(24)
            lbl.setStyleSheet(
                "QLabel { color: #555; background: #f7f7f7; border: 1px solid #ddd; padding-left: 6px; }"
            )
            lbl.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
            info_row.addWidget(lbl)

        info_row.addStretch()

        info_widget = QWidget()
        info_widget.setFixedHeight(26)
        info_widget.setLayout(info_row)
        outer_layout.addWidget(info_widget)

        # --- Azioni relazione corrente ---
        relation_actions_row = QHBoxLayout()
        relation_actions_row.setContentsMargins(0, 2, 0, 0)
        relation_actions_row.setSpacing(4)

        btn_verify = QPushButton("Verifica integrità")
        btn_export = QPushButton("Esporta")
        btn_open_attachments = QPushButton("Apri Allegati")
        btn_manage_attachments = QPushButton("Gestisci Allegati")
        btn_layout = QPushButton("Layout pagina")
        btn_cover = QPushButton("Copertina documento")

        btn_verify.clicked.connect(self.verify_project_integrity)
        btn_export.clicked.connect(self.export_document)
        btn_open_attachments.clicked.connect(self.open_allegati_folder)
        btn_manage_attachments.clicked.connect(self.manage_attachments)
        btn_layout.clicked.connect(self.open_layout_dialog)
        btn_cover.clicked.connect(self.open_cover_dialog)

        for btn in (
            btn_cover,
            btn_verify,
            btn_export,
            btn_open_attachments,
            btn_manage_attachments,
            btn_layout,
        ):
            btn.setFixedSize(140, 26)
            style_small_button(btn)

        relation_actions_row.addStretch()

        relation_actions_row.addWidget(btn_cover)
        relation_actions_row.addWidget(btn_verify)
        relation_actions_row.addWidget(btn_layout)
        relation_actions_row.addWidget(btn_manage_attachments)
        relation_actions_row.addWidget(btn_open_attachments)
        relation_actions_row.addWidget(btn_export)

        relation_actions_row.addStretch()

        outer_layout.addSpacing(4)

        separator = QFrame()
        separator.setFrameShape(QFrame.HLine)
        separator.setFrameShadow(QFrame.Plain)
        separator.setStyleSheet("""
        QFrame {
            background-color: #d0d0d0;
            border: none;
        }
        """)
        separator.setFixedHeight(1)
        separator.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

        outer_layout.addWidget(separator)

        outer_layout.addSpacing(4)

        relation_actions_widget = QWidget()
        relation_actions_widget.setFixedHeight(28)
        relation_actions_widget.setLayout(relation_actions_row)
        outer_layout.addWidget(relation_actions_widget)

        self.workspace_dependent_actions.extend([
            btn_cover,
            btn_verify,
            btn_export,
            btn_open_attachments,
            btn_manage_attachments,
            btn_layout,
        ])

        outer_layout.addSpacing(4)

        splitter = QSplitter(Qt.Horizontal)
        outer_layout.addWidget(splitter)

        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(8)

        title = QLabel("Struttura documento")
        title.setStyleSheet("font-size: 18px; font-weight: 600;")
        left_layout.addWidget(title)

        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText("Filtra capitoli e paragrafi...")
        self.search_edit.textChanged.connect(self.filter_tree)
        left_layout.addWidget(self.search_edit)

        self.tree = QTreeWidget()
        self.tree.setHeaderHidden(True)
        self.tree.setIndentation(18)
        self.tree.setSelectionMode(QTreeWidget.SingleSelection)
        self.tree.setDragEnabled(True)
        self.tree.viewport().setAcceptDrops(True)
        self.tree.setDropIndicatorShown(True)
        self.tree.setDragDropMode(QTreeWidget.InternalMove)
        self.tree.setDefaultDropAction(Qt.MoveAction)
        self.tree.setContextMenuPolicy(Qt.CustomContextMenu)
        self.tree.customContextMenuRequested.connect(self.open_context_menu)
        self.tree.itemSelectionChanged.connect(self.on_tree_selection_changed)
        self.tree.model().rowsMoved.connect(self.on_tree_rows_moved)
        left_layout.addWidget(self.tree)

        right_panel = QWidget()
        right_layout = QHBoxLayout(right_panel)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_layout.setSpacing(8)

        editor_column = QWidget()
        editor_column_layout = QVBoxLayout(editor_column)
        editor_column_layout.setContentsMargins(0, 0, 0, 0)
        editor_column_layout.setSpacing(8)

        editor_title = QLabel("Contenuto")
        editor_title.setStyleSheet("font-size: 18px; font-weight: 600;")
        editor_column_layout.addWidget(editor_title)

        self.selected_label = QLabel("Seleziona un nodo nella struttura")
        self.selected_label.setStyleSheet("color: #666;")
        editor_column_layout.addWidget(self.selected_label)

        self.block_editor = BlockEditorWidget(self)

        self.blocks_navigator = BlocksNavigatorWidget()
        self.blocks_navigator.setMinimumWidth(260)
        self.blocks_navigator.setStyleSheet("""
            QListWidget {
                background: #fafafa;
                border: 1px solid #d6d6d6;
                color: black;
            }
            QListWidget::item {
                padding: 6px;
            }
            QListWidget::item:selected {
                background: #dcecff;
                color: black;
            }
        """)
        self.blocks_navigator.itemClicked.connect(self.go_to_block_from_navigator)
        self.blocks_navigator.itemDoubleClicked.connect(self.focus_block_from_navigator)
        self.blocks_navigator.orderChanged.connect(
            self.on_blocks_navigator_order_changed
        )

        self.navigator_focus_timer = QTimer(self)
        self.navigator_focus_timer.setSingleShot(True)
        self._navigator_focus_token = 0
        self._navigator_selected_widget = None

        self.blocks_navigator.setDragEnabled(True)
        self.blocks_navigator.setAcceptDrops(True)
        self.blocks_navigator.setDropIndicatorShown(True)

        editor_column_layout.addWidget(self.block_editor, 1)
        right_layout.addWidget(editor_column, 1)

        blocks_nav_dock_content = QWidget()
        blocks_nav_dock_layout = QVBoxLayout(blocks_nav_dock_content)
        blocks_nav_dock_layout.setContentsMargins(8, 8, 8, 8)
        blocks_nav_dock_layout.setSpacing(8)

        blocks_nav_header = QHBoxLayout()

        blocks_nav_title = QLabel("Navigatore Blocchi")
        blocks_nav_title.setStyleSheet("font-size: 18px; font-weight: 600;")

        self.btn_pin_navigator = QPushButton("📌")
        self.btn_pin_navigator.setFixedSize(28, 24)
        self.btn_pin_navigator.setToolTip("Ancora / sgancia navigatore blocchi")
        style_icon_button(self.btn_pin_navigator)
        self.btn_pin_navigator.clicked.connect(self.toggle_navigator_pinned)

        blocks_nav_header.addWidget(blocks_nav_title)
        blocks_nav_header.addStretch()
        blocks_nav_header.addWidget(self.btn_pin_navigator)

        blocks_nav_dock_layout.addLayout(blocks_nav_header)
        blocks_nav_dock_layout.addWidget(self.blocks_navigator, 1)

        self.navigator_pinned = False
        self.navigator_dragging = False
        self.navigator_close_timer = QTimer(self)
        self.navigator_close_timer.setSingleShot(True)
        self.navigator_close_timer.timeout.connect(self.close_navigator_if_inactive)

        self.navigator_side_container = QWidget()
        self.navigator_side_layout = QHBoxLayout(self.navigator_side_container)
        self.navigator_side_layout.setContentsMargins(0, 0, 0, 0)
        self.navigator_side_layout.setSpacing(0)

        self.navigator_placeholder = QLabel("◀")
        self.navigator_placeholder.setFixedWidth(18)
        self.navigator_placeholder.setAlignment(Qt.AlignCenter)
        self.navigator_placeholder.setStyleSheet("""
            QLabel {
                background: #dcdcdc;
                color: #444;
                font-weight: bold;
                border-left: 1px solid #bcbcbc;
            }
        """)

        self.navigator_panel = blocks_nav_dock_content
        self.navigator_placeholder.enterEvent = self.on_navigator_placeholder_enter
        self.navigator_side_container.installEventFilter(self)
        self.navigator_panel.installEventFilter(self)

        self.navigator_panel.setFixedWidth(0)

        self.navigator_side_layout.addWidget(self.navigator_placeholder)
        self.navigator_side_layout.addWidget(self.navigator_panel)

        right_layout.addWidget(self.navigator_side_container, 0)

        splitter.addWidget(left_panel)
        splitter.addWidget(right_panel)
        splitter.setSizes([350, 830])

    def project_assets_dir(self):
        if not self.project_dir:
            return None
        assets = Path(self.project_dir) / "assets"
        assets.mkdir(parents=True, exist_ok=True)
        return assets

    def assets_texts_dir(self):
        assets = self.project_assets_dir()
        if assets is None:
            return None
        path = assets / "texts"
        path.mkdir(parents=True, exist_ok=True)
        return path

    def block_text_file_path(self, block_id):
        texts_dir = self.assets_texts_dir()
        if texts_dir is None:
            return None
        return texts_dir / f"{block_id}.txt"

    def save_text_block_to_file(self, block):
        block_id = block.get("id")
        if not block_id:
            return block
        path = self.block_text_file_path(block_id)
        if path is None:
            return block
        path.parent.mkdir(parents=True, exist_ok=True)
        content = block.get("text", "")
        path.write_text(content, encoding="utf-8")
        block = deepcopy(block)
        block["text_path"] = f"assets/texts/{block_id}.txt"
        return block

    def load_text_block_from_file(self, block):
        block = deepcopy(block)
        text_path = block.get("text_path", "")
        if text_path and self.project_dir:
            abs_path = Path(self.project_dir) / text_path
            if abs_path.exists():
                block["text"] = abs_path.read_text(encoding="utf-8")
            else:
                block["text"] = ""
        else:
            block["text"] = block.get("text", "")
        return block

    def remove_text_file_for_block(self, block):
        block_id = block.get("id")
        if not block_id or not self.project_dir:
            return
        path = self.block_text_file_path(block_id)
        if path and path.exists():
            try:
                path.unlink()
            except Exception:
                pass

    def collect_referenced_image_paths(self):
        refs = set()

        def collect_from_blocks(blocks):
            for block in blocks or []:
                if block.get("type") == "image":
                    p = block.get("path", "")
                    if p:
                        refs.add(as_posix_path(p))

                elif block.get("type") == "images":
                    for img in block.get("images", []):
                        p = img.get("path", "")
                        if p:
                            refs.add(as_posix_path(p))

        def visit(item):
            blocks = item.data(0, CONTENT_ROLE) or []
            collect_from_blocks(blocks)

            for i in range(item.childCount()):
                visit(item.child(i))

        for i in range(self.tree.topLevelItemCount()):
            visit(self.tree.topLevelItem(i))

        cover_blocks = self.cover_config.get("blocks", [])
        collect_from_blocks(cover_blocks)

        return refs

    def collect_text_block_ids(self):
        ids = set()

        def visit(item):
            blocks = item.data(0, CONTENT_ROLE) or []
            for block in blocks:
                if block.get("type") == "text":
                    block_id = block.get("id")
                    if block_id:
                        ids.add(block_id)
            for i in range(item.childCount()):
                visit(item.child(i))

        for i in range(self.tree.topLevelItemCount()):
            visit(self.tree.topLevelItem(i))
        return ids

    def sync_all_text_blocks_to_files(self):
        if not self.project_dir:
            return
        texts_dir = self.assets_texts_dir()
        if texts_dir is None:
            return

        used_ids = set()

        def visit(item):
            blocks = self.ensure_blocks_defaults_list(item.data(0, CONTENT_ROLE) or [])
            for block in blocks:
                if block.get("type") == "text":
                    self.save_text_block_to_file(block)
                    block_id = block.get("id")
                    if block_id:
                        used_ids.add(block_id)
            for i in range(item.childCount()):
                visit(item.child(i))

        for i in range(self.tree.topLevelItemCount()):
            visit(self.tree.topLevelItem(i))

        for file in texts_dir.glob("*.txt"):
            if file.stem not in used_ids:
                try:
                    file.unlink()
                except Exception:
                    pass

    def sync_all_images_to_assets(self, source_project_dir=None):
        if not self.project_dir:
            return
        used_paths = set()

        def visit(item):
            blocks = self.ensure_blocks_defaults_list(item.data(0, CONTENT_ROLE) or [])
            updated_blocks = []
            for block in blocks:
                b = self.ensure_block_defaults(block)
                if b.get("type") == "image":
                    block_id = b.get("id")
                    b["path"] = self.store_image_in_assets(
                        b.get("path", ""),
                        block_id or new_block_id(),
                        source_project_dir=source_project_dir,
                    )
                    if as_posix_path(b.get("path", "")).startswith("assets/images/"):
                        used_paths.add(as_posix_path(b["path"]))
                elif b.get("type") == "images":
                    images = []
                    for img in self.ensure_group_images_defaults(b.get("images", [])):
                        img_copy = deepcopy(img)
                        img_id = img_copy.get("id") or new_block_id()
                        img_copy["id"] = img_id
                        img_copy["path"] = self.store_image_in_assets(
                            img_copy.get("path", ""),
                            img_id,
                            source_project_dir=source_project_dir,
                        )
                        if as_posix_path(img_copy.get("path", "")).startswith("assets/images/"):
                            used_paths.add(as_posix_path(img_copy["path"]))
                        images.append(img_copy)
                    b["images"] = images
                updated_blocks.append(b)
            item.setData(0, CONTENT_ROLE, updated_blocks)
            for i in range(item.childCount()):
                visit(item.child(i))

        for i in range(self.tree.topLevelItemCount()):
            visit(self.tree.topLevelItem(i))

        img_dir = Path(self.project_dir) / "assets" / "images"
        if img_dir.exists():
            for file in img_dir.iterdir():
                if file.is_file():
                    rel = f"assets/images/{file.name}"
                    if rel not in used_paths:
                        try:
                            file.unlink()
                        except Exception:
                            pass

    def prune_unused_asset_files(self):
        if not self.project_dir:
            return
        refs = self.collect_referenced_image_paths()
        img_dir = Path(self.project_dir) / "assets" / "images"
        if img_dir.exists():
            for file in img_dir.iterdir():
                if file.is_file():
                    rel = f"assets/images/{file.name}"
                    if rel not in refs:
                        try:
                            file.unlink()
                        except Exception:
                            pass

    def make_image_asset_name(self, source_path, preferred_id):
        ext = Path(source_path).suffix or ".png"
        return f"{preferred_id}{ext}"

    def store_image_in_assets(self, source_path, preferred_id, source_project_dir=None):
        if not source_path or not self.project_dir:
            return source_path
        img_dir = Path(self.project_dir) / "assets" / "images"
        img_dir.mkdir(parents=True, exist_ok=True)

        normalized = as_posix_path(source_path)
        if normalized.startswith("assets/images/"):
            existing_abs = self.resolve_runtime_path_with_base(normalized, self.project_dir)
            if os.path.exists(existing_abs):
                return normalized

        abs_source = Path(self.resolve_runtime_path_with_base(source_path, source_project_dir or self.project_dir))
        if not abs_source.exists():
            return source_path

        filename = self.make_image_asset_name(str(abs_source), preferred_id)
        dest = img_dir / filename
        try:
            if os.path.normpath(str(abs_source)) != os.path.normpath(str(dest)):
                shutil.copy2(abs_source, dest)
            return f"assets/images/{filename}"
        except Exception:
            return source_path

    def ingest_image_into_project(self, source_path):
        return source_path

    def clone_block_data(self, data):
        data = deepcopy(data)
        data["id"] = new_block_id()
        data["meta"] = deepcopy(data.get("meta", {}))
        if data.get("type") == "images":
            cloned_images = []
            for img in self.ensure_group_images_defaults(data.get("images", [])):
                img_copy = deepcopy(img)
                img_copy["id"] = new_block_id()
                cloned_images.append(img_copy)
            data["images"] = cloned_images
        return data

    def delete_image_asset_if_local(self, path):
        if not path or not self.project_dir:
            return
        normalized = as_posix_path(path)
        if not normalized.startswith("assets/images/"):
            return
        abs_path = Path(self.project_dir) / normalized
        if abs_path.exists():
            try:
                abs_path.unlink()
            except Exception:
                pass

    def mark_dirty(self):
        self.project_dirty = True
        self.update_window_title()

    def set_project_dir(self, project_dir):
        self.project_dir = project_dir
        self.document_path = os.path.join(project_dir, "document.json")
        os.makedirs(os.path.join(project_dir, "assets", "images"), exist_ok=True)
        os.makedirs(os.path.join(project_dir, "assets", "images", "cover"), exist_ok=True)
        os.makedirs(os.path.join(project_dir, "assets", "texts"), exist_ok=True)
        self.source_project_dir_for_save = None

    def assets_images_dir(self):
        if not self.project_dir:
            return None
        path = os.path.join(self.project_dir, "assets", "images")
        os.makedirs(path, exist_ok=True)
        return path

    def save_clipboard_image_to_assets(self):
        clipboard = QApplication.clipboard()
        image = clipboard.image()

        if image.isNull():
            return ""

        image_id = new_block_id()

        if self.project_dir:
            images_dir = self.assets_images_dir()
            file_path = os.path.join(images_dir, f"{image_id}.png")
            stored_path = f"assets/images/{image_id}.png"
        else:
            temp_dir = tempfile.gettempdir()
            file_path = os.path.join(temp_dir, f"{image_id}.png")
            stored_path = file_path

        if not image.save(file_path, "PNG"):
            return ""

        return stored_path

    def resolve_runtime_path(self, path):
        if not path:
            return ""
        if os.path.isabs(path):
            return path
        if self.project_dir:
            return os.path.normpath(os.path.join(self.project_dir, path))
        return path

    def resolve_runtime_path_with_base(self, path, base_dir=None):
        if not path:
            return ""
        if os.path.isabs(path):
            return path
        if base_dir:
            return os.path.normpath(os.path.join(base_dir, path))
        return self.resolve_runtime_path(path)

    def maybe_save_before_destructive_action(self):
        if not self.project_dirty:
            return True
        box = QMessageBox(self)
        box.setWindowTitle("Salvare le modifiche?")
        box.setText("Il progetto contiene modifiche non salvate.")
        box.setStandardButtons(QMessageBox.Save | QMessageBox.Discard | QMessageBox.Cancel)
        box.setDefaultButton(QMessageBox.Save)
        choice = box.exec()
        if choice == QMessageBox.Cancel:
            return False
        if choice == QMessageBox.Save:
            return self.save_project()
        return True

    def populate_empty_structure(self):
        self.tree.clear()
        self.current_item = None
        self.selected_label.setText("Seleziona un nodo nella struttura")
        self.block_editor.clear_blocks()
        self.block_editor.set_enabled_for_node(False)
        self.figure_reference_map = {}
        self.equation_reference_map = {}
        self.source_project_dir_for_save = None

    def create_item(self, title):
        item = QTreeWidgetItem([title])
        item.setData(0, TITLE_ROLE, title)
        item.setData(0, CONTENT_ROLE, [])
        return item

    def ensure_block_defaults(self, block):
        block = deepcopy(block)

        if "id" not in block or not block["id"]:
            block["id"] = new_block_id()

        if "meta" not in block or not isinstance(block["meta"], dict):
            block["meta"] = {}

        if "collapsed" not in block:
            block["collapsed"] = False

        if block.get("type") == "images":
            images = []
            for img in block.get("images", []):
                img = deepcopy(img)
                if "id" not in img or not img["id"]:
                    img["id"] = new_block_id()
                images.append(img)
            block["images"] = images

        elif block.get("type") == "table":
            if "spans" not in block or not isinstance(block.get("spans"), list):
                block["spans"] = []

            if "cell_formats" not in block or not isinstance(block.get("cell_formats"), dict):
                block["cell_formats"] = {}
        elif block.get("type") == "equation":
            if "numbering_mode" not in block:
                block["numbering_mode"] = "none"

            if block.get("numbering_mode") not in ("none", "number", "number_caption"):
                block["numbering_mode"] = "none"

            if "caption" not in block:
                block["caption"] = ""
        return block

    def ensure_blocks_defaults_list(self, blocks):
        return [self.ensure_block_defaults(b) for b in (blocks or [])]

    def ensure_group_images_defaults(self, images):
        normalized = []
        for img in (images or []):
            item = deepcopy(img)
            if "id" not in item or not item["id"]:
                item["id"] = new_block_id()
            if "caption" not in item:
                item["caption"] = ""
            normalized.append(item)
        return normalized

    def save_current_item_content(self):
        if self.current_item is None:
            return
        exported = deepcopy(self.block_editor.export_blocks())
        exported = self.ensure_blocks_defaults_list(exported)
        self.current_item.setData(0, CONTENT_ROLE, exported)

    def add_chapter(self):
        self.save_current_item_content()
        title, ok = QInputDialog.getText(self, "Nuovo capitolo", "Titolo capitolo:")
        if not ok or not title.strip():
            return
        item = self.create_item(title.strip())
        self.tree.addTopLevelItem(item)
        self.renumber_tree()
        self.tree.setCurrentItem(item)
        self.mark_dirty()
        self.statusBar().showMessage("Capitolo aggiunto", 2000)

    def add_paragraph(self):
        self.save_current_item_content()
        current = self.tree.currentItem()
        if current is None:
            self.statusBar().showMessage("Seleziona un nodo prima di aggiungere un paragrafo", 2500)
            return
        title, ok = QInputDialog.getText(self, "Nuovo paragrafo", "Titolo paragrafo:")
        if not ok or not title.strip():
            return
        item = self.create_item(title.strip())
        current.addChild(item)
        current.setExpanded(True)
        self.renumber_tree()
        self.tree.setCurrentItem(item)
        self.mark_dirty()
        self.statusBar().showMessage("Paragrafo aggiunto", 2000)


    def on_tree_rows_moved(self, *args):
        self.renumber_tree()
        self.update_figure_references()

    def renumber_tree(self, *args):
        for i in range(self.tree.topLevelItemCount()):
            item = self.tree.topLevelItem(i)
            self.renumber_item(item, [i + 1])

    def renumber_item(self, item, numbers):
        base_title = item.data(0, TITLE_ROLE)
        if not base_title:
            base_title = item.text(0)
            item.setData(0, TITLE_ROLE, base_title)
        prefix = ".".join(str(n) for n in numbers)
        item.setText(0, f"{prefix}. {base_title}")
        for i in range(item.childCount()):
            child = item.child(i)
            self.renumber_item(child, numbers + [i + 1])

    def block_navigator_icon(self, widget):
        if isinstance(widget, TextBlockWidget):
            return "📝"
        if isinstance(widget, TableBlockWidget):
            return "▦"
        if isinstance(widget, EquationBlockWidget):
            return "∑"
        if isinstance(widget, ImageBlockWidget):
            return "🖼"
        if isinstance(widget, MultiImageBlockWidget):
            return "🖼🖼"
        return "•"

    def eventFilter(self, obj, event):
        if (
            hasattr(self, "navigator_panel")
            and obj == self.navigator_panel
            and event.type() == QEvent.Leave
        ):
            if not self.navigator_pinned:
                self.navigator_close_timer.start(300)

        return super().eventFilter(obj, event)

    def close_navigator_if_inactive(self):
        if self.navigator_pinned:
            return

        if self.navigator_dragging:
            return

        pos = QCursor.pos()
        local_pos = self.navigator_side_container.mapFromGlobal(pos)

        if self.navigator_side_container.rect().contains(local_pos):
            return

        self.navigator_panel.setFixedWidth(0)

    def toggle_navigator_pinned(self):
        self.navigator_pinned = not self.navigator_pinned

        if self.navigator_pinned:
            self._animate_navigator(280)
            self.btn_pin_navigator.setText("📍")
        else:
            self._animate_navigator(0)
            self.btn_pin_navigator.setText("📌")

    def _animate_navigator(self, target_width):
        anim = QPropertyAnimation(self.navigator_panel, b"maximumWidth")
        anim.setDuration(180)
        anim.setStartValue(self.navigator_panel.maximumWidth() if self.navigator_panel.maximumWidth() < 16777215 else 0)
        anim.setEndValue(target_width)
        anim.setEasingCurve(QEasingCurve.Type.OutCubic)
        self.navigator_panel.setMinimumWidth(0)
        self.navigator_panel.setMaximumWidth(16777215)
        self._nav_anim = anim
        anim.start()

    def on_navigator_placeholder_enter(self, event):
        self._animate_navigator(280)


    def on_navigator_panel_leave(self, event):
        if self.navigator_pinned:
            return

        self._animate_navigator(0)

    def refresh_blocks_navigator(self):
        if not hasattr(self, "blocks_navigator"):
            return

        if getattr(self, "_refreshing_blocks_navigator", False):
            return

        self._refreshing_blocks_navigator = True

        try:
            self.blocks_navigator.clear()

            if not hasattr(self, "block_editor"):
                return

            for index, widget in enumerate(self.block_editor.current_blocks_widgets):
                try:
                    label = widget.build_summary()
                except Exception:
                    label = widget.block_type_label

                icon = self.block_navigator_icon(widget)
                block_id = widget.data.get("id", "")

                item = QListWidgetItem(f"{index + 1}. {icon} {label}")
                item.setData(Qt.UserRole, block_id)
                self.blocks_navigator.addItem(item)

        finally:
            self._refreshing_blocks_navigator = False


    def widget_from_navigator_item(self, item):
        block_id = item.data(Qt.UserRole)

        if not block_id:
            return None

        for widget in self.block_editor.current_blocks_widgets:
            if widget.data.get("id") == block_id:
                return widget

        return None


    def go_to_block_from_navigator(self, item):
        self.blocks_navigator.clearSelection()
        self.blocks_navigator.setCurrentItem(item)
        item.setSelected(True)

        widget = self.widget_from_navigator_item(item)

        if widget is None:
            return

        self._navigator_selected_widget = widget
        self._navigator_focus_token += 1
        current_token = self._navigator_focus_token

        for other_widget in self.block_editor.current_blocks_widgets:
            other_widget.set_collapsed_state(other_widget is not widget)

        self.block_editor.scroll_area.ensureWidgetVisible(
            widget,
            20,
            20
        )

        QTimer.singleShot(
            1200,
            lambda token=current_token: self.check_navigator_block_focus(token)
        )


    def focus_block_from_navigator(self, item):
        widget = self.widget_from_navigator_item(item)

        if widget is None:
            return

        self.go_to_block_from_navigator(item)

        if hasattr(widget, "editor"):
            widget.editor.setFocus()
        elif hasattr(widget, "latex_edit"):
            widget.latex_edit.setFocus()
        elif hasattr(widget, "table"):
            widget.table.setFocus()
        elif hasattr(widget, "title_edit"):
            widget.title_edit.setFocus()
        elif hasattr(widget, "caption_edit"):
            widget.caption_edit.setFocus()

    def check_navigator_block_focus(self, token):
        if token != getattr(self, "_navigator_focus_token", None):
            return
        widget = getattr(
            self,
            "_navigator_selected_widget",
            None
        )

        if widget is None:
            self.blocks_navigator.clearSelection()
            return

        focused = QApplication.focusWidget()

        if focused is None:
            self.blocks_navigator.clearSelection()
            return

        if widget.isAncestorOf(focused):
            return

        self.blocks_navigator.clearSelection()

    def on_blocks_navigator_order_changed(self, ordered_ids):
        id_to_widget = {
            widget.data.get("id"): widget
            for widget in self.block_editor.current_blocks_widgets
        }

        new_order = []

        for block_id in ordered_ids:
            widget = id_to_widget.get(block_id)

            if widget is not None:
                new_order.append(widget)

        if len(new_order) != len(self.block_editor.current_blocks_widgets):
            return

        self.block_editor.current_blocks_widgets = new_order

        for widget in new_order:
            self.block_editor.blocks_layout.removeWidget(widget)

        for index, widget in enumerate(new_order):
            self.block_editor.blocks_layout.insertWidget(index, widget)

        self.mark_dirty()
        self.update_figure_references()
        self.update_equation_references()
        self.refresh_blocks_navigator()

    def on_tree_selection_changed(self):
        self.save_current_item_content()
        items = self.tree.selectedItems()
        if not items:
            self.current_item = None
            self.selected_label.setText("Seleziona un nodo nella struttura")
            self.block_editor.clear_blocks()
            self.block_editor.set_enabled_for_node(False)
            self.refresh_blocks_navigator()
            return
        self.current_item = items[0]
        title = self.current_item.data(0, TITLE_ROLE) or self.current_item.text(0)
        blocks = self.ensure_blocks_defaults_list(self.current_item.data(0, CONTENT_ROLE) or [])
        self.current_item.setData(0, CONTENT_ROLE, blocks)
        self.selected_label.setText(f"Elemento selezionato: {title}")
        self.block_editor.load_blocks(deepcopy(blocks))
        self.block_editor.set_enabled_for_node(True)
        self.statusBar().showMessage(f"Selezionato: {title}", 2000)
        self.update_figure_references()
        self.refresh_blocks_navigator()

    def open_context_menu(self, position):
        item = self.tree.itemAt(position)
        if item is not None:
            self.save_current_item_content()
            self.tree.setCurrentItem(item)
        menu = QMenu(self)
        a1 = menu.addAction("Nuovo capitolo")
        a2 = menu.addAction("Nuovo paragrafo")
        menu.addSeparator()
        a3 = menu.addAction("Rinomina")
        a4 = menu.addAction("Cancella")
        menu.addSeparator()
        a_expand_tree = menu.addAction("Espandi tutti i nodi")
        a_collapse_tree = menu.addAction("Collassa tutti i nodi")
        menu.addSeparator()
        a5 = menu.addAction("Aggiorna numerazione")
        selected = menu.exec(self.tree.viewport().mapToGlobal(position))
        if selected == a1:
            self.add_chapter()
        elif selected == a2:
            self.add_paragraph()
        elif selected == a3:
            self.rename_selected_item()
        elif selected == a4:
            self.delete_selected_item()
        elif selected == a_expand_tree:
            self.tree.expandAll()
        elif selected == a_collapse_tree:
            self.tree.collapseAll()
        elif selected == a5:
            self.renumber_tree()
            self.mark_dirty()
            self.update_figure_references()
            self.statusBar().showMessage("Numerazione aggiornata", 2000)

    def rename_selected_item(self):
        item = self.tree.currentItem()
        if item is None:
            return
        current_title = item.data(0, TITLE_ROLE) or item.text(0)
        new_title, ok = QInputDialog.getText(self, "Rinomina", "Nuovo titolo:", text=current_title)
        if not ok or not new_title.strip():
            return
        item.setData(0, TITLE_ROLE, new_title.strip())
        self.renumber_tree()
        self.on_tree_selection_changed()
        self.mark_dirty()
        self.statusBar().showMessage("Elemento rinominato", 2000)

    def delete_selected_item(self):
        item = self.tree.currentItem()
        if item is None:
            return
        parent = item.parent()
        if parent is None:
            index = self.tree.indexOfTopLevelItem(item)
            self.tree.takeTopLevelItem(index)
        else:
            parent.removeChild(item)
        self.current_item = None
        self.selected_label.setText("Seleziona un nodo nella struttura")
        self.block_editor.clear_blocks()
        self.block_editor.set_enabled_for_node(False)
        self.renumber_tree()
        self.mark_dirty()
        self.update_figure_references()
        self.statusBar().showMessage("Elemento cancellato", 2000)

    def filter_tree(self, text):
        text = text.strip().lower()

        def filter_item(item):
            own_match = text in item.text(0).lower() if text else True
            child_match = False
            for i in range(item.childCount()):
                if filter_item(item.child(i)):
                    child_match = True
            visible = own_match or child_match
            item.setHidden(not visible)
            return visible

        for i in range(self.tree.topLevelItemCount()):
            filter_item(self.tree.topLevelItem(i))

    def collect_nodes_in_order(self):
        ordered = []

        def visit(item):
            ordered.append(item)
            for i in range(item.childCount()):
                visit(item.child(i))

        for i in range(self.tree.topLevelItemCount()):
            visit(self.tree.topLevelItem(i))
        return ordered

    def update_figure_references(self):
        self.save_current_item_content()
        figure_map = {}
        figure_number = 1
        for node in self.collect_nodes_in_order():
            blocks = self.ensure_blocks_defaults_list(node.data(0, CONTENT_ROLE) or [])
            node.setData(0, CONTENT_ROLE, blocks)
            for block in blocks:
                if block.get("type") == "image":
                    block_id = block.get("id")
                    figure_map[block_id] = {"number": figure_number, "label": f"Figura {figure_number}"}
                    figure_number += 1
                elif block.get("type") == "images":
                    images = self.ensure_group_images_defaults(block.get("images", []))
                    block["images"] = images
                    for img in images:
                        img_id = img.get("id")
                        figure_map[img_id] = {"number": figure_number, "label": f"Figura {figure_number}"}
                        figure_number += 1
        self.figure_reference_map = figure_map
        for widget in self.block_editor.iter_widgets():
            if isinstance(widget, (ImageBlockWidget, MultiImageBlockWidget)):
                widget.refresh_figure_info()

    def update_equation_references(self):
        self.save_current_item_content()

        equation_map = {}
        equation_number = 1

        for node in self.collect_nodes_in_order():
            blocks = self.ensure_blocks_defaults_list(
                node.data(0, CONTENT_ROLE) or []
            )

            node.setData(0, CONTENT_ROLE, blocks)

            for block in blocks:
                if block.get("type") != "equation":
                    continue

                block_id = block.get("id")
                if not block_id:
                    continue

                numbering_mode = block.get("numbering_mode", "none")

                if numbering_mode == "none":
                    continue

                equation_map[block_id] = {
                    "number": equation_number,
                    "label": f"Eq. ({equation_number})",
                    "numbering_mode": numbering_mode,
                    "caption": block.get("caption", "").strip(),
                }

                equation_number += 1

        self.equation_reference_map = equation_map

    def collect_attachment_references(self):
        pattern = re.compile(r"\[\[ALL:(.*?)\]\]")
        found = []

        self.save_current_item_content()

        def visit_item(item):
            blocks = self.ensure_blocks_defaults_list(item.data(0, CONTENT_ROLE) or [])

            for block in blocks:
                if block.get("type") == "text":
                    text = block.get("text", "") or ""
                    html = block.get("html", "") or ""
                    found.extend(pattern.findall(text))
                    found.extend(pattern.findall(html))

            for i in range(item.childCount()):
                visit_item(item.child(i))

        for i in range(self.tree.topLevelItemCount()):
            visit_item(self.tree.topLevelItem(i))

        unique = []
        for filename in found:
            filename = filename.strip()
            if filename and filename not in unique:
                unique.append(filename)

        return unique

    def build_attachment_reference_map(self):
        files = self.collect_attachment_references()

        return {
            filename: f"Allegato {idx}"
            for idx, filename in enumerate(files, start=1)
        }

    def build_attachment_list_for_export(self):
        attachment_map = self.build_attachment_reference_map()

        if not attachment_map:
            return []

        # restituisce lista ordinata tipo:
        # [("Allegato 1", "file.pdf"), ...]
        return sorted(
            [(label, filename) for filename, label in attachment_map.items()],
            key=lambda x: int(x[0].split()[-1])
        )

    def _add_docx_attachments_list(self, doc):
        attachments = self.build_attachment_list_for_export()

        if not attachments:
            return

        doc.add_page_break()

        title = doc.add_heading("Allegati", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for label, filename in attachments:
            p = doc.add_paragraph()

            run_label = p.add_run(f"{label} — ")
            run_label.bold = True

            # percorso assoluto file allegato
            allegati_dir = self.current_allegati_dir()
            file_path = os.path.join(allegati_dir, filename)

            # hyperlink
            part = doc.part
            r_id = part.relate_to(
                file_path,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True
            )

            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), r_id)

            new_run = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")

            # stile link (blu + underline)
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "0000FF")

            rPr.append(u)
            rPr.append(color)
            new_run.append(rPr)

            text = OxmlElement("w:t")
            text.text = filename
            new_run.append(text)

            hyperlink.append(new_run)
            p._p.append(hyperlink)

    def _add_pdf_attachments_list(self, story, title_style, body_style):
        attachments = self.build_attachment_list_for_export()

        if not attachments:
            return

        story.append(PageBreak())

        heading = Paragraph("Allegati", title_style)
        heading.toc_level = 0
        heading.toc_text = "Allegati"
        story.append(heading)

        story.append(Spacer(1, 0.25 * cm))

        for label, filename in attachments:
            allegati_dir = self.current_allegati_dir()
            file_path = os.path.join(allegati_dir, filename)

            # normalizza percorso per PDF
            file_path = file_path.replace("\\", "/")

            link = f"file:///{file_path}"

            text = (
                f"<b>{self._pdf_escape_inline(label)} — </b>"
                f'<link href="{link}">{self._pdf_escape_inline(filename)}</link>'
            )

            story.append(Paragraph(text, body_style))

    def render_attachment_references(self, text):
        attachment_map = self.build_attachment_reference_map()

        def repl(match):
            filename = match.group(1).strip()
            return attachment_map.get(filename, f"[Allegato non trovato: {filename}]")

        return re.sub(r"\[\[ALL:(.*?)\]\]", repl, text)

    def render_text_references(self, text):

        # === FIGURE ===

        def repl_fig(match):
            block_id = match.group(1).strip()

            info = self.figure_reference_map.get(block_id)

            return (
                info["label"]
                if info
                else f"[FIG:{block_id} non trovata]"
            )

        text = re.sub(
            r"\[\[FIG:([^\]]+)\]\]",
            repl_fig,
            text
        )

        # === EQUAZIONI ===

        def repl_eq(match):
            block_id = match.group(1).strip()

            info = self.equation_reference_map.get(block_id)

            if info:
                return info["label"]

            return f"[EQ:{block_id} non numerata o non trovata]"

        text = re.sub(
            r"\[\[EQ:([^\]]+)\]\]",
            repl_eq,
            text
        )

        # === ALLEGATI ===

        text = self.render_attachment_references(text)

        return text


    def verify_project_integrity(self):
        self.save_current_item_content()

        issues = []
        warnings = []
        referenced_text_files = set()
        referenced_image_files = set()
        checked_texts = 0
        checked_images = 0

        def visit(item):
            nonlocal checked_texts, checked_images
            blocks = self.ensure_blocks_defaults_list(item.data(0, CONTENT_ROLE) or [])
            for block in blocks:
                btype = block.get("type")
                if btype == "text":
                    block_id = block.get("id")
                    if not block_id:
                        issues.append(f"Nodo '{item.text(0)}': blocco testo senza ID.")
                    text_path = block.get("text_path", "")
                    if text_path:
                        rel = as_posix_path(text_path)
                        referenced_text_files.add(rel)
                        abs_path = self.resolve_runtime_path(text_path)
                        if not abs_path or not os.path.exists(abs_path):
                            issues.append(f"Nodo '{item.text(0)}': file testo mancante '{text_path}'.")
                    else:
                        warnings.append(f"Nodo '{item.text(0)}': blocco testo non ancora serializzato su file.")
                    checked_texts += 1
                elif btype == "image":
                    block_id = block.get("id")
                    if not block_id:
                        issues.append(f"Nodo '{item.text(0)}': blocco immagine senza ID.")
                    img_path = block.get("path", "")
                    if img_path:
                        rel = as_posix_path(img_path)
                        referenced_image_files.add(rel)
                        abs_path = self.resolve_runtime_path(img_path)
                        if not abs_path or not os.path.exists(abs_path):
                            issues.append(f"Nodo '{item.text(0)}': immagine mancante '{img_path}'.")
                    checked_images += 1
                elif btype == "images":
                    for idx, img in enumerate(self.ensure_group_images_defaults(block.get("images", [])), start=1):
                        img_id = img.get("id")
                        if not img_id:
                            issues.append(f"Nodo '{item.text(0)}': immagine {idx} del gruppo senza ID.")
                        img_path = img.get("path", "")
                        if img_path:
                            rel = as_posix_path(img_path)
                            referenced_image_files.add(rel)
                            abs_path = self.resolve_runtime_path(img_path)
                            if not abs_path or not os.path.exists(abs_path):
                                issues.append(f"Nodo '{item.text(0)}': immagine gruppo mancante '{img_path}'.")
                        checked_images += 1
            for i in range(item.childCount()):
                visit(item.child(i))

        for i in range(self.tree.topLevelItemCount()):
            visit(self.tree.topLevelItem(i))

        if self.project_dir:
            texts_dir = Path(self.project_dir) / "assets" / "texts"
            if texts_dir.exists():
                for file in texts_dir.glob("*.txt"):
                    rel = f"assets/texts/{file.name}"
                    if rel not in referenced_text_files:
                        warnings.append(f"File testo orfano: {rel}")
            images_dir = Path(self.project_dir) / "assets" / "images"
            if images_dir.exists():
                for file in images_dir.iterdir():
                    if file.is_file():
                        rel = f"assets/images/{file.name}"
                        if rel not in referenced_image_files:
                            warnings.append(f"Immagine orfana: {rel}")

        summary = [
            f"Blocchi testo verificati: {checked_texts}",
            f"Immagini verificate: {checked_images}",
        ]
        if self.project_dir:
            summary.append(f"Cartella progetto: {self.project_dir}")

        details = "\n".join(summary)
        if issues:
            details += "\n\nProblemi trovati:\n- " + "\n- ".join(issues)
        if warnings:
            details += "\n\nAvvisi:\n- " + "\n- ".join(warnings)

        if issues:
            QMessageBox.warning(self, "Verifica integrità progetto", details)
            self.statusBar().showMessage("Verifica completata: problemi trovati", 4000)
        elif warnings:
            QMessageBox.information(self, "Verifica integrità progetto", details)
            self.statusBar().showMessage("Verifica completata: solo avvisi", 4000)
        else:
            QMessageBox.information(
                self,
                "Verifica integrità progetto",
                details + "\n\nNessun problema rilevato."
            )
            self.statusBar().showMessage("Verifica completata: progetto integro", 4000)

    def normalize_blocks_for_save(self, blocks, source_project_dir=None):
        normalized = []
        for block in blocks:
            b = self.ensure_block_defaults(block)
            t = b.get("type")
            if t == "text":
                block_id = b.get("id")
                if block_id:
                    b["text_path"] = f"assets/texts/{block_id}.txt"
                b.pop("text", None)
            elif t == "image":
                if b.get("path"):
                    b["path"] = as_posix_path(b["path"])
            elif t == "images":
                imgs = []
                for img in self.ensure_group_images_defaults(b.get("images", [])):
                    img_copy = deepcopy(img)
                    if img_copy.get("path"):
                        img_copy["path"] = as_posix_path(img_copy["path"])
                    imgs.append(img_copy)
                b["images"] = imgs
            normalized.append(b)
        return normalized

    def serialize_item(self, item, source_project_dir=None):
        node = {
            "title": item.data(0, TITLE_ROLE),
            "content": self.normalize_blocks_for_save(item.data(0, CONTENT_ROLE) or [], source_project_dir),
            "children": [],
        }
        for i in range(item.childCount()):
            node["children"].append(self.serialize_item(item.child(i), source_project_dir))
        return node

    def deserialize_item(self, node):
        item = self.create_item(node.get("title", "Senza titolo"))
        blocks = self.ensure_blocks_defaults_list(node.get("content", []))
        loaded = []
        for block in blocks:
            if block.get("type") == "text":
                loaded.append(self.load_text_block_from_file(block))
            else:
                loaded.append(block)
        item.setData(0, CONTENT_ROLE, loaded)
        for child in node.get("children", []):
            item.addChild(self.deserialize_item(child))
        return item

    def _document_title_from_tree(self):
        return "RELAZIONE TECNICA"

    def _add_docx_cover(self, doc):
        cover_blocks = self.ensure_blocks_defaults_list(
            self.cover_config.get("blocks", [])
        )

        if not cover_blocks:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = p.add_run("RELAZIONE TECNICA")
            run.bold = True
            run.font.size = Pt(24)

            return

        for block in cover_blocks:
            self._export_block_to_docx(doc, block)

            doc.add_paragraph()

    def _add_pdf_cover(self, story, title_style, meta_style):
        cover_blocks = self.ensure_blocks_defaults_list(
            self.cover_config.get("blocks", [])
        )

        if not cover_blocks:
            story.append(Spacer(1, 8 * cm))
            story.append(Paragraph("RELAZIONE TECNICA", title_style))
            return

        body_style = ParagraphStyle(
            "CoverBody",
            parent=getSampleStyleSheet()["BodyText"],
            fontSize=11,
            leading=14,
            alignment=TA_CENTER,
            spaceAfter=8,
        )

        caption_style = ParagraphStyle(
            "CoverCaption",
            parent=getSampleStyleSheet()["BodyText"],
            fontSize=9,
            leading=11,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#444444"),
            spaceAfter=6,
        )

        for block in cover_blocks:
            self._export_block_to_pdf(
                story,
                block,
                body_style,
                caption_style
            )
            story.append(Spacer(1, 0.25 * cm))

    def _add_docx_toc_page(self, doc):
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run("Sommario")
        title_run.bold = True
        title_run.font.size = Pt(16)

        p = doc.add_paragraph()
        run = p.add_run()

        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_begin)

        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '
        run._r.append(instr)

        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')
        run._r.append(fld_sep)

        txt = OxmlElement('w:t')
        txt.text = "Aggiorna il campo in Word con F9 per visualizzare il sommario."
        run._r.append(txt)

        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_end)

    def _add_pdf_toc_page(self, story, title_style, toc):
        story.append(Paragraph("Sommario", title_style))
        story.append(Spacer(1, 0.6 * cm))
        story.append(toc)

    def _set_paragraph_border(self, paragraph, *, top=False, bottom=False):
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)

        def add_edge(tag_name):
            edge = p_bdr.find(qn(f"w:{tag_name}"))
            if edge is None:
                edge = OxmlElement(f"w:{tag_name}")
                p_bdr.append(edge)
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), "6")
            edge.set(qn("w:space"), "1")
            edge.set(qn("w:color"), "000000")

        if top:
            add_edge("top")
        if bottom:
            add_edge("bottom")

    def _add_page_number_run(self, paragraph):
        run = paragraph.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run._r.append(instr)

        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_sep)

        text_node = OxmlElement("w:t")
        text_node.text = "1"
        run._r.append(text_node)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)

    def _restart_word_page_numbering(self, section, start=1):
        sect_pr = section._sectPr
        pg_num_type = sect_pr.find(qn("w:pgNumType"))

        if pg_num_type is None:
            pg_num_type = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num_type)

        pg_num_type.set(qn("w:start"), str(start))

    def _continue_word_page_numbering(self, section):
        sect_pr = section._sectPr
        pg_num_type = sect_pr.find(qn("w:pgNumType"))

        if pg_num_type is not None:
            sect_pr.remove(pg_num_type)

    def _configure_noncontent_section_word(self, section):
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        self._apply_layout_to_word_header(section, "cover", chapter_title="")
        self._apply_layout_to_word_footer(section, "cover", show_page_number=False)
    
    def _configure_toc_section_word(self, section):
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        self._apply_layout_to_word_header(
            section,
            "content",
            chapter_title=""
        )

        self._apply_layout_to_word_footer(
            section,
            "content",
            show_page_number=False
        )

    def _configure_blank_section_word(self, section):
        section.different_first_page_header_footer = False

        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        for header in (
            section.header,
            section.first_page_header,
            section.even_page_header,
        ):
            header.is_linked_to_previous = False
            for p in header.paragraphs:
                p.clear()

        for footer in (
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            footer.is_linked_to_previous = False
            for p in footer.paragraphs:
                p.clear()

    def _configure_content_section_word(self, section, chapter_title=""):
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        self._apply_layout_to_word_header(section, "content", chapter_title=chapter_title)
        self._apply_layout_to_word_footer(section, "content", show_page_number=True)

    def _extract_base_title(self, numbered_title: str) -> str:
        if not numbered_title:
            return ""
        parts = numbered_title.split(". ", 1)
        if len(parts) == 2 and parts[0].replace(".", "").isdigit():
            return parts[1]
        return numbered_title

    def _chapter_header_text(self, item) -> str:
        title = item.data(0, TITLE_ROLE) or item.text(0)
        return self._extract_base_title(title)
    
    def _get_layout_entry(self, section_name, part_name):
        return (
            self.layout_config.get(section_name, {})
            .get(part_name, {"type": "none", "value": ""})
        )

    def _get_layout_alignment(self, section_name, part_name, default="center"):
        return (
            self.layout_config.get(section_name, {})
            .get(part_name, {})
            .get("align", default)
        )

    def _alignment_to_word(self, align_value):
        if align_value == "left":
            return WD_ALIGN_PARAGRAPH.LEFT
        if align_value == "center":
            return WD_ALIGN_PARAGRAPH.CENTER
        return WD_ALIGN_PARAGRAPH.RIGHT

    def _apply_layout_to_word_header(self, section, section_name, chapter_title=""):
        cfg = self._get_layout_entry(section_name, "header")
        alignment = self._alignment_to_word(
            self._get_layout_alignment(section_name, "header", default="right")
        )

        header = section.header
        header.is_linked_to_previous = False

        # pulizia header
        while len(header.paragraphs) > 0:
            p = header.paragraphs[0]
            p._element.getparent().remove(p._element)

        has_custom = False

        # ===== BLOCCO CUSTOM =====
        if cfg["type"] == "text":
            html = cfg.get("html", "")
            value = cfg.get("value", "").strip()

            if html or value:
                p = header.add_paragraph()
                p.alignment = alignment
                p.paragraph_format.space_after = Pt(3)

                if html:
                    qdoc = QTextDocument()
                    qdoc.setHtml(html)

                    block = qdoc.begin()
                    first_block = True

                    while block.isValid():
                        if not first_block:
                            p.add_run().add_break()

                        it = block.begin()
                        while not it.atEnd():
                            fragment = it.fragment()
                            if fragment.isValid():
                                txt = fragment.text()
                                if txt:
                                    char_fmt = fragment.charFormat()
                                    run = p.add_run(txt)
                                    run.bold = char_fmt.fontWeight() >= 700
                                    run.italic = char_fmt.fontItalic()
                                    run.underline = char_fmt.fontUnderline()

                                    size = char_fmt.fontPointSize()
                                    if size > 0:
                                        run.font.size = Pt(size)
                            it += 1

                        first_block = False
                        block = block.next()

                else:
                    p.add_run(value)

                has_custom = True

        elif cfg["type"] == "image" and cfg["value"] and os.path.exists(cfg["value"]):
            p = header.add_paragraph()
            p.alignment = alignment
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run()
            run.add_picture(cfg["value"], width=Inches(1.8))
            has_custom = True

        # ===== TITOLO CAPITOLO =====
        p = header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(2 if has_custom else 0)
        p.paragraph_format.space_after = Pt(0)

        if chapter_title:
            p.add_run(chapter_title)

        self._set_paragraph_border(p, bottom=True)

        # ===== RIGA VUOTA =====
        blank_p = header.add_paragraph()
        blank_p.paragraph_format.space_before = Pt(0)
        blank_p.paragraph_format.space_after = Pt(0)
        blank_p.add_run("")

    def _apply_layout_to_word_footer(self, section, section_name, show_page_number):
        cfg = self._get_layout_entry(section_name, "footer")
        alignment = self._alignment_to_word(
            self._get_layout_alignment(section_name, "footer", default="center")
        )

        footer = section.footer
        footer.is_linked_to_previous = False

        # pulizia paragrafi esistenti
        while len(footer.paragraphs) > 0:
            p = footer.paragraphs[0]
            p._element.getparent().remove(p._element)

        # unico paragrafo con filetto + eventuale contenuto personalizzato
        p = footer.add_paragraph()
        p.alignment = alignment
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        self._set_paragraph_border(p, top=True)

        if cfg["type"] == "text":
            html = cfg.get("html", "")
            value = cfg.get("value", "").strip()

            if html:
                qdoc = QTextDocument()
                qdoc.setHtml(html)

                block = qdoc.begin()
                first_block = True

                while block.isValid():
                    if not first_block:
                        p.add_run().add_break()

                    it = block.begin()
                    while not it.atEnd():
                        fragment = it.fragment()
                        if fragment.isValid():
                            txt = fragment.text()
                            if txt:
                                char_fmt = fragment.charFormat()
                                run = p.add_run(txt)
                                run.bold = char_fmt.fontWeight() >= 700
                                run.italic = char_fmt.fontItalic()
                                run.underline = char_fmt.fontUnderline()

                                size = char_fmt.fontPointSize()
                                if size > 0:
                                    run.font.size = Pt(size)
                        it += 1

                    first_block = False
                    block = block.next()

            elif value:
                for idx, line in enumerate(value.splitlines()):
                    line = line.rstrip()
                    if not line:
                        continue
                    if idx > 0:
                        p.add_run().add_break()
                    p.add_run(line)

        elif cfg["type"] == "image" and cfg["value"] and os.path.exists(cfg["value"]):
            run = p.add_run()
            run.add_picture(cfg["value"], width=Inches(1.8))

        # numero pagina sempre sull'ultima riga
        if show_page_number:
            page_p = footer.add_paragraph()
            page_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            page_p.paragraph_format.space_before = Pt(0)
            page_p.paragraph_format.space_after = Pt(0)
            page_p.add_run("- ")
            self._add_page_number_run(page_p)
            page_p.add_run(" -")

    def export_document(self):
        self.save_current_item_content()
        self.update_figure_references()
        self.update_equation_references()

        fmt, ok = QInputDialog.getItem(self, "Esporta documento", "Formato:", ["Word (.docx)", "PDF (.pdf)"], 0, False)
        if not ok:
            return

        if fmt.startswith("Word"):
            start = os.path.join(self.project_dir or "", "documento.docx")
            file_path, _ = QFileDialog.getSaveFileName(self, "Esporta in Word", start, "Word (*.docx)")
            if not file_path:
                return
            try:
                self.export_to_docx(file_path)
                self.statusBar().showMessage("Esportazione Word completata", 3000)
                QMessageBox.information(self, "Esportazione completata", "Il documento Word è stato generato correttamente.")
            except Exception as exc:
                QMessageBox.critical(self, "Errore export Word", f"{exc}\n\n{traceback.format_exc()}")
        else:
            start = os.path.join(self.project_dir or "", "documento.pdf")
            file_path, _ = QFileDialog.getSaveFileName(self, "Esporta in PDF", start, "PDF (*.pdf)")
            if not file_path:
                return
            try:
                self.export_to_pdf(file_path)
                self.statusBar().showMessage("Esportazione PDF completata", 3000)
                QMessageBox.information(self, "Esportazione completata", "Il PDF è stato generato correttamente.")
            except Exception as exc:
                QMessageBox.critical(self, "Errore export PDF", f"{exc}\n\n{traceback.format_exc()}")

    def iter_nodes_for_export(self):
        def visit(item, level):
            yield item, level
            for i in range(item.childCount()):
                yield from visit(item.child(i), level + 1)

        for i in range(self.tree.topLevelItemCount()):
            yield from visit(self.tree.topLevelItem(i), 1)

    def _set_doc_table_borders(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = tblPr.first_child_found_in("w:tblBorders")
        if tblBorders is None:
            tblBorders = OxmlElement("w:tblBorders")
            tblPr.append(tblBorders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = tblBorders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tblBorders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "6")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "808080")


    def _set_doc_table_no_borders(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = tblPr.first_child_found_in("w:tblBorders")
        if tblBorders is None:
            tblBorders = OxmlElement("w:tblBorders")
            tblPr.append(tblBorders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = tblBorders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tblBorders.append(element)
            element.set(qn("w:val"), "nil")

    def _set_row_cant_split(self, row):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)

    def _estimate_caption_height_inches(self, caption_text, usable_width_inches):
        if not caption_text:
            return 0.0
        approx_chars_per_line = max(25, int(usable_width_inches * 11))
        lines = max(1, math.ceil(len(caption_text) / approx_chars_per_line))
        return 0.22 * lines + 0.08

    def _compute_word_figure_width(self, path, section, caption_text=""):
        pix = QPixmap(path)
        if pix.isNull() or pix.height() <= 0 or pix.width() <= 0:
            return Inches(5.8)

        page_w_in = section.page_width.inches
        page_h_in = section.page_height.inches
        usable_w_in = max(1.0, page_w_in - section.left_margin.inches - section.right_margin.inches)
        usable_h_in = max(1.0, page_h_in - section.top_margin.inches - section.bottom_margin.inches)

        header_footer_reserve_in = 0.9
        caption_reserve_in = self._estimate_caption_height_inches(caption_text, usable_w_in)
        figure_max_h_in = max(1.0, usable_h_in - header_footer_reserve_in - caption_reserve_in - 0.15)

        aspect = pix.width() / pix.height()
        width_from_height = figure_max_h_in * aspect
        final_w_in = min(usable_w_in, width_from_height)

        min_target_w_in = min(usable_w_in * 0.70, usable_w_in)
        if min_target_w_in / aspect <= figure_max_h_in:
            final_w_in = max(final_w_in, min_target_w_in)

        final_w_in = min(final_w_in, usable_w_in)
        final_w_in = max(1.0, final_w_in)
        return Inches(final_w_in)


    def _word_image_aspect(self, path):
        pix = QPixmap(path)
        if pix.isNull() or pix.height() <= 0 or pix.width() <= 0:
            return 1.0
        return pix.width() / pix.height()

    def _word_page_usable_metrics(self, section):
        usable_w_in = max(1.0, section.page_width.inches - section.left_margin.inches - section.right_margin.inches)
        usable_h_in = max(1.0, section.page_height.inches - section.top_margin.inches - section.bottom_margin.inches)
        header_footer_reserve_in = 0.9
        return usable_w_in, max(1.0, usable_h_in - header_footer_reserve_in)

    def _word_single_figure_total_height(self, path, section, caption_text=""):
        width = self._compute_word_figure_width(path, section, caption_text).inches
        aspect = self._word_image_aspect(path)
        image_h = width / aspect if aspect > 0 else width
        caption_h = self._estimate_caption_height_inches(caption_text, width)
        return image_h + caption_h + 0.20

    def _figure_caption_text(self, block_id, caption):
        info = self.figure_reference_map.get(block_id)
        label = info["label"] if info else "Figura"
        caption = (caption or "").strip()
        return f"{label} — {caption}" if caption else (label if info else "")

    def _can_pair_word_image_blocks(self, doc, block1, block2):
        if block1.get("type") != "image" or block2.get("type") != "image":
            return False

        path1 = self.resolve_runtime_path(block1.get("path", ""))
        path2 = self.resolve_runtime_path(block2.get("path", ""))
        if not path1 or not path2 or not os.path.exists(path1) or not os.path.exists(path2):
            return False

        section = doc.sections[-1]
        usable_w_in, usable_h_in = self._word_page_usable_metrics(section)

        cap1 = self._figure_caption_text(block1.get("id"), block1.get("caption", ""))
        cap2 = self._figure_caption_text(block2.get("id"), block2.get("caption", ""))

        h1 = self._word_single_figure_total_height(path1, section, cap1)
        h2 = self._word_single_figure_total_height(path2, section, cap2)

        if h1 >= 0.60 * usable_h_in:
            return False

        aspect1 = self._word_image_aspect(path1)
        aspect2 = self._word_image_aspect(path2)

        # Avoid pairing two very wide panoramas
        if aspect1 > 2.6 and aspect2 > 2.6:
            return False

        caps_h = self._estimate_caption_height_inches(cap1, usable_w_in) + self._estimate_caption_height_inches(cap2, usable_w_in)
        spacing_h = 0.30
        available_fig_h = usable_h_in - caps_h - spacing_h
        if available_fig_h <= 1.0:
            return False

        natural_h1 = usable_w_in / aspect1 if aspect1 > 0 else usable_w_in
        natural_h2 = usable_w_in / aspect2 if aspect2 > 0 else usable_w_in

        # Keep the smaller figure at full width, shrink the larger one if needed.
        if natural_h1 <= natural_h2:
            used_h1 = min(natural_h1, available_fig_h)
            remaining = available_fig_h - used_h1
            if remaining <= 0.9:
                return False
            used_h2 = min(natural_h2, remaining)
        else:
            used_h2 = min(natural_h2, available_fig_h)
            remaining = available_fig_h - used_h2
            if remaining <= 0.9:
                return False
            used_h1 = min(natural_h1, remaining)

        total = used_h1 + used_h2 + caps_h + spacing_h
        return total <= usable_h_in + 1e-6

    def _compute_word_pair_widths(self, doc, path1, caption1, path2, caption2):
        section = doc.sections[-1]
        usable_w_in, usable_h_in = self._word_page_usable_metrics(section)
        aspect1 = self._word_image_aspect(path1)
        aspect2 = self._word_image_aspect(path2)

        caps_h = self._estimate_caption_height_inches(caption1, usable_w_in) + self._estimate_caption_height_inches(caption2, usable_w_in)
        spacing_h = 0.30
        available_fig_h = max(1.0, usable_h_in - caps_h - spacing_h)

        natural_h1 = usable_w_in / aspect1 if aspect1 > 0 else usable_w_in
        natural_h2 = usable_w_in / aspect2 if aspect2 > 0 else usable_w_in

        if natural_h1 <= natural_h2:
            h1 = min(natural_h1, available_fig_h)
            h2 = min(natural_h2, max(1.0, available_fig_h - h1))
        else:
            h2 = min(natural_h2, available_fig_h)
            h1 = min(natural_h1, max(1.0, available_fig_h - h2))

        w1 = min(usable_w_in, h1 * aspect1)
        w2 = min(usable_w_in, h2 * aspect2)
        return Inches(max(1.0, w1)), Inches(max(1.0, w2))

    def _add_two_figure_blocks_to_docx(self, doc, block1, block2):
        path1 = self.resolve_runtime_path(block1.get("path", ""))
        path2 = self.resolve_runtime_path(block2.get("path", ""))
        caption1 = self._figure_caption_text(block1.get("id"), block1.get("caption", ""))
        caption2 = self._figure_caption_text(block2.get("id"), block2.get("caption", ""))

        width1, width2 = self._compute_word_pair_widths(doc, path1, caption1, path2, caption2)

        table = doc.add_table(rows=2, cols=1)
        table.autofit = True
        self._set_doc_table_no_borders(table)
        self._set_row_cant_split(table.rows[0])
        self._set_row_cant_split(table.rows[1])

        for row, path, caption_text, width in [
            (table.rows[0], path1, caption1, width1),
            (table.rows[1], path2, caption2, width2),
        ]:
            cell = row.cells[0]
            image_p = cell.paragraphs[0]
            image_p.paragraph_format.keep_with_next = True
            image_p.paragraph_format.keep_together = True
            self._add_picture_to_paragraph(image_p, path, width=width)

            if caption_text:
                cap_p = cell.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.paragraph_format.keep_together = True
                run = cap_p.add_run(caption_text)
                run.italic = True
                run.font.size = Pt(9)

        doc.add_paragraph()

    def _add_picture_to_paragraph(self, paragraph, path, width):
        if not path or not os.path.exists(path):
            return False
        try:
            run = paragraph.add_run()
            run.add_picture(path, width=width)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        except Exception:
            pass

        try:
            img = QPixmap(path)
            if img.isNull():
                raise ValueError("Immagine non caricabile")
            fd, tmp_path = tempfile.mkstemp(suffix=".png")
            os.close(fd)
            if not img.save(tmp_path, "PNG"):
                raise ValueError("Impossibile salvare immagine temporanea")
            run = paragraph.add_run()
            run.add_picture(tmp_path, width=width)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        except Exception:
            paragraph.add_run(f"[Immagine non esportabile: {os.path.basename(path)}]")
            return False
        finally:
            try:
                if "tmp_path" in locals() and os.path.exists(tmp_path):
                    os.remove(tmp_path)
            except Exception:
                pass

    def _add_figure_block_to_docx(self, doc, path, caption_text="", scale_percent=100):
        section = doc.sections[-1]
        figure_width = self._compute_word_figure_width(path, section, caption_text)

        base_width = figure_width

        scale_factor = max(10, min(200, int(scale_percent or 100))) / 100.0
        figure_width = base_width * scale_factor

        max_width = section.page_width - section.left_margin - section.right_margin

        if figure_width > max_width:
            figure_width = max_width

        table = doc.add_table(rows=1, cols=1)
        table.autofit = True
        self._set_doc_table_no_borders(table)
        self._set_row_cant_split(table.rows[0])

        cell = table.cell(0, 0)
        image_p = cell.paragraphs[0]
        image_p.paragraph_format.keep_with_next = True
        image_p.paragraph_format.keep_together = True
        self._add_picture_to_paragraph(image_p, path, width=figure_width)

        if caption_text:
            cap_p = cell.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.keep_together = True
            run = cap_p.add_run(caption_text)
            run.italic = True
            run.font.size = Pt(9)

        doc.add_paragraph()

    def _add_docx_caption(self, doc, text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = False
        p.paragraph_format.keep_together = True
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(9)

    def _safe_add_picture_to_docx(self, doc, path, width=Inches(5.8)):
        if not path or not os.path.exists(path):
            return
        try:
            doc.add_picture(path, width=width)
            if doc.paragraphs:
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
        except Exception:
            pass

        try:
            img = QPixmap(path)
            if img.isNull():
                raise ValueError("Immagine non caricabile")
            fd, tmp_path = tempfile.mkstemp(suffix=".png")
            os.close(fd)
            if not img.save(tmp_path, "PNG"):
                raise ValueError("Impossibile salvare immagine temporanea")
            doc.add_picture(tmp_path, width=width)
            if doc.paragraphs:
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            doc.add_paragraph(f"[Immagine non esportabile: {os.path.basename(path)}]")
        finally:
            try:
                if 'tmp_path' in locals() and os.path.exists(tmp_path):
                    os.remove(tmp_path)
            except Exception:
                pass

    def export_to_docx(self, file_path):
        doc = Document()
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)

        self._configure_noncontent_section_word(doc.sections[0])

        self._add_docx_cover(doc)

        blank_after_cover = doc.add_section(WD_SECTION.NEW_PAGE)
        self._configure_blank_section_word(blank_after_cover)

        toc_section = doc.add_section(WD_SECTION.NEW_PAGE)
        self._configure_toc_section_word(toc_section)

        self._add_docx_toc_page(doc)

        blank_after_toc = doc.add_section(WD_SECTION.NEW_PAGE)
        self._configure_blank_section_word(blank_after_toc)

        first_content_section = True

        for item, level in self.iter_nodes_for_export():
            if level == 1:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                self._configure_content_section_word(section, chapter_title=self._chapter_header_text(item))

                if first_content_section:
                    self._restart_word_page_numbering(section, 1)
                    first_content_section = False
                else:
                    self._continue_word_page_numbering(section)

            doc.add_heading(item.text(0), level=min(level, 3))
            blocks = self.ensure_blocks_defaults_list(item.data(0, CONTENT_ROLE) or [])

            i = 0
            while i < len(blocks):
                block = blocks[i]
                if i + 1 < len(blocks) and self._can_pair_word_image_blocks(doc, block, blocks[i + 1]):
                    self._add_two_figure_blocks_to_docx(doc, block, blocks[i + 1])
                    i += 2
                    continue

                self._export_block_to_docx(doc, block)
                i += 1

            doc.add_paragraph()

        self._add_docx_attachments_list(doc)

        try:
            doc.save(file_path)
        except PermissionError:
            QMessageBox.warning(
                self,
                "File in uso",
                "Chiudi il documento Word aperto prima di esportare."
            )
            return

    def _export_rich_text_to_docx(self, doc, html):
        qdoc = QTextDocument()
        qdoc.setHtml(html or "")

        block = qdoc.begin()

        while block.isValid():
            text_list = block.textList()
            block_alignment = block.blockFormat().alignment()

            if text_list:
                fmt = text_list.format()
                indent_level = max(1, fmt.indent())

                if fmt.style() == QTextListFormat.ListDecimal:
                    p = doc.add_paragraph(style="List Number")
                else:
                    p = doc.add_paragraph(style="List Bullet")

                p.paragraph_format.left_indent = Pt(18 * indent_level)
                p.paragraph_format.first_line_indent = Pt(-9)
            else:
                p = doc.add_paragraph()
                #p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            it = block.begin()
            while not it.atEnd():
                fragment = it.fragment()
                if fragment.isValid():
                    frag_text = self.render_text_references(fragment.text())
                    if frag_text:
                        char_fmt = fragment.charFormat()
                        run = p.add_run(frag_text)
                        run.bold = char_fmt.fontWeight() >= 700
                        run.italic = char_fmt.fontItalic()
                        run.underline = char_fmt.fontUnderline()

                        vertical = char_fmt.verticalAlignment()
                        if vertical == QTextCharFormat.AlignSuperScript:
                            run.font.superscript = True
                        elif vertical == QTextCharFormat.AlignSubScript:
                            run.font.subscript = True

                        font_size = char_fmt.fontPointSize()
                        if font_size > 0:
                            if vertical in (
                                QTextCharFormat.AlignSuperScript,
                                QTextCharFormat.AlignSubScript,
                            ):
                                run.font.size = Pt(font_size * 0.75)
                            else:
                                run.font.size = Pt(font_size)

                it += 1

            # --- ALLINEAMENTO PARAGRAFO ---
            text_content = p.text.strip()

            if block_alignment & Qt.AlignHCenter:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif block_alignment & Qt.AlignRight:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif block_alignment & Qt.AlignJustify:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif len(text_content) < 60:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            block = block.next()

    def _pdf_hyphenate_text(self, text):
        if not text:
            return text

        words = text.split(" ")
        result = []

        protected_patterns = [
            r"/",                 # unità composte: kN/m², N/mm², C25/30
            r"\\",                # eventuali percorsi o simboli tecnici
            r"\d",                # numeri o mix lettere/numeri: C25, S235
            r"²|³",               # unità già formattate
            r"^[A-ZÀ-Ý]{2,}$",    # sigle: SLU, SLE, NTC
        ]

        for w in words:
            clean = w.strip(".,;:()[]{}")

            protect = False
            for pattern in protected_patterns:
                if re.search(pattern, clean):
                    protect = True
                    break

            if protect:
                result.append(w)
                continue

            if len(clean) < 9:
                result.append(w)
                continue

            hyp_clean = self.hyphenator.inserted(clean, hyphen="\u00AD")

            if clean != w:
                hyp_word = w.replace(clean, hyp_clean, 1)
            else:
                hyp_word = hyp_clean

            result.append(hyp_word)

        return " ".join(result)

    def _pdf_escape_inline_with_unicode_super_sub(self, text):
        supers = {
            "⁰": "0", "¹": "1", "²": "2", "³": "3", "⁴": "4",
            "⁵": "5", "⁶": "6", "⁷": "7", "⁸": "8", "⁹": "9",
            "⁺": "+", "⁻": "-", "⁼": "=", "⁽": "(", "⁾": ")",
        }

        subs = {
            "₀": "0", "₁": "1", "₂": "2", "₃": "3", "₄": "4",
            "₅": "5", "₆": "6", "₇": "7", "₈": "8", "₉": "9",
            "₊": "+", "₋": "-", "₌": "=", "₍": "(", "₎": ")",
        }

        out = []

        for ch in str(text):
            if ch in supers:
                out.append(f"<super>{supers[ch]}</super>")
            elif ch in subs:
                out.append(f"<sub>{subs[ch]}</sub>")
            else:
                out.append(self._pdf_escape_inline(ch))

        return "".join(out)

    def _pdf_escape_inline(self, text):
        return (
            str(text)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\t", "&nbsp;&nbsp;&nbsp;&nbsp;")
            .replace("    ", "&nbsp;&nbsp;&nbsp;&nbsp;")
        )

    def _export_rich_text_to_pdf(self, story, html, body_style):
        qdoc = QTextDocument()
        qdoc.setHtml(html or "")

        block = qdoc.begin()

        while block.isValid():
            text_list = block.textList()
            block_alignment = block.blockFormat().alignment()

            prefix = ""
            indent_level = 0

            if text_list:
                fmt = text_list.format()
                indent_level = max(1, fmt.indent())
                number = text_list.itemNumber(block) + 1

                if fmt.style() == QTextListFormat.ListDecimal:
                    prefix = f"{number}. "
                else:
                    prefix = "• "

            parts = []

            if prefix:
                parts.append(self._pdf_escape_inline(prefix))

            it = block.begin()
            while not it.atEnd():
                fragment = it.fragment()
                if fragment.isValid():
                    frag_text = self.render_text_references(fragment.text())
                    frag_text = self._pdf_hyphenate_text(frag_text)
                    if frag_text:
                        char_fmt = fragment.charFormat()
                        leading_spaces = len(frag_text) - len(frag_text.lstrip(" "))
                        trailing_spaces = len(frag_text) - len(frag_text.rstrip(" "))

                        leading = "&nbsp;" * leading_spaces
                        trailing = "&nbsp;" * trailing_spaces

                        core_text = frag_text.strip(" ")
                        txt = self._pdf_escape_inline_with_unicode_super_sub(core_text)

                        if char_fmt.fontWeight() >= 700:
                            txt = f"<b>{txt}</b>"
                        if char_fmt.fontItalic():
                            txt = f"<i>{txt}</i>"
                        if char_fmt.fontUnderline():
                            txt = f"<u>{txt}</u>"

                        txt = leading + txt + trailing

                        font_size = char_fmt.fontPointSize()
                        vertical = char_fmt.verticalAlignment()

                        if vertical == QTextCharFormat.AlignSuperScript:
                            txt = f"<super>{txt}</super>"
                        elif vertical == QTextCharFormat.AlignSubScript:
                            txt = f"<sub>{txt}</sub>"

                        if font_size > 0:
                            if vertical in (
                                QTextCharFormat.AlignSuperScript,
                                QTextCharFormat.AlignSubScript,
                            ):
                                txt = f'<font size="{font_size * 0.75}">{txt}</font>'
                            else:
                                txt = f'<font size="{font_size}">{txt}</font>'

                        parts.append(txt)

                it += 1

            if parts:
                paragraph_text = "".join(parts)
                plain_text = re.sub(r"<[^>]+>", "", paragraph_text).strip()

                if block_alignment & Qt.AlignHCenter:
                    paragraph_alignment = TA_CENTER
                elif block_alignment & Qt.AlignRight:
                    paragraph_alignment = TA_RIGHT
                elif block_alignment & Qt.AlignJustify:
                    paragraph_alignment = TA_JUSTIFY
                elif len(plain_text) < 60:
                    paragraph_alignment = TA_LEFT
                else:
                    paragraph_alignment = TA_JUSTIFY

                if indent_level > 0:
                    final_style = ParagraphStyle(
                        f"BodyListIndent_{indent_level}_{paragraph_alignment}",
                        parent=body_style,
                        leftIndent=18 * indent_level,
                        firstLineIndent=-9,
                        alignment=paragraph_alignment,
                    )
                else:
                    final_style = ParagraphStyle(
                        f"BodyAligned_{paragraph_alignment}",
                        parent=body_style,
                        alignment=paragraph_alignment,
                    )

                story.append(Paragraph(paragraph_text, final_style))

            block = block.next()
    
    def render_equation_to_png(self, latex, block_id="equation"):
        raw_text = (latex or "").strip()

        if not raw_text:
            return ""

        lines = [
            line.strip()
            for line in raw_text.splitlines()
            if line.strip()
        ]

        if not lines:
            return ""

        equations_dir = os.path.join(tempfile.gettempdir(), "EditorTecnico_equations")
        os.makedirs(equations_dir, exist_ok=True)

        safe_id = block_id or new_block_id()
        file_path = os.path.join(equations_dir, f"{safe_id}.png")

        try:
            line_count = len(lines)

            fig_height = max(0.30, 0.30 * line_count)

            fig = Figure(figsize=(4.2, fig_height), dpi=220)
            fig.patch.set_facecolor("white")

            canvas = FigureCanvasAgg(fig)

            ax = fig.add_subplot(111)
            fig.subplots_adjust(left=0, right=1, top=1, bottom=0)
            ax.axis("off")

            for idx, line in enumerate(lines):
                latex_line = line

                if not (latex_line.startswith("$") and latex_line.endswith("$")):
                    latex_line = f"${latex_line}$"

                y = 1.0 - ((idx + 0.5) / line_count)

                ax.text(
                    0.5,
                    y,
                    latex_line,
                    fontsize=10,
                    ha="center",
                    va="center",
                    color="black",
                    transform=ax.transAxes,
                )

            canvas.draw()

            fig.savefig(
                file_path,
                dpi=220,
                facecolor="white",
                edgecolor="white",
                transparent=False,
                bbox_inches="tight",
                pad_inches=0,
            )

            return file_path if os.path.exists(file_path) else ""

        except Exception as exc:
            QMessageBox.warning(
                self,
                "Errore equazione",
                f"Impossibile renderizzare l'equazione:\n\n{raw_text}\n\n{exc}"
            )
            return ""

    def _export_block_to_docx(self, doc, block):
        btype = block.get("type")

        if btype == "text":
            html = block.get("html", "")
            if html:
                self._export_rich_text_to_docx(doc, html)
            else:
                rendered = self.render_text_references(block.get("text", ""))
                for part in rendered.splitlines():
                    if part.strip():
                        p = doc.add_paragraph(part)
                        if len(part.strip()) < 60:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        elif btype == "image":
            path = self.resolve_runtime_path(block.get("path", ""))
            info = self.figure_reference_map.get(block.get("id"))
            caption = block.get("caption", "").strip()
            label = info["label"] if info else "Figura"
            caption_text = f"{label} — {caption}" if caption else (label if info else "")

            scale_percent = int(block.get("scale_percent", 100) or 100)

            self._add_figure_block_to_docx(
                doc,
                path,
                caption_text,
                scale_percent=scale_percent
            )

        elif btype == "equation":
            latex = block.get("latex", "").strip()
            caption = block.get("caption", "").strip()
            numbering_mode = block.get("numbering_mode", "none")

            eq_info = self.equation_reference_map.get(block.get("id"))
            eq_label = eq_info["label"] if eq_info else ""

            path = self.render_equation_to_png(
                latex,
                block.get("id", new_block_id())
            )

            if not path or not os.path.exists(path):
                doc.add_paragraph(f"[Equazione non esportabile: {latex}]")
                return

            if numbering_mode == "number" and eq_label:
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False

                left_cell = table.cell(0, 0)
                right_cell = table.cell(0, 1)

                left_cell.width = Inches(5.6)
                right_cell.width = Inches(1.0)

                left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                p_eq = left_cell.paragraphs[0]
                p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_picture_to_paragraph(p_eq, path, width=Inches(4.8))

                p_num = right_cell.paragraphs[0]
                p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r_num = p_num.add_run(eq_label)
                r_num.font.size = Pt(10)

                for row in table.rows:
                    for cell in row.cells:
                        tc_pr = cell._tc.get_or_add_tcPr()
                        tc_borders = OxmlElement("w:tcBorders")

                        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                            border = OxmlElement(f"w:{border_name}")
                            border.set(qn("w:val"), "nil")
                            tc_borders.append(border)

                        tc_pr.append(tc_borders)

            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_picture_to_paragraph(p, path, width=Inches(4.8))

                if numbering_mode == "number_caption" and eq_label and caption:
                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    run = cap_p.add_run(f"{eq_label} – {caption}")
                    run.italic = True
                    run.font.size = Pt(9)

        elif btype == "images":
            title = block.get("title", "").strip()
            if title:
                p = doc.add_paragraph()
                r = p.add_run(title)
                r.bold = True

            for img in self.ensure_group_images_defaults(block.get("images", [])):
                path = self.resolve_runtime_path(img.get("path", ""))
                info = self.figure_reference_map.get(img.get("id"))
                caption = img.get("caption", "").strip()
                label = info["label"] if info else "Figura"
                caption_text = f"{label} — {caption}" if caption else (label if info else "")
                self._add_figure_block_to_docx(doc, path, caption_text)

        elif btype == "table":
            data = block.get("data", [])
            rows = max(1, block.get("rows", len(data) or 1))
            cols = max(1, block.get("cols", len(data[0]) if data else 1))
            spans = block.get("spans", [])
            cell_formats = block.get("cell_formats", {})

            table = doc.add_table(rows=rows, cols=cols)
            table.style = "Table Grid"
            table.autofit = True
            self._set_doc_table_borders(table)

            for row in table.rows:
                for cell in row.cells:
                    cell.width = Inches(6.3 / cols)
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)

            merged_cells = set()

            for span in spans:
                r1 = span.get("r1", 0)
                c1 = span.get("c1", 0)
                r2 = span.get("r2", r1)
                c2 = span.get("c2", c1)

                if (
                    0 <= r1 < rows and
                    0 <= r2 < rows and
                    0 <= c1 < cols and
                    0 <= c2 < cols and
                    (r2 > r1 or c2 > c1)
                ):
                    try:
                        table.cell(r1, c1).merge(table.cell(r2, c2))

                        for rr in range(r1, r2 + 1):
                            for cc in range(c1, c2 + 1):
                                if not (rr == r1 and cc == c1):
                                    merged_cells.add((rr, cc))

                    except Exception:
                        pass

            for r in range(rows):
                for c in range(cols):
                    if (r, c) in merged_cells:
                        continue

                    value = ""
                    if r < len(data) and c < len(data[r]):
                        value = str(data[r][c])

                    cell = table.cell(r, c)
                    cell.text = ""

                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)

                    fmt = cell_formats.get(f"{r},{c}", {})
                    halign = fmt.get("halign", "left")
                    valign = fmt.get("valign", "middle")

                    if halign == "center":
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif halign == "right":
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif halign == "justify":
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    if valign == "top":
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    elif valign == "bottom":
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
                    else:
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                    run = p.add_run(value)
                    run.font.size = Pt(9)

            doc.add_paragraph()

    def _scaled_image_for_pdf(self, path, max_w=16*cm, max_h=12*cm, scale_percent=100):
        img = RLImage(path)
        iw, ih = img.imageWidth, img.imageHeight

        if iw <= 0 or ih <= 0:
            return img

        base_scale = min(max_w / iw, max_h / ih, 1.0)
        user_scale = max(10, min(200, int(scale_percent or 100))) / 100.0

        final_scale = base_scale * user_scale

        draw_w = iw * final_scale
        draw_h = ih * final_scale

        if draw_w > max_w:
            ratio = max_w / draw_w
            draw_w *= ratio
            draw_h *= ratio

        if draw_h > max_h:
            ratio = max_h / draw_h
            draw_w *= ratio
            draw_h *= ratio

        img.drawWidth = draw_w
        img.drawHeight = draw_h
        img.hAlign = "CENTER"

        return img

    def _pdf_layout_text_height_points(self, cfg):
        if not cfg or cfg.get("type") != "text":
            return 0
        value = str(cfg.get("value", "")).strip()
        if not value:
            return 0
        lines = [line for line in value.splitlines() if line.strip()]
        return len(lines) * 10


    def _pdf_layout_image_height_points(self, cfg, max_height_cm=2.0):
        if not cfg or cfg.get("type") != "image":
            return 0
        value = cfg.get("value", "")
        if not value or not os.path.exists(value):
            return 0

        pix = QPixmap(value)
        if pix.isNull() or pix.height() <= 0:
            return 0

        max_h = max_height_cm * cm
        return min(max_h, pix.height())


    def _pdf_layout_block_height_points(self, cfg):
        if not cfg:
            return 0
        if cfg.get("type") == "text":
            return self._pdf_layout_text_height_points(cfg)
        if cfg.get("type") == "image":
            return self._pdf_layout_image_height_points(cfg)
        return 0


    def _compute_pdf_dynamic_margins(self):
        base_top = 2.6 * cm
        base_bottom = 2.2 * cm

        cover_header = self.layout_config.get("cover", {}).get("header", {})
        content_header = self.layout_config.get("content", {}).get("header", {})
        cover_footer = self.layout_config.get("cover", {}).get("footer", {})
        content_footer = self.layout_config.get("content", {}).get("footer", {})

        max_header_h = max(
            self._pdf_layout_block_height_points(cover_header),
            self._pdf_layout_block_height_points(content_header),
        )

        max_footer_h = max(
            self._pdf_layout_block_height_points(cover_footer),
            self._pdf_layout_block_height_points(content_footer),
        )

        top_margin = base_top + max_header_h
        bottom_margin = base_bottom + max_footer_h

        return top_margin, bottom_margin

    def export_to_pdf(self, file_path):
        styles = getSampleStyleSheet()
        title1 = ParagraphStyle("Title1", parent=styles["Heading1"], spaceAfter=10)
        title2 = ParagraphStyle("Title2", parent=styles["Heading2"], spaceAfter=8)
        title3 = ParagraphStyle("Title3", parent=styles["Heading3"], spaceAfter=6)
        cover_title = ParagraphStyle("CoverTitle", parent=styles["Title"], alignment=TA_CENTER, fontSize=22, spaceAfter=18)
        cover_meta = ParagraphStyle("CoverMeta", parent=styles["BodyText"], alignment=TA_CENTER, fontSize=10, italic=True, spaceAfter=6)
        caption_style = ParagraphStyle("FigureCaption", parent=styles["BodyText"], alignment=TA_CENTER, fontSize=9, italic=True, spaceAfter=8)
        body = ParagraphStyle(
            "BodyAligned",
            parent=styles["BodyText"],
            spaceAfter=4,
            alignment=TA_JUSTIFY
        )

        toc = TableOfContents()
        toc.levelStyles = [
            ParagraphStyle("TOC1", parent=styles["BodyText"], leftIndent=0, rightIndent=20, firstLineIndent=0, spaceAfter=4),
            ParagraphStyle("TOC2", parent=styles["BodyText"], leftIndent=12, rightIndent=20, firstLineIndent=0, spaceAfter=3),
            ParagraphStyle("TOC3", parent=styles["BodyText"], leftIndent=24, rightIndent=20, firstLineIndent=0, spaceAfter=2),
        ]
        toc.dotsMinLevel = 0

        story = []
        dynamic_top_margin, dynamic_bottom_margin = self._compute_pdf_dynamic_margins()

        doc = TechnicalDocTemplate(
            file_path,
            owner=self,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=dynamic_top_margin,
            bottomMargin=dynamic_bottom_margin,
        )

        doc.pdf_blank_pages = {2}

        self._add_pdf_cover(story, cover_title, cover_meta)

        # Pagina bianca fissa sul retro della copertina
        story.append(PageBreak())
        story.append(Spacer(1, 1))
        story.append(PageBreak())

        self._add_pdf_toc_page(story, title1, toc)

        # Gestione dinamica:
        # se il sommario termina su pagina dispari -> 1 bianca
        # se termina su pagina pari -> 2 bianche
        story.append(DynamicTocEndBreak())

        first_chapter = True

        for item, level in self.iter_nodes_for_export():
            if level == 1 and not first_chapter:
                story.append(PageBreak())
            if level == 1:
                first_chapter = False

            style = title1 if level == 1 else title2 if level == 2 else title3
            heading = Paragraph(item.text(0), style)
            heading.toc_level = max(0, min(level - 1, 2))
            heading.toc_text = item.text(0)
            story.append(heading)

            story.append(Spacer(1, 0.15 * cm))
            blocks = self.ensure_blocks_defaults_list(item.data(0, CONTENT_ROLE) or [])
            for block in blocks:
                self._export_block_to_pdf(story, block, body, caption_style)
                story.append(Spacer(1, 0.12 * cm))

        self._add_pdf_attachments_list(story, title1, body)
        doc.multiBuild(story)

    def _export_block_to_pdf(self, story, block, body_style, caption_style):
        def esc(s):
            return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        btype = block.get("type")

        if btype == "text":
            html = block.get("html", "")
            text = block.get("text", "")

            if html:
                plain_lines = text.splitlines()

                top_blank = 0
                bottom_blank = 0

                for line in plain_lines:
                    if line.strip():
                        break
                    top_blank += 1

                for line in reversed(plain_lines):
                    if line.strip():
                        break
                    bottom_blank += 1

                if top_blank > 0:
                    story.append(Spacer(1, top_blank * 0.35 * cm))

                self._export_rich_text_to_pdf(story, html, body_style)

                if bottom_blank > 0:
                    story.append(Spacer(1, bottom_blank * 0.35 * cm))

            else:
                rendered = self.render_text_references(text)
                for part in rendered.splitlines():
                    if part.strip():
                        story.append(Paragraph(esc(part), body_style))
                    else:
                        story.append(Spacer(1, 0.35 * cm))

        elif btype == "image":
            path = self.resolve_runtime_path(block.get("path", ""))
            if path and os.path.exists(path):
                story.append(
                    self._scaled_image_for_pdf(
                        path,
                        scale_percent=block.get("scale_percent", 100)
                    )
                )

            info = self.figure_reference_map.get(block.get("id"))
            caption = block.get("caption", "").strip()
            label = info["label"] if info else "Figura"

            if caption:
                story.append(Paragraph(esc(f"{label} — {caption}"), caption_style))
            elif info:
                story.append(Paragraph(esc(label), caption_style))

        elif btype == "equation":
            latex = block.get("latex", "").strip()
            caption = block.get("caption", "").strip()
            numbering_mode = block.get("numbering_mode", "none")

            eq_info = self.equation_reference_map.get(block.get("id"))
            eq_label = eq_info["label"] if eq_info else ""

            path = self.render_equation_to_png(
                latex,
                block.get("id", new_block_id())
            )

            if not path or not os.path.exists(path):
                story.append(
                    Paragraph(
                        esc(f"[Equazione non esportabile: {latex}]"),
                        body_style
                    )
                )
                return

            equation_image = self._scaled_image_for_pdf(path)

            if numbering_mode == "number" and eq_label:
                number_style = ParagraphStyle(
                    "EquationNumberRight",
                    parent=body_style,
                    alignment=TA_RIGHT,
                    fontSize=10,
                    leading=12,
                    spaceAfter=0,
                    spaceBefore=0,
                )

                table = RLTable(
                    [
                        [
                            equation_image,
                            Paragraph(esc(eq_label), number_style),
                        ]
                    ],
                    colWidths=[14.2 * cm, 2.0 * cm],
                )

                table.setStyle(
                    TableStyle([
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("ALIGN", (0, 0), (0, 0), "CENTER"),
                        ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                        ("TOPPADDING", (0, 0), (-1, -1), 0),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                    ])
                )

                story.append(table)

            else:
                story.append(equation_image)

                if numbering_mode == "number_caption" and eq_label and caption:
                    story.append(
                        Paragraph(
                            esc(f"{eq_label} – {caption}"),
                            caption_style
                        )
                    )

        elif btype == "images":
            title = block.get("title", "").strip()
            if title:
                story.append(Paragraph(f"<b>{esc(title)}</b>", body_style))

            for img in self.ensure_group_images_defaults(block.get("images", [])):
                path = self.resolve_runtime_path(img.get("path", ""))
                if path and os.path.exists(path):
                    story.append(self._scaled_image_for_pdf(path))

                info = self.figure_reference_map.get(img.get("id"))
                caption = img.get("caption", "").strip()
                label = info["label"] if info else "Figura"

                if caption:
                    story.append(Paragraph(esc(f"{label} — {caption}"), caption_style))
                elif info:
                    story.append(Paragraph(esc(label), caption_style))

        elif btype == "table":
            data = block.get("data", [])
            rows = max(1, block.get("rows", len(data) or 1))
            cols = max(1, block.get("cols", len(data[0]) if data else 1))
            cell_formats = block.get("cell_formats", {})
            spans = block.get("spans", [])

            matrix = []

            for r in range(rows):
                row = []
                for c in range(cols):
                    value = ""
                    if r < len(data) and c < len(data[r]):
                        value = str(data[r][c])

                    fmt = cell_formats.get(f"{r},{c}", {})
                    halign = fmt.get("halign", "left")

                    if halign == "center":
                        cell_alignment = TA_CENTER
                    elif halign == "right":
                        cell_alignment = TA_RIGHT
                    elif halign == "justify":
                        cell_alignment = TA_JUSTIFY
                    else:
                        cell_alignment = TA_LEFT

                    cell_style = ParagraphStyle(
                        f"TableCell_{r}_{c}",
                        parent=body_style,
                        alignment=cell_alignment,
                        fontSize=9,
                        leading=11,
                        spaceAfter=0,
                        spaceBefore=0,
                    )

                    value = esc(value).replace("\n", "<br/>")
                    row.append(Paragraph(value, cell_style))
                matrix.append(row)

            available_width = A4[0] - 4 * cm
            col_width = available_width / cols

            row_height = 0.75 * cm

            table = RLTable(
                matrix,
                colWidths=[col_width] * cols,
                minRowHeights=[row_height] * rows,
                hAlign="LEFT"
            )

            table_styles = [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("LEADING", (0, 0), (-1, -1), 11),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
            ]
            for span in spans:
                r1 = span.get("r1", 0)
                c1 = span.get("c1", 0)
                r2 = span.get("r2", r1)
                c2 = span.get("c2", c1)

                if (
                    0 <= r1 < rows and
                    0 <= r2 < rows and
                    0 <= c1 < cols and
                    0 <= c2 < cols and
                    (r2 > r1 or c2 > c1)
                ):
                    table_styles.append(("SPAN", (c1, r1), (c2, r2)))

            for r in range(rows):
                for c in range(cols):
                    fmt = cell_formats.get(f"{r},{c}", {})
                    halign = fmt.get("halign", "left")
                    valign = fmt.get("valign", "middle")

                    # orizzontale
                    if halign == "center":
                        table_styles.append(("ALIGN", (c, r), (c, r), "CENTER"))
                    elif halign == "right":
                        table_styles.append(("ALIGN", (c, r), (c, r), "RIGHT"))
                    else:
                        table_styles.append(("ALIGN", (c, r), (c, r), "LEFT"))

                    # verticale
                    if valign == "top":
                        table_styles.append(("VALIGN", (c, r), (c, r), "TOP"))
                    elif valign == "bottom":
                        table_styles.append(("VALIGN", (c, r), (c, r), "BOTTOM"))
                    else:
                        table_styles.append(("VALIGN", (c, r), (c, r), "MIDDLE"))

            table.setStyle(TableStyle(table_styles))
            story.append(table)

    def new_project(self):
        if not self.workspace_ready:
            QMessageBox.warning(
                self,
                "Workspace mancante",
                "Seleziona prima un workspace."
            )
            return

        if not self.maybe_save_before_destructive_action():
            return

        dialog = NewRelationDialog(self.workspace_dir, self)

        if dialog.exec() != QDialog.Accepted:
            return

        data = dialog.values()

        relation_dir = os.path.join(
            self.workspace_dir,
            "ARCHIVIO LAVORI",
            data["cliente"],
            data["commessa"],
            "Relazioni tecniche",
            data["relazione"]
        )

        allegati_dir = os.path.join(relation_dir, "Allegati")

        if os.path.exists(relation_dir):
            QMessageBox.warning(
                self,
                "Relazione già esistente",
                f"Esiste già una relazione con questo nome.\n\n"
                f"Percorso:\n{relation_dir}\n\n"
                "Scegli un nome diverso."
            )
            return

        os.makedirs(relation_dir, exist_ok=True)
        os.makedirs(allegati_dir, exist_ok=True)

        self.set_project_dir(relation_dir)
        self.populate_empty_structure()
        self.project_dirty = True
        self.update_window_title()
        self.update_window_title()
        self.save_project()
        self.update_project_info_labels()

        self.statusBar().showMessage(
            f"Nuova relazione creata: {relation_dir}",
            4000
        )

    def save_project(self):
        if self.project_dir is None:
            QMessageBox.warning(
                self,
                "Nessuna relazione aperta",
                "Apri prima una relazione."
            )
            return False
        self.save_current_item_content()
        os.makedirs(os.path.join(self.project_dir, "assets", "texts"), exist_ok=True)
        os.makedirs(os.path.join(self.project_dir, "assets", "images"), exist_ok=True)
        self.sync_all_text_blocks_to_files()
        self.sync_all_images_to_assets(self.source_project_dir_for_save or self.project_dir)
        data = {
            "format_version": 2,
            "nodes": [],
            "layout": self.layout_config,
            "cover": self.cover_config,
        }
        for i in range(self.tree.topLevelItemCount()):
            data["nodes"].append(self.serialize_item(self.tree.topLevelItem(i), self.source_project_dir_for_save or self.project_dir))
        with open(self.document_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        self.prune_unused_asset_files()

        recovery_path = os.path.join(
            self.project_dir,
            "document.recovery.json"
        )

        if os.path.exists(recovery_path):
            try:
                os.remove(recovery_path)
            except Exception:
                pass

        self.project_dirty = False
        self.update_window_title()
        self.update_figure_references()
        self.statusBar().showMessage("Progetto salvato", 2000)
        return True

    def save_project_as_new_relation(self):
        if not self.workspace_ready:
            QMessageBox.warning(
                self,
                "Workspace mancante",
                "Seleziona prima un workspace."
            )
            return

        if not self.project_dir or not os.path.exists(self.project_dir):
            QMessageBox.warning(
                self,
                "Nessuna relazione",
                "Apri o crea prima una relazione da salvare con nome."
            )
            return

        dialog = SaveAsRelationDialog(self.workspace_dir, self)

        if dialog.exec() != QDialog.Accepted:
            return

        data = dialog.values()

        target_relation_dir = os.path.join(
            self.workspace_dir,
            "ARCHIVIO LAVORI",
            data["cliente"],
            data["commessa"],
            "Relazioni tecniche",
            data["relazione"]
        )

        if os.path.exists(target_relation_dir):
            QMessageBox.warning(
                self,
                "Relazione già esistente",
                (
                    "Esiste già una relazione con questo nome:\n\n"
                    f"{target_relation_dir}"
                )
            )
            return

        try:
            self.save_project()

            shutil.copytree(self.project_dir, target_relation_dir)

            new_document_path = os.path.join(
                target_relation_dir,
                os.path.basename(self.document_path)
            )

            self.set_project_dir(target_relation_dir)
            self.document_path = new_document_path

            self.save_project()

            self.project_dirty = False
            self.update_window_title()
            self.update_project_info_labels()

            QMessageBox.information(
                self,
                "Salva con Nome",
                "Nuova relazione creata correttamente."
            )

        except Exception as exc:
            QMessageBox.critical(
                self,
                "Errore",
                f"Impossibile salvare la relazione:\n\n{exc}"
            )

    def unique_system_archive_dir(self, parent_dir, base_name):
        """
        Restituisce una cartella libera dentro l'archivio di sistema.

        Se base_name non esiste, restituisce:
            parent_dir/base_name

        Se esiste già, restituisce:
            parent_dir/base_name REV. 01
            parent_dir/base_name REV. 02
            ecc.
        """

        candidate = os.path.join(parent_dir, base_name)

        if not os.path.exists(candidate):
            return candidate

        rev = 1

        while True:
            candidate = os.path.join(
                parent_dir,
                f"{base_name} REV. {rev:02d}"
            )

            if not os.path.exists(candidate):
                return candidate

            rev += 1

    def save_project_copy_to_system_archive(self):
        if not self.workspace_ready:
            QMessageBox.warning(
                self,
                "Workspace mancante",
                "Seleziona prima un workspace."
            )
            return

        if not self.project_dir or not os.path.exists(self.project_dir):
            QMessageBox.warning(
                self,
                "Nessuna relazione",
                "Apri o crea prima una relazione da copiare."
            )
            return

        if not self.save_project():
            return

        dialog = QDialog(self)
        dialog.setWindowTitle("Salva in archivio di sistema")
        dialog.resize(460, 220)

        layout = QVBoxLayout(dialog)

        layout.addWidget(QLabel("Nome copia/template:"))

        name_combo = QComboBox()
        name_combo.setEditable(True)
        name_combo.setInsertPolicy(QComboBox.NoInsert)
        name_combo.setPlaceholderText("Seleziona o inserisci nome")
        layout.addWidget(name_combo)

        def archive_name_source_dir():
            return os.path.join(
                self.workspace_dir,
                "ARCHIVIO DI SISTEMA",
                "TEMPLATE",
                "RELAZIONI TECNICHE"
            )


        layout.addWidget(QLabel("Percorso di salvataggio:"))

        archive_path_label = QLabel("")
        archive_path_label.setWordWrap(True)
        archive_path_label.setStyleSheet(
            "color: #404040; background: #f5f5f5; "
            "border: 1px solid #d0d0d0; padding: 6px;"
        )
        layout.addWidget(archive_path_label)


        def refresh_archive_path_label():
            archive_path_label.setText(
                archive_name_source_dir()
            )


        def refresh_name_combo():
            current_text = name_combo.currentText().strip()

            source_dir = archive_name_source_dir()
            refresh_archive_path_label()

            existing_names = []

            if os.path.exists(source_dir):
                existing_names = [
                    name for name in os.listdir(source_dir)
                    if os.path.isdir(os.path.join(source_dir, name))
                ]

            name_combo.blockSignals(True)
            name_combo.clear()
            name_combo.addItems(sorted(existing_names))

            if existing_names:
                name_combo.setCurrentIndex(0)
            elif current_text:
                name_combo.setCurrentText(current_text)
            else:
                name_combo.setCurrentText("Template relazione tecnica")

            name_combo.blockSignals(False)

        refresh_name_combo()

        btn_row = QHBoxLayout()
        btn_ok = QPushButton("Salva")
        btn_cancel = QPushButton("Annulla")

        btn_ok.clicked.connect(dialog.accept)
        btn_cancel.clicked.connect(dialog.reject)

        btn_row.addStretch()
        btn_row.addWidget(btn_ok)
        btn_row.addWidget(btn_cancel)
        layout.addStretch()
        layout.addLayout(btn_row)

        if dialog.exec() != QDialog.Accepted:
            return

        copy_name = name_combo.currentText().strip()

        valid, message = validate_folder_name(copy_name)
        if not valid:
            QMessageBox.warning(
                self,
                "Nome non valido",
                message
            )
            return

        archive_root = os.path.join(
            self.workspace_dir,
            "ARCHIVIO DI SISTEMA",
            "TEMPLATE",
            "RELAZIONI TECNICHE"
        )

        success_message = "Template relazione salvato nell'archivio di sistema"

        os.makedirs(archive_root, exist_ok=True)

        destination_dir = self.unique_system_archive_dir(
            archive_root,
            copy_name
        )

        try:
            shutil.copytree(self.project_dir, destination_dir)
        except Exception as exc:
            QMessageBox.critical(
                self,
                "Errore copia",
                f"Impossibile creare la copia in archivio:\n{exc}"
            )
            return

        self.statusBar().showMessage(success_message, 4000)

        QMessageBox.information(
            self,
            "Copia salvata",
            f"La copia è stata salvata in:\n\n{destination_dir}"
        )

    def create_system_archive_document(self):
        if not self.workspace_ready:
            QMessageBox.warning(
                self,
                "Workspace mancante",
                "Seleziona prima un workspace."
            )
            return

        if not self.maybe_save_before_destructive_action():
            return

        dialog = QDialog(self)
        dialog.setWindowTitle("Crea documento archivio")
        dialog.resize(460, 240)

        layout = QVBoxLayout(dialog)

        layout.addWidget(QLabel("Nome documento archivio:"))

        name_combo = QComboBox()
        name_combo.setEditable(True)
        name_combo.setInsertPolicy(QComboBox.NoInsert)
        name_combo.setPlaceholderText("Inserisci nome documento")
        layout.addWidget(name_combo)

        def archive_root_dir():
            return os.path.join(
                self.workspace_dir,
                "ARCHIVIO DI SISTEMA",
                "TEMPLATE",
                "RELAZIONI TECNICHE"
            )

        def default_document_name():
            return "Nuovo template relazione"

        def refresh_name_combo(force_default=False):
            current_text = name_combo.currentText().strip()
            source_dir = archive_root_dir()

            existing_names = []

            if os.path.exists(source_dir):
                existing_names = [
                    name for name in os.listdir(source_dir)
                    if os.path.isdir(os.path.join(source_dir, name))
                ]

            name_combo.blockSignals(True)
            name_combo.clear()
            name_combo.addItems(sorted(existing_names))

            if force_default:
                name_combo.setEditText(default_document_name())
            elif current_text:
                name_combo.setEditText(current_text)
            else:
                name_combo.setEditText(default_document_name())

            name_combo.blockSignals(False)

        btn_row = QHBoxLayout()
        btn_ok = QPushButton("Crea")
        btn_cancel = QPushButton("Annulla")

        btn_cancel.clicked.connect(dialog.reject)
        btn_ok.clicked.connect(dialog.accept)

        btn_row.addStretch()
        btn_row.addWidget(btn_ok)
        btn_row.addWidget(btn_cancel)

        layout.addStretch()
        layout.addLayout(btn_row)

        refresh_name_combo()

        if dialog.exec() != QDialog.Accepted:
            return

        document_name = name_combo.currentText().strip()

        valid, message = validate_folder_name(document_name)
        if not valid:
            QMessageBox.warning(
                self,
                "Nome non valido",
                message
            )
            return

        archive_root = archive_root_dir()
        os.makedirs(archive_root, exist_ok=True)

        destination_dir = self.unique_system_archive_dir(
            archive_root,
            document_name
        )

        try:
            os.makedirs(destination_dir, exist_ok=False)
            os.makedirs(os.path.join(destination_dir, "assets", "images"), exist_ok=True)
            os.makedirs(os.path.join(destination_dir, "assets", "texts"), exist_ok=True)

            document_path = os.path.join(destination_dir, "document.json")

            data = {
                "format_version": 2,
                "nodes": [],
                "layout": self.layout_config,
            }

            with open(document_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

            self.set_project_dir(destination_dir)
            self.document_path = document_path
            self.populate_empty_structure()
            self.project_dirty = False
            self.update_window_title()
            self.update_project_info_labels()

        except Exception as exc:
            QMessageBox.critical(
                self,
                "Errore creazione documento",
                f"Impossibile creare il documento archivio:\n\n{exc}"
            )
            return

        QMessageBox.information(
            self,
            "Documento archivio creato",
            f"Il documento è stato creato in:\n\n{destination_dir}"
        )

        self.statusBar().showMessage("Documento archivio creato", 4000)

    def open_system_archive_document(self):
        if not self.workspace_dir:
            QMessageBox.warning(
                self,
                "Workspace mancante",
                "Seleziona prima un workspace."
            )
            return

        dialog = QDialog(self)
        dialog.setWindowTitle("Apri documento archivio")
        dialog.resize(520, 160)

        layout = QVBoxLayout(dialog)

        # -----------------------------
        # Documento
        # -----------------------------

        layout.addWidget(QLabel("Documento archivio:"))

        name_combo = QComboBox()
        name_combo.setEditable(False)

        layout.addWidget(name_combo)

        # -----------------------------
        # Path helper
        # -----------------------------

        def archive_source_dir():
            return os.path.join(
                self.workspace_dir,
                "ARCHIVIO DI SISTEMA",
                "TEMPLATE",
                "RELAZIONI TECNICHE"
            )

        # -----------------------------
        # Refresh combo
        # -----------------------------

        def refresh_name_combo():
            name_combo.clear()

            source_dir = archive_source_dir()

            if not os.path.exists(source_dir):
                return

            names = [
                name for name in os.listdir(source_dir)
                if os.path.isdir(
                    os.path.join(source_dir, name)
                )
            ]

            name_combo.addItems(sorted(names))

        # -----------------------------
        # Buttons
        # -----------------------------

        buttons_row = QHBoxLayout()
        buttons_row.addStretch()

        btn_open = QPushButton("Apri")
        btn_cancel = QPushButton("Annulla")

        buttons_row.addWidget(btn_open)
        buttons_row.addWidget(btn_cancel)

        layout.addStretch()
        layout.addLayout(buttons_row)

        btn_cancel.clicked.connect(dialog.reject)

        # -----------------------------
        # Open logic
        # -----------------------------

        def open_selected_document():
            name = name_combo.currentText().strip()

            if not name:
                QMessageBox.warning(
                    dialog,
                    "Documento mancante",
                    "Seleziona un documento."
                )
                return

            relation_dir = os.path.join(
                archive_source_dir(),
                name
            )

            document_path = os.path.join(
                relation_dir,
                "document.json"
            )

            if not os.path.exists(document_path):
                QMessageBox.warning(
                    dialog,
                    "Documento non trovato",
                    "Il file document.json non esiste."
                )
                return

            try:
                with open(document_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
            except Exception as exc:
                QMessageBox.critical(
                    dialog,
                    "Errore",
                    f"Impossibile aprire il documento:\n{exc}"
                )
                return

            try:
                self.set_project_dir(os.path.dirname(document_path))
                self.document_path = document_path
                self.populate_empty_structure()
                self.layout_config = data.get("layout", self.layout_config)
                self.cover_config = data.get("cover", {"blocks": []})

                for node in data.get("nodes", []):
                    self.tree.addTopLevelItem(
                        self.deserialize_item(node)
                    )

                self.renumber_tree()
                self.project_dirty = False
                self.update_window_title()
                self.update_figure_references()
                self.update_project_info_labels()

                self.statusBar().showMessage(
                    "Documento archivio aperto",
                    2000
                )

            except Exception as exc:
                QMessageBox.critical(
                    dialog,
                    "Errore apertura documento",
                    f"{exc}\n\n{traceback.format_exc()}"
                )
                return

            dialog.accept()

        btn_open.clicked.connect(
            open_selected_document
        )

        # -----------------------------
        # Init
        # -----------------------------

        refresh_name_combo()

        dialog.exec()

    def auto_save_project(self):
        if not self.project_dir:
            return

        if not self.project_dirty:
            return

        try:
            self.save_project()
            self.statusBar().showMessage("Salvataggio automatico completato", 2000)
        except Exception:
            self.statusBar().showMessage("Errore durante il salvataggio automatico", 4000)

    def update_window_title(self):
        title = "Editor Tecnico (designed by Ing. Luca Monardi)"

        if self.project_dir:
            relazione = os.path.basename(self.project_dir)
            title += f" - {relazione}"

        if getattr(self, "project_dirty", False):
            #title += "  (ATTENZIONE! MODIFICHE NON SALVATE)"
            title += " - ⚠ MODIFICHE NON SALVATE - "

        self.setWindowTitle(title)

    def save_project_as(self):
        start_path = os.path.join(self.project_dir, "document.json") if self.project_dir else ""
        file_path, _ = QFileDialog.getSaveFileName(self, "Salva con nome", start_path, "File JSON (*.json)")
        if not file_path:
            return False
        new_project_dir = os.path.dirname(file_path)
        old_project_dir = self.project_dir
        if old_project_dir and os.path.normpath(new_project_dir) != os.path.normpath(old_project_dir):
            choice = QMessageBox.question(
                self,
                "Cambio cartella progetto",
                "Stai cambiando la cartella progetto. Il nuovo percorso diventerà la cartella progetto della sessione. Vuoi procedere?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No,
            )
            if choice != QMessageBox.Yes:
                return False
        self.set_project_dir(new_project_dir)
        self.document_path = file_path
        self.source_project_dir_for_save = old_project_dir
        try:
            return self.save_project()
        finally:
            self.source_project_dir_for_save = None
        
        self.project_dirty = False
        self.update_window_title()

    def open_project(self):
        if not self.workspace_ready:
            QMessageBox.warning(
                self,
                "Workspace mancante",
                "Seleziona prima un workspace."
            )
            return

        if not self.maybe_save_before_destructive_action():
            return

        dialog = OpenRelationDialog(self.workspace_dir, self)

        if dialog.exec() != QDialog.Accepted:
            return

        json_path = dialog.selected_document_path()

        try:
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception as exc:
            QMessageBox.critical(self, "Errore", f"Impossibile aprire la relazione:\n{exc}")
            return

        try:
            self.set_project_dir(os.path.dirname(json_path))
            self.document_path = json_path
            self.populate_empty_structure()
            self.layout_config = data.get("layout", self.layout_config)
            self.cover_config = data.get("cover", self.cover_config)

            for node in data.get("nodes", []):
                self.tree.addTopLevelItem(self.deserialize_item(node))

            self.renumber_tree()
            self.project_dirty = False
            self.update_window_title()
            self.update_figure_references()
            self.update_project_info_labels()

            self.statusBar().showMessage("Relazione aperta", 2000)

        except Exception as exc:
            QMessageBox.critical(
                self,
                "Errore apertura relazione",
                f"{exc}\n\n{traceback.format_exc()}"
            )
            return

    def closeEvent(self, event):
        self.save_current_item_content()
        if not self.maybe_save_before_destructive_action():
            event.ignore()
            return
        super().closeEvent(event)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
